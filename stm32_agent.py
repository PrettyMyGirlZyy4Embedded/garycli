#!/usr/bin/env python3
"""
Gary Dev Agent
===============
融合 Gary（编译/烧录/调试闭环）与 ClaudeTerminal（对话 UI + 工具框架）的 STM32 专属 AI 助手。

硬件后端：pyocd（比 OpenOCD 更易用，支持 ST-Link / CMSIS-DAP / J-Link USB 探针）
AI 前端  ：流式对话 + 函数调用工具链

功能：
  - 自然语言 → 生成完整 STM32 HAL C 代码
  - 一键编译（arm-none-eabi-gcc）
  - 一键烧录（pyocd）
  - 读寄存器 / 串口监控 / HardFault 分析
  - 对话式修改历史项目代码
  - 完整自动调试闭环

用法：
  python3 stm32_agent.py               # 启动（不连接硬件）
  python3 stm32_agent.py --connect     # 启动并自动连接第一个可用探针
  python3 stm32_agent.py --chip STM32F407VET6   # 指定芯片
"""

import atexit
import signal
import sys, os, json, re, time, shutil, subprocess, threading, shlex
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Any
import requests
from stm32_extra_tools import EXTRA_TOOLS_MAP, EXTRA_TOOL_SCHEMAS
from gary_skills import (
    init_skills,
    handle_skill_command,
    SKILL_TOOLS_MAP,
    SKILL_TOOL_SCHEMAS,
    _get_manager,
)

# ─────────────────────────────────────────────────────────────
# 将本文件所在目录加入路径，使能 import compiler / config
# ─────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent.resolve()
sys.path.insert(0, str(_HERE))

# TUI 依赖
from rich.console import Console
from rich.panel import Panel
from rich.text import Text
from rich.table import Table
from rich.markdown import Markdown
from rich.rule import Rule
from rich import box

from prompt_toolkit import PromptSession
from prompt_toolkit.auto_suggest import AutoSuggestFromHistory
from prompt_toolkit.completion import Completer, Completion
from prompt_toolkit.formatted_text import HTML
from prompt_toolkit.styles import Style
from prompt_toolkit.history import FileHistory, InMemoryHistory

from openai import OpenAI

# Flash 项目内部模块
import config as _cfg
import compiler as _compiler_module

Compiler = _compiler_module.Compiler

AI_API_KEY = getattr(_cfg, "AI_API_KEY", "")
AI_BASE_URL = getattr(_cfg, "AI_BASE_URL", "https://api.openai.com/v1")
AI_MODEL = getattr(_cfg, "AI_MODEL", "gpt-4o")
AI_TEMPERATURE = getattr(_cfg, "AI_TEMPERATURE", 1)

WORKSPACE = _cfg.WORKSPACE
BUILD_DIR = _cfg.BUILD_DIR
PROJECTS_DIR = _cfg.PROJECTS_DIR
DEFAULT_CHIP = getattr(_cfg, "DEFAULT_CHIP", "")
DEFAULT_CLOCK = getattr(_cfg, "DEFAULT_CLOCK", "HSI_internal")
CLI_LANGUAGE = getattr(_cfg, "CLI_LANGUAGE", "zh")
SERIAL_PORT = getattr(_cfg, "SERIAL_PORT", "")
SERIAL_BAUD = getattr(_cfg, "SERIAL_BAUD", 115200)
POST_FLASH_DELAY = getattr(_cfg, "POST_FLASH_DELAY", 1.5)
REGISTER_READ_DELAY = getattr(_cfg, "REGISTER_READ_DELAY", 0.3)

# ─────────────────────────────────────────────────────────────
# AI 接口管理（动态读写 config.py，支持运行时切换）
# ─────────────────────────────────────────────────────────────
_AI_PRESETS = [
    # (显示名,                base_url,                                                    默认 model)
    ("OpenAI", "https://api.openai.com/v1", "gpt-4o"),
    ("DeepSeek", "https://api.deepseek.com/v1", "deepseek-chat"),
    ("Kimi / Moonshot", "https://api.moonshot.cn/v1", "kimi-k2.5"),
    (
        "Google Gemini",
        "https://generativelanguage.googleapis.com/v1beta/openai/",
        "gemini-2.0-flash",
    ),
    ("通义千问 (阿里云)", "https://dashscope.aliyuncs.com/compatible-mode/v1", "qwen-plus"),
    ("智谱 GLM", "https://open.bigmodel.cn/api/paas/v4/", "glm-4-flash"),
    ("Ollama (本地)", "http://127.0.0.1:11434/v1", "qwen2.5-coder:14b"),
    ("自定义 / Other", "", ""),
]

_CONFIG_KEYS_TO_RELOAD = (
    "AI_API_KEY",
    "AI_BASE_URL",
    "AI_MODEL",
    "AI_TEMPERATURE",
    "DEFAULT_CHIP",
    "DEFAULT_CLOCK",
    "CLI_LANGUAGE",
    "SERIAL_PORT",
    "SERIAL_BAUD",
    "POST_FLASH_DELAY",
    "REGISTER_READ_DELAY",
)


def _parse_cli_language(value: Any, default: Optional[str] = None) -> Optional[str]:
    raw = str(value or "").strip().lower()
    if not raw:
        return default
    if raw in {"en", "eng", "english", "英文"}:
        return "en"
    if raw in {"zh", "cn", "zh-cn", "zh_cn", "chinese", "中文"}:
        return "zh"
    return None


def _normalize_cli_language(value: Any) -> str:
    return _parse_cli_language(value, default="zh") or "zh"


CLI_LANGUAGE = _normalize_cli_language(CLI_LANGUAGE)


def _is_cli_english() -> bool:
    return CLI_LANGUAGE == "en"


def _cli_text(zh: str, en: str) -> str:
    return en if _is_cli_english() else zh


def _upsert_config_assignment(text: str, key: str, value: Any) -> str:
    line = f"{key} = {json.dumps(value, ensure_ascii=False)}"
    pattern = rf"^{key}\s*=.*$"
    if re.search(pattern, text, re.MULTILINE):
        return re.sub(pattern, line, text, flags=re.MULTILINE)
    if text and not text.endswith("\n"):
        text += "\n"
    return text + line + "\n"


def _read_ai_config() -> tuple:
    """从 config.py 读取 (api_key, base_url, model)"""
    p = _HERE / "config.py"
    if not p.exists():
        return AI_API_KEY, AI_BASE_URL, AI_MODEL
    text = p.read_text(encoding="utf-8")

    def _get(pat):
        m = re.search(pat, text, re.MULTILINE)
        return m.group(1).strip() if m else ""

    return (
        _get(r'^AI_API_KEY\s*=\s*["\']([^"\']*)["\']') or AI_API_KEY,
        _get(r'^AI_BASE_URL\s*=\s*["\']([^"\']*)["\']') or AI_BASE_URL,
        _get(r'^AI_MODEL\s*=\s*["\']([^"\']*)["\']') or AI_MODEL,
    )


def _write_config_assignments(updates: Dict[str, Any]) -> bool:
    """原地修改 config.py 中的若干配置项"""
    p = _HERE / "config.py"
    if not p.exists():
        return False
    text = p.read_text(encoding="utf-8")
    for key, value in updates.items():
        text = _upsert_config_assignment(text, key, value)
    p.write_text(text, encoding="utf-8")
    return True


def _write_ai_config(api_key: str, base_url: str, model: str) -> bool:
    """原地修改 config.py 的三行 AI 配置"""
    return _write_config_assignments(
        {
            "AI_API_KEY": api_key,
            "AI_BASE_URL": base_url,
            "AI_MODEL": model,
        }
    )


def _write_cli_language_config(language: str) -> bool:
    return _write_config_assignments(
        {
            "CLI_LANGUAGE": _normalize_cli_language(language),
        }
    )


def _reload_ai_globals():
    """写入 config.py 后重新加载 AI 相关全局变量"""
    import importlib, config as _cfg

    importlib.reload(_cfg)
    defaults = {
        "AI_API_KEY": "",
        "AI_BASE_URL": "https://api.openai.com/v1",
        "AI_MODEL": "gpt-4o",
        "AI_TEMPERATURE": 1,
        "DEFAULT_CHIP": "",
        "DEFAULT_CLOCK": "HSI_internal",
        "CLI_LANGUAGE": "zh",
        "SERIAL_PORT": "",
        "SERIAL_BAUD": 115200,
        "POST_FLASH_DELAY": 1.5,
        "REGISTER_READ_DELAY": 0.3,
    }
    for name in _CONFIG_KEYS_TO_RELOAD:
        globals()[name] = getattr(_cfg, name, defaults[name])
    globals()["CLI_LANGUAGE"] = _normalize_cli_language(globals().get("CLI_LANGUAGE", "zh"))


def _mask_key(key: str) -> str:
    if not key:
        return "(未设置)"
    return key[:6] + "..." + key[-4:] if len(key) > 12 else "***"


def _api_key_is_placeholder(api_key: str) -> bool:
    key = (api_key or "").strip()
    if not key:
        return True
    placeholder_prefixes = ("YOUR_API_KEY", "sk-YOUR")
    return any(key.startswith(prefix) for prefix in placeholder_prefixes)


def _ai_is_configured() -> bool:
    api_key, base_url, model = _read_ai_config()
    return bool(api_key and base_url and model and not _api_key_is_placeholder(api_key))


# ─────────────────────────────────────────────────────────────
# Gary member.md 经验库（自动记忆 + 系统提示注入）
# ─────────────────────────────────────────────────────────────
MEMBER_MD_PATH = _HERE / "member.md"
MEMBER_PROMPT_CHAR_LIMIT = 12000
MEMBER_PROMPT_MAX_DYNAMIC = 24
MEMBER_MAX_FILE_CHARS = 40000
MEMBER_MAX_DYNAMIC_ENTRIES = 120
_MEMBER_LOCK = threading.RLock()


def _default_member_content() -> str:
    return """# Gary Member Memory

## Focus
- 这里只记录高价值、可复用、能提高成功率的经验。
- 自动写入：成功编译、成功运行闭环。
- 主动写入：遇到关键初始化顺序、硬件坑、寄存器判定经验、稳定模板时，调用 `gary_save_member_memory`。
- 经验必须短、具体、可执行，不要粘贴大段原始日志。

## Memories

### [Pinned] 启动标记优先
- UART 初始化后立刻打印 `Gary:BOOT`，再初始化 I2C/SPI/TIM/OLED 等外设。
- 这样即使后续外设卡死，也能先确认程序已启动。

### [Pinned] 裸机 HAL_Delay 依赖 SysTick_Handler
- 裸机代码必须定义 `void SysTick_Handler(void) { HAL_IncTick(); }`。
- 否则 `HAL_Delay()` 会永久阻塞。

### [Pinned] I2C 外设先探测再使用
- 初始化后先 `HAL_I2C_IsDeviceReady()` 检查从设备是否应答。
- 无应答优先怀疑接线或地址，不要盲改业务逻辑。

### [Pinned] 增量修改优先精确替换
- 修改已有工程时优先 `str_replace_edit` + `stm32_recompile`。
- 不要无必要地整文件重写。

### [Pinned] 裸机禁止 sprintf/printf/malloc
- 裸机项目优先手写轻量调试输出，避免 `_sbrk/end` 链接错误。
"""


def _ensure_member_file() -> Path:
    with _MEMBER_LOCK:
        if not MEMBER_MD_PATH.exists():
            MEMBER_MD_PATH.write_text(_default_member_content(), encoding="utf-8")
    return MEMBER_MD_PATH


def _normalize_member_text(value: Any, limit: int = 220) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" -\t\r\n")
    if len(text) > limit:
        text = text[: limit - 3].rstrip() + "..."
    return text


def _member_text_to_lines(value: Any, per_line_limit: int = 220, max_lines: int = 8) -> list[str]:
    source = str(value or "").replace("\r", "\n")
    lines = []
    for raw in source.splitlines():
        line = _normalize_member_text(raw, limit=per_line_limit)
        if line:
            lines.append(line)
        if len(lines) >= max_lines:
            break
    if not lines:
        single = _normalize_member_text(source, limit=per_line_limit)
        if single:
            lines.append(single)
    return lines


def _split_member_content(text: str) -> tuple[str, list[str]]:
    source = (text or "").strip()
    if not source:
        source = _default_member_content().strip()
    if "## Memories" not in source:
        source = _default_member_content().strip() + "\n\n" + source
    before, _, after = source.partition("## Memories")
    header = before.rstrip() + "\n\n## Memories\n"
    entries = [
        chunk.strip() for chunk in re.split(r"(?m)(?=^### )", after.strip()) if chunk.strip()
    ]
    return header, entries


def _prune_member_content(text: str) -> str:
    header, entries = _split_member_content(text)
    pinned = [entry for entry in entries if entry.startswith("### [Pinned]")]
    dynamic = [entry for entry in entries if not entry.startswith("### [Pinned]")]
    dynamic = dynamic[-MEMBER_MAX_DYNAMIC_ENTRIES:]

    kept = []
    total = len(header)
    for entry in pinned:
        needed = len(entry) + 2
        if kept and total + needed > MEMBER_MAX_FILE_CHARS:
            break
        kept.append(entry)
        total += needed

    recent = []
    for entry in reversed(dynamic):
        needed = len(entry) + 2
        if recent and total + needed > MEMBER_MAX_FILE_CHARS:
            break
        recent.append(entry)
        total += needed
    recent.reverse()
    kept.extend(recent)

    content = header.rstrip() + "\n\n" + "\n\n".join(kept).strip()
    return content.strip() + "\n"


def _append_member_memory(
    title: str,
    experience: str,
    tags: Optional[list[str]] = None,
    source: str = "manual",
    importance: str = "high",
) -> dict:
    clean_title = _normalize_member_text(title, limit=120)
    lines = _member_text_to_lines(experience, per_line_limit=220, max_lines=10)
    if not clean_title or not lines:
        return {"success": False, "message": "title 或 experience 为空"}

    tags = [
        _normalize_member_text(tag, limit=24)
        for tag in (tags or [])
        if _normalize_member_text(tag, limit=24)
    ][:10]
    source = _normalize_member_text(source, limit=40) or "manual"
    importance = _normalize_member_text(importance, limit=16) or "high"

    fingerprint = _normalize_member_text(clean_title + " " + " ".join(lines), limit=500).lower()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    entry_lines = [
        f"### [{timestamp}] {clean_title}",
        f"- importance: {importance}",
        f"- source: {source}",
    ]
    if tags:
        entry_lines.append(f"- tags: {', '.join(tags)}")
    entry_lines.extend(f"- {line}" for line in lines)
    entry = "\n".join(entry_lines)

    with _MEMBER_LOCK:
        path = _ensure_member_file()
        current = path.read_text(encoding="utf-8")
        normalized_current = re.sub(r"\s+", " ", current).lower()
        if fingerprint and fingerprint in normalized_current:
            return {
                "success": True,
                "deduplicated": True,
                "path": str(path),
                "message": f"member.md 已存在相似经验: {clean_title}",
            }
        updated = _prune_member_content(current.rstrip() + "\n\n" + entry + "\n")
        path.write_text(updated, encoding="utf-8")
    return {
        "success": True,
        "deduplicated": False,
        "path": str(MEMBER_MD_PATH),
        "message": f"已写入 member.md: {clean_title}",
    }


def gary_save_member_memory(
    title: str,
    experience: str,
    tags: list = None,
    importance: str = "high",
) -> dict:
    """将高价值经验写入 Gary 的 member.md 经验库。"""
    return _append_member_memory(
        title=title,
        experience=experience,
        tags=tags or [],
        source="model",
        importance=importance,
    )


def _infer_code_tags(code: str) -> list[str]:
    text = code or ""
    checks = [
        (
            "rtos",
            any(token in text for token in ("FreeRTOS.h", "xTaskCreate", "vTaskDelay(", "task.h")),
        ),
        ("uart", "HAL_UART_" in text or "USART" in text),
        ("debug_print", "Debug_Print(" in text or "Debug_PrintInt(" in text),
        ("boot_marker", "Gary:BOOT" in text),
        ("systick", "SysTick_Handler" in text),
        ("i2c", "HAL_I2C_" in text or "I2C_HandleTypeDef" in text),
        ("spi", "HAL_SPI_" in text or "SPI_HandleTypeDef" in text),
        ("adc", "HAL_ADC_" in text or "ADC_HandleTypeDef" in text),
        ("pwm", "HAL_TIM_PWM_" in text or "TIM_HandleTypeDef" in text),
        ("oled", "OLED_" in text),
        ("sensor_check", "HAL_I2C_IsDeviceReady" in text),
        ("fault_analysis", "HardFault" in text or "SCB_CFSR" in text),
    ]
    tags = [name for name, matched in checks if matched]
    if "rtos" not in tags:
        tags.insert(0, "baremetal")
    return tags[:10]


def _derive_success_patterns_from_code(code: str) -> list[str]:
    text = code or ""
    patterns = []
    if "Gary:BOOT" in text and ("MX_USART" in text or "HAL_UART_Init" in text):
        patterns.append("UART 初始化后会尽早打印 Gary:BOOT，启动链路更容易确认。")
    if "SysTick_Handler" in text and "FreeRTOS.h" not in text:
        patterns.append("裸机代码显式定义 SysTick_Handler，HAL_Delay 不会卡死。")
    if "vApplicationTickHook" in text:
        patterns.append("FreeRTOS 项目通过 vApplicationTickHook 维持 HAL tick。")
    if "HAL_I2C_IsDeviceReady" in text:
        patterns.append("I2C 设备在使用前会先做在线探测。")
    if "Debug_Print(" in text or "Debug_PrintInt(" in text:
        patterns.append("代码保留了轻量串口调试输出，便于运行时定位问题。")
    if "xTaskCreate" in text:
        patterns.append("当前成功模板已经包含可编译的 FreeRTOS 任务结构。")
    if "str_replace_edit" in text:
        patterns.append("此模板来自增量修改链路，适合继续做精准替换。")
    return patterns[:5]


def _record_success_memory(
    event_type: str,
    code: str,
    result: Optional[dict] = None,
    request: str = "",
    steps: Optional[list] = None,
):
    try:
        tags = _infer_code_tags(code)
        chip = _current_chip or DEFAULT_CHIP or "UNKNOWN"
        mode = "rtos" if "rtos" in tags else "baremetal"
        short_request = _normalize_member_text(request, limit=60)
        if event_type == "runtime_success":
            title = f"运行成功闭环 | {chip} | {mode}"
            if short_request:
                title += f" | {short_request}"
        else:
            title = f"编译成功模板 | {chip} | {mode}"

        lines = []
        if short_request:
            lines.append(f"需求: {short_request}")
        if result and result.get("bin_size"):
            lines.append(f"bin_size: {result['bin_size']} B")
        if tags:
            lines.append(f"特征: {', '.join(tags[:8])}")
        if event_type == "runtime_success":
            lines.append("烧录、启动、串口/寄存器验证通过，无 HardFault、无硬件缺失。")
            uart_step = next((step for step in (steps or []) if step.get("step") == "uart"), None)
            reg_step = next(
                (step for step in (steps or []) if step.get("step") == "registers"), None
            )
            if uart_step and uart_step.get("boot_ok"):
                lines.append("串口已看到 Gary:BOOT，启动链路正常。")
            if reg_step and reg_step.get("key_regs"):
                reg_names = list(reg_step["key_regs"].keys())[:6]
                lines.append(f"运行时已回读关键寄存器: {', '.join(reg_names)}")
        else:
            lines.append("当前代码已在本机工具链上成功编译通过。")
        lines.extend(_derive_success_patterns_from_code(code))
        _append_member_memory(
            title=title,
            experience="\n".join(lines),
            tags=tags,
            source=event_type,
            importance="critical" if event_type == "runtime_success" else "high",
        )
    except Exception as e:
        _telegram_log(f"member auto_save error={str(e)[:160]}")


def _member_prompt_section() -> str:
    with _MEMBER_LOCK:
        path = _ensure_member_file()
        current = path.read_text(encoding="utf-8")
    header, entries = _split_member_content(current)
    pinned = [entry for entry in entries if entry.startswith("### [Pinned]")]
    dynamic = [entry for entry in entries if not entry.startswith("### [Pinned]")]

    selected = []
    total = len(header)
    for entry in pinned:
        needed = len(entry) + 2
        if selected and total + needed > MEMBER_PROMPT_CHAR_LIMIT:
            break
        selected.append(entry)
        total += needed

    recent = []
    for entry in reversed(dynamic):
        needed = len(entry) + 2
        if recent and (
            total + needed > MEMBER_PROMPT_CHAR_LIMIT or len(recent) >= MEMBER_PROMPT_MAX_DYNAMIC
        ):
            break
        recent.append(entry)
        total += needed
    recent.reverse()
    selected.extend(recent)

    excerpt = header.rstrip()
    if selected:
        excerpt += "\n\n" + "\n\n".join(selected)
    return (
        "## Gary Member Memory（重点）\n"
        "以下内容来自 member.md，是 Gary 的长期经验库。优先复用这些成功经验；"
        "遇到新的高价值经验时，调用 `gary_save_member_memory` 写进去。\n\n"
        f"{excerpt.strip()}"
    )


def _member_preview_markdown(max_dynamic: int = 10) -> str:
    with _MEMBER_LOCK:
        path = _ensure_member_file()
        current = path.read_text(encoding="utf-8")
    header, entries = _split_member_content(current)
    pinned = [entry for entry in entries if entry.startswith("### [Pinned]")]
    dynamic = [entry for entry in entries if not entry.startswith("### [Pinned]")]
    selected = pinned + dynamic[-max_dynamic:]
    path_line = _cli_text("路径", "Path")
    body = header.rstrip()
    if selected:
        body += "\n\n" + "\n\n".join(selected)
    return f"**{path_line}:** `{path}`\n\n{body.strip()}"


# ─────────────────────────────────────────────────────────────
# Telegram 机器人配置 / 守护进程 / 长轮询
# ─────────────────────────────────────────────────────────────
GARY_HOME = Path.home() / ".gary"
TELEGRAM_CONFIG_PATH = GARY_HOME / "telegram_bot.json"
TELEGRAM_PID_PATH = GARY_HOME / "telegram_bot.pid"
TELEGRAM_LOG_PATH = GARY_HOME / "telegram_bot.log"
TELEGRAM_MESSAGE_LIMIT = 3800
_TELEGRAM_CONFIG_LOCK = threading.RLock()


def _ensure_gary_home():
    GARY_HOME.mkdir(parents=True, exist_ok=True)


def _default_telegram_config() -> dict:
    return {
        "bot_token": "",
        "bot_id": None,
        "bot_username": "",
        "bot_name": "",
        "allow_all_chats": False,
        "allowed_chat_ids": [],
        "allowed_user_ids": [],
        "last_update_id": 0,
        "updated_at": "",
    }


def _unique_int_list(values: Any) -> list[int]:
    result = []
    seen = set()
    for value in values or []:
        try:
            num = int(str(value).strip())
        except (TypeError, ValueError):
            continue
        if num in seen:
            continue
        seen.add(num)
        result.append(num)
    return result


def _normalize_telegram_config(config: Optional[dict]) -> dict:
    merged = _default_telegram_config()
    if isinstance(config, dict):
        merged.update(config)
    merged["allow_all_chats"] = bool(merged.get("allow_all_chats", False))
    merged["allowed_chat_ids"] = _unique_int_list(merged.get("allowed_chat_ids"))
    merged["allowed_user_ids"] = _unique_int_list(merged.get("allowed_user_ids"))
    try:
        merged["last_update_id"] = int(merged.get("last_update_id", 0) or 0)
    except (TypeError, ValueError):
        merged["last_update_id"] = 0
    return merged


def _read_telegram_config() -> dict:
    with _TELEGRAM_CONFIG_LOCK:
        if not TELEGRAM_CONFIG_PATH.exists():
            return _default_telegram_config()
        try:
            raw = json.loads(TELEGRAM_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return _default_telegram_config()
        return _normalize_telegram_config(raw)


def _write_telegram_config(config: dict) -> dict:
    with _TELEGRAM_CONFIG_LOCK:
        _ensure_gary_home()
        clean = _normalize_telegram_config(config)
        clean["updated_at"] = datetime.now().isoformat(timespec="seconds")
        TELEGRAM_CONFIG_PATH.write_text(
            json.dumps(clean, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return clean


def _telegram_is_configured(config: Optional[dict] = None) -> bool:
    config = config or _read_telegram_config()
    token = str(config.get("bot_token", "")).strip()
    placeholders = ("", "YOUR_TELEGRAM_BOT_TOKEN", "123456:ABC")
    return bool(token and token not in placeholders and ":" in token)


def _mask_telegram_token(token: str) -> str:
    if not token:
        return "(未设置)"
    return token[:8] + "..." + token[-6:] if len(token) > 20 else "***"


def _split_tokens(raw: str) -> list[str]:
    return [tok for tok in re.split(r"[\s,]+", raw.strip()) if tok]


def _parse_telegram_targets(raw: str) -> dict:
    parsed = {"chat_ids": [], "user_ids": [], "invalid": []}
    for token in _split_tokens(raw):
        kind = "chat"
        value = token
        lower = token.lower()
        if lower.startswith("user:"):
            kind, value = "user", token.split(":", 1)[1]
        elif lower.startswith("chat:"):
            kind, value = "chat", token.split(":", 1)[1]
        try:
            num = int(value)
        except ValueError:
            parsed["invalid"].append(token)
            continue
        if kind == "user":
            parsed["user_ids"].append(num)
        else:
            parsed["chat_ids"].append(num)
    parsed["chat_ids"] = _unique_int_list(parsed["chat_ids"])
    parsed["user_ids"] = _unique_int_list(parsed["user_ids"])
    return parsed


def _telegram_set_permissions(
    *,
    add_chat_ids: Optional[list[int]] = None,
    remove_chat_ids: Optional[list[int]] = None,
    add_user_ids: Optional[list[int]] = None,
    remove_user_ids: Optional[list[int]] = None,
    allow_all_chats: Optional[bool] = None,
) -> dict:
    config = _read_telegram_config()
    chats = set(config.get("allowed_chat_ids", []))
    users = set(config.get("allowed_user_ids", []))
    chats.update(add_chat_ids or [])
    users.update(add_user_ids or [])
    chats.difference_update(remove_chat_ids or [])
    users.difference_update(remove_user_ids or [])
    config["allowed_chat_ids"] = sorted(chats)
    config["allowed_user_ids"] = sorted(users)
    if allow_all_chats is not None:
        config["allow_all_chats"] = bool(allow_all_chats)
    return _write_telegram_config(config)


def _read_pid_file() -> Optional[int]:
    try:
        return int(TELEGRAM_PID_PATH.read_text(encoding="utf-8").strip())
    except Exception:
        return None


def _pid_is_alive(pid: Optional[int]) -> bool:
    if not pid:
        return False
    proc_stat = Path(f"/proc/{pid}/stat")
    if proc_stat.exists():
        try:
            stat_text = proc_stat.read_text(encoding="utf-8", errors="ignore")
            fields = stat_text.split()
            if len(fields) >= 3 and fields[2] == "Z":
                return False
            return True
        except Exception:
            pass
    try:
        os.kill(pid, 0)
    except OSError:
        return False
    return True


def _write_pid_file(pid: int):
    _ensure_gary_home()
    TELEGRAM_PID_PATH.write_text(str(pid), encoding="utf-8")


def _clear_pid_file(expected_pid: Optional[int] = None):
    if not TELEGRAM_PID_PATH.exists():
        return
    if expected_pid is not None:
        current = _read_pid_file()
        if current and current != expected_pid:
            return
    try:
        TELEGRAM_PID_PATH.unlink()
    except FileNotFoundError:
        pass


def _telegram_daemon_status() -> dict:
    pid = _read_pid_file()
    running = _pid_is_alive(pid)
    if pid and not running:
        _clear_pid_file(pid)
        pid = None
    return {
        "running": running,
        "pid": pid,
        "log_path": str(TELEGRAM_LOG_PATH),
    }


def _telegram_api_call(
    token: str, method: str, payload: Optional[dict] = None, timeout: float = 30.0
):
    url = f"https://api.telegram.org/bot{token}/{method}"
    try:
        response = requests.post(url, json=payload or {}, timeout=timeout)
    except requests.RequestException as e:
        raise RuntimeError(f"Telegram API 网络错误: {e}") from e

    try:
        data = response.json()
    except ValueError as e:
        raise RuntimeError(f"Telegram API 返回非法 JSON: HTTP {response.status_code}") from e

    if response.status_code >= 400 or not data.get("ok"):
        detail = data.get("description") or response.text[:200] or f"HTTP {response.status_code}"
        raise RuntimeError(f"Telegram API 调用失败: {detail}")
    return data.get("result")


def _telegram_get_me(token: str) -> dict:
    result = _telegram_api_call(token, "getMe", timeout=15.0)
    return result if isinstance(result, dict) else {}


def _telegram_set_my_commands(token: str):
    commands = [
        {"command": "start", "description": "查看 Gary 机器人状态和用法"},
        {"command": "help", "description": "查看 Telegram 侧命令"},
        {"command": "clear", "description": "清空当前聊天上下文"},
        {"command": "status", "description": "查看 Gary 当前硬件状态"},
        {"command": "connect", "description": "连接探针和串口，可带芯片型号"},
        {"command": "disconnect", "description": "断开当前硬件连接"},
        {"command": "chip", "description": "查看或切换芯片型号"},
        {"command": "projects", "description": "查看最近历史项目"},
    ]
    _telegram_api_call(token, "setMyCommands", {"commands": commands}, timeout=20.0)


def _telegram_split_text(text: str, limit: int = TELEGRAM_MESSAGE_LIMIT) -> list[str]:
    source = (text or "").strip() or "Gary 已处理，但没有返回文本。"
    chunks = []
    while len(source) > limit:
        cut = source.rfind("\n\n", 0, limit)
        if cut < limit // 3:
            cut = source.rfind("\n", 0, limit)
        if cut < limit // 3:
            cut = source.rfind(" ", 0, limit)
        if cut < limit // 3:
            cut = limit
        chunks.append(source[:cut].rstrip())
        source = source[cut:].lstrip()
    if source:
        chunks.append(source)
    return chunks


def _telegram_send_text(
    token: str, chat_id: int, text: str, reply_to_message_id: Optional[int] = None
):
    for index, chunk in enumerate(_telegram_split_text(text)):
        payload = {
            "chat_id": chat_id,
            "text": chunk,
            "disable_web_page_preview": True,
        }
        if index == 0 and reply_to_message_id is not None:
            payload["reply_to_message_id"] = reply_to_message_id
        _telegram_api_call(token, "sendMessage", payload, timeout=20.0)


def _telegram_log(message: str):
    try:
        _ensure_gary_home()
        with TELEGRAM_LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write(f"[{datetime.now().isoformat(timespec='seconds')}] {message}\n")
    except Exception:
        pass


def _telegram_send_chat_action(token: str, chat_id: int, action: str = "typing"):
    _telegram_api_call(
        token,
        "sendChatAction",
        {"chat_id": chat_id, "action": action},
        timeout=10.0,
    )


class _TelegramTypingPulse:
    def __init__(self, token: str, chat_id: int, interval: float = 4.0):
        self.token = token
        self.chat_id = chat_id
        self.interval = interval
        self._stop_event = threading.Event()
        self._thread: Optional[threading.Thread] = None

    def __enter__(self):
        self.start()
        return self

    def __exit__(self, exc_type, exc, tb):
        self.stop()

    def start(self):
        if self._thread and self._thread.is_alive():
            return
        self._stop_event.clear()
        self._thread = threading.Thread(
            target=self._run,
            name=f"TelegramTypingPulse:{self.chat_id}",
            daemon=True,
        )
        self._thread.start()

    def stop(self):
        self._stop_event.set()
        if self._thread and self._thread.is_alive():
            self._thread.join(timeout=1.0)
        self._thread = None

    def _run(self):
        while not self._stop_event.is_set():
            try:
                _telegram_send_chat_action(self.token, self.chat_id, action="typing")
            except Exception:
                pass
            if self._stop_event.wait(self.interval):
                break


class _TelegramPhaseReporter:
    def __init__(
        self,
        token: str,
        chat_id: int,
        reply_to_message_id: Optional[int] = None,
        idle_notice_delay: float = 60.0,
        heartbeat_interval: float = 60.0,
    ):
        self.token = token
        self.chat_id = chat_id
        self.reply_to_message_id = reply_to_message_id
        self.idle_notice_delay = idle_notice_delay
        self.heartbeat_interval = heartbeat_interval
        self.preface_text = ""
        self.preface_sent = False
        self._idle_notice_sent = False
        self._last_tool_notice = ""
        self._last_tool_notice_ts = 0.0
        self._current_stage = "分析需求"
        self._last_progress_ts = time.time()
        self._stop_event = threading.Event()
        self._progress_thread: Optional[threading.Thread] = None

    def start(self):
        if self._progress_thread and self._progress_thread.is_alive():
            return
        self._stop_event.clear()
        self._progress_thread = threading.Thread(
            target=self._progress_loop,
            name=f"TelegramProgress:{self.chat_id}",
            daemon=True,
        )
        self._progress_thread.start()

    def stop(self):
        self._stop_event.set()
        if self._progress_thread and self._progress_thread.is_alive():
            self._progress_thread.join(timeout=0.5)
        self._progress_thread = None

    def capture_preface(self, text: str):
        cleaned = (text or "").strip()
        if cleaned:
            self.preface_text = cleaned

    def send_preface_if_needed(self):
        if self.preface_sent or not self.preface_text:
            return
        try:
            _telegram_send_text(
                self.token,
                self.chat_id,
                self.preface_text,
                reply_to_message_id=self.reply_to_message_id,
            )
            self.preface_sent = True
            self._note_progress("前置说明")
            _telegram_log(f"telegram preface chat={self.chat_id} len={len(self.preface_text)}")
        except Exception:
            pass

    def tool_start(self, name: str):
        self.send_preface_if_needed()
        tool_name = (name or "").strip() or "unknown"
        now = time.time()
        if tool_name == self._last_tool_notice and now - self._last_tool_notice_ts < 1.5:
            return
        self._last_tool_notice = tool_name
        self._last_tool_notice_ts = now
        self._current_stage = f"执行工具 {tool_name}"
        try:
            _telegram_send_text(
                self.token,
                self.chat_id,
                f"🔧 正在调用工具: {tool_name}",
                reply_to_message_id=self.reply_to_message_id,
            )
            self._note_progress(self._current_stage)
            _telegram_log(f"telegram tool_start chat={self.chat_id} name={tool_name}")
        except Exception:
            pass

    def tool_finish(self, name: str, preview: str = ""):
        tool_name = (name or "").strip() or "unknown"
        self._current_stage = f"等待下一步 ({tool_name} 已完成)"
        text = f"✅ 工具完成: {tool_name}"
        snippet = (preview or "").strip()
        if snippet:
            snippet = snippet[:100]
            text += f"\n{snippet}"
        try:
            _telegram_send_text(
                self.token,
                self.chat_id,
                text,
                reply_to_message_id=self.reply_to_message_id,
            )
            self._note_progress(self._current_stage)
            _telegram_log(f"telegram tool_finish chat={self.chat_id} name={tool_name}")
        except Exception:
            pass

    def tool_error(self, name: str, error: str):
        self.send_preface_if_needed()
        self._current_stage = f"工具失败 {(name or '').strip() or 'unknown'}"
        try:
            detail = (error or "").strip()
            if len(detail) > 160:
                detail = detail[:157] + "..."
            _telegram_send_text(
                self.token,
                self.chat_id,
                f"❌ 工具失败: {(name or '').strip() or 'unknown'}\n{detail or '未知错误'}",
                reply_to_message_id=self.reply_to_message_id,
            )
            self._note_progress(self._current_stage)
            _telegram_log(
                f"telegram tool_error chat={self.chat_id} name={(name or '').strip() or 'unknown'} detail={detail[:80]}"
            )
        except Exception:
            pass

    def strip_preface_from_reply(self, reply: str) -> str:
        cleaned_reply = (reply or "").strip()
        preface = self.preface_text.strip()
        if not self.preface_sent or not preface or not cleaned_reply:
            return cleaned_reply
        if cleaned_reply == preface:
            return ""
        prefix = preface + "\n\n"
        if cleaned_reply.startswith(prefix):
            return cleaned_reply[len(prefix) :].lstrip()
        return cleaned_reply

    def _note_progress(self, stage: str):
        self._current_stage = stage
        self._last_progress_ts = time.time()

    def _progress_loop(self):
        while not self._stop_event.wait(1.0):
            now = time.time()
            if (
                not self._idle_notice_sent
                and not self.preface_sent
                and not self._last_tool_notice
                and now - self._last_progress_ts >= self.idle_notice_delay
            ):
                try:
                    _telegram_send_text(
                        self.token,
                        self.chat_id,
                        "Gary 正在分析需求，准备开始处理。",
                        reply_to_message_id=self.reply_to_message_id,
                    )
                    self._idle_notice_sent = True
                    self._note_progress("分析需求")
                    _telegram_log(f"telegram idle_notice chat={self.chat_id}")
                except Exception:
                    pass
                continue

            if now - self._last_progress_ts < self.heartbeat_interval:
                continue

            stage = self._current_stage or "处理中"
            try:
                _telegram_send_text(
                    self.token,
                    self.chat_id,
                    f"⏳ Gary 仍在处理: {stage}",
                    reply_to_message_id=self.reply_to_message_id,
                )
                self._note_progress(stage)
                _telegram_log(f"telegram heartbeat chat={self.chat_id} stage={stage}")
            except Exception:
                pass


def _telegram_is_authorized(config: dict, chat_id: int, user_id: int) -> bool:
    if config.get("allow_all_chats"):
        return True
    if chat_id in set(config.get("allowed_chat_ids", [])):
        return True
    if user_id in set(config.get("allowed_user_ids", [])):
        return True
    return False


def _telegram_unauthorized_text(chat_id: int, user_id: int) -> str:
    return (
        "这个 Telegram 会话还没授权。\n\n"
        f"chat_id = {chat_id}\n"
        f"user_id = {user_id}\n\n"
        "请在终端执行任一命令后再试：\n"
        f"gary telegram allow {chat_id}\n"
        f"gary telegram allow user:{user_id}"
    )


def _telegram_status_lines(include_commands: bool = True) -> list[str]:
    config = _read_telegram_config()
    daemon = _telegram_daemon_status()
    lines = [
        "Telegram 机器人状态",
        f"AI 接口: {'已配置' if _ai_is_configured() else '未配置'}",
        f"机器人: {'已配置' if _telegram_is_configured(config) else '未配置'}",
        f"Token: {_mask_telegram_token(config.get('bot_token', ''))}",
        f"Bot: @{config.get('bot_username') or '-'}",
        f"Bot 名称: {config.get('bot_name') or '-'}",
        f"授权模式: {'允许所有 chat' if config.get('allow_all_chats') else '白名单'}",
        f"允许的 chat_id: {config.get('allowed_chat_ids') or '[]'}",
        f"允许的 user_id: {config.get('allowed_user_ids') or '[]'}",
        f"守护进程: {'运行中' if daemon['running'] else '未运行'}",
        f"PID: {daemon.get('pid') or '-'}",
        f"日志: {daemon['log_path']}",
        f"last_update_id: {config.get('last_update_id', 0)}",
    ]
    if include_commands:
        lines.extend(
            [
                "",
                "常用命令:",
                "gary telegram start        启动后台机器人",
                "gary telegram stop         停止后台机器人",
                "gary telegram allow <id>   添加 chat_id 白名单",
                "gary telegram allow user:<id>   添加 user_id 白名单",
                "gary telegram remove <id>  删除白名单",
                "gary telegram allow-all    允许所有 chat",
                "gary telegram whitelist    切回白名单模式",
                "gary telegram reset        删除 Telegram 配置并停止机器人",
            ]
        )
    return lines


def _print_telegram_status(include_commands: bool = True):
    CONSOLE.print("\n".join(_telegram_status_lines(include_commands=include_commands)))
    CONSOLE.print()


def _telegram_help_text() -> str:
    return "\n".join(
        [
            "Telegram 侧可用命令：",
            "/start  查看状态和授权信息",
            "/help   查看帮助",
            "/clear  清空当前聊天上下文",
            "/status 查看 Gary 当前硬件状态",
            "/connect [芯片] 连接探针和串口",
            "/disconnect 断开硬件",
            "/chip [型号] 查看或切换芯片",
            "/projects 查看最近历史项目",
            "",
            "其他普通文本会直接发给 Gary，保持当前聊天上下文连续对话。",
        ]
    )


def _format_hw_status_for_text(status: dict) -> str:
    return "\n".join(
        [
            "Gary 当前状态",
            f"chip: {status.get('chip', '-')}",
            f"hw_connected: {status.get('hw_connected', False)}",
            f"serial_connected: {status.get('serial_connected', False)}",
            f"gcc_ok: {status.get('gcc_ok', False)}",
            f"gcc_version: {status.get('gcc_version', '-')}",
            f"hal_ok: {status.get('hal_ok', False)}",
            f"hal_lib_ok: {status.get('hal_lib_ok', False)}",
            f"workspace: {status.get('workspace', '-')}",
        ]
    )


def _format_projects_for_text(projects: list[dict]) -> str:
    if not projects:
        return "暂无历史项目"
    lines = ["最近项目："]
    for item in projects[:10]:
        lines.append(
            f"- {item.get('name', '-')}"
            f" | {item.get('chip', '-')}"
            f" | {str(item.get('request', ''))[:40]}"
        )
    return "\n".join(lines)


def _normalize_telegram_incoming_text(text: str, bot_username: str) -> Optional[str]:
    stripped = (text or "").strip()
    if not stripped.startswith("/"):
        return stripped
    head, sep, tail = stripped.partition(" ")
    command = head
    if "@" in head:
        base, mention = head.split("@", 1)
        if bot_username and mention.lower() != bot_username.lower():
            return None
        command = base
    return f"{command}{sep}{tail}".strip()


def configure_telegram_cli() -> dict:
    CONSOLE.print()
    CONSOLE.rule("[bold cyan]  配置 Telegram 机器人[/]")
    CONSOLE.print()

    config = _read_telegram_config()
    if _telegram_is_configured(config):
        CONSOLE.print(f"  [dim]当前 Token:[/] {_mask_telegram_token(config.get('bot_token', ''))}")
        CONSOLE.print(f"  [dim]当前 Bot  :[/] @{config.get('bot_username') or '-'}")
        CONSOLE.print(
            f"  [dim]授权模式 :[/] {'允许所有 chat' if config.get('allow_all_chats') else '白名单'}"
        )
        CONSOLE.print(f"  [dim]chat_id   :[/] {config.get('allowed_chat_ids') or '[]'}")
        CONSOLE.print(f"  [dim]user_id   :[/] {config.get('allowed_user_ids') or '[]'}")
        CONSOLE.print()

    token = ""
    current_token = str(config.get("bot_token", "")).strip()
    while True:
        try:
            prompt = "  Bot Token"
            if current_token:
                prompt += "（回车保留当前）"
            prompt += ": "
            entered = input(prompt).strip()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[dim]已取消[/]")
            return {"success": False, "message": "已取消"}

        token = entered or current_token
        if not token:
            CONSOLE.print("[yellow]  Bot Token 不能为空[/]")
            continue
        try:
            me = _telegram_get_me(token)
            config["bot_token"] = token
            config["bot_id"] = me.get("id")
            config["bot_username"] = me.get("username", "")
            config["bot_name"] = me.get("first_name", "")
            CONSOLE.print(f"[green]  ✓ Token 有效，机器人: @{config['bot_username']}[/]")
            break
        except Exception as e:
            CONSOLE.print(f"[red]  ✗ Token 校验失败: {e}[/]")

    CONSOLE.print()
    CONSOLE.print("[bold cyan]  授权模式[/]")
    CONSOLE.print("    [yellow]1[/]. 白名单（推荐）")
    CONSOLE.print("    [yellow]2[/]. 允许所有 chat")
    CONSOLE.print()

    default_choice = "2" if config.get("allow_all_chats") else "1"
    try:
        choice = input(f"  输入序号 [1-2]（默认 {default_choice}）: ").strip() or default_choice
    except (EOFError, KeyboardInterrupt):
        choice = default_choice

    config["allow_all_chats"] = choice == "2"

    if not config["allow_all_chats"]:
        try:
            raw_targets = input(
                "  输入白名单（chat_id 或 user:123，多项用空格/逗号分隔，回车保留当前）: "
            ).strip()
        except (EOFError, KeyboardInterrupt):
            raw_targets = ""
        if raw_targets:
            parsed = _parse_telegram_targets(raw_targets)
            if parsed["invalid"]:
                CONSOLE.print(f"[yellow]  忽略非法项: {parsed['invalid']}[/]")
            config["allowed_chat_ids"] = parsed["chat_ids"]
            config["allowed_user_ids"] = parsed["user_ids"]

    saved = _write_telegram_config(config)
    try:
        _telegram_set_my_commands(saved["bot_token"])
    except Exception as e:
        CONSOLE.print(f"[yellow]  命令菜单注册失败，可忽略: {e}[/]")

    CONSOLE.print()
    CONSOLE.print("[green]  ✓ Telegram 配置已保存[/]")
    CONSOLE.print(f"  [green]✓[/] Bot       @{saved.get('bot_username') or '-'}")
    CONSOLE.print(f"  [green]✓[/] Token     {_mask_telegram_token(saved.get('bot_token', ''))}")
    CONSOLE.print(
        f"  [green]✓[/] 授权模式  {'允许所有 chat' if saved.get('allow_all_chats') else '白名单'}"
    )
    if not saved.get("allow_all_chats"):
        CONSOLE.print(f"  [green]✓[/] chat_id   {saved.get('allowed_chat_ids') or '[]'}")
        CONSOLE.print(f"  [green]✓[/] user_id   {saved.get('allowed_user_ids') or '[]'}")
        CONSOLE.print(
            "  [dim]若还不知道 chat_id，可先在 Telegram 给机器人发 /start，再按提示执行 allow 命令[/]"
        )
    CONSOLE.print()
    return {"success": True, "config": saved}


def _ensure_ai_for_telegram() -> bool:
    if _ai_is_configured():
        return True
    CONSOLE.print("[yellow]Telegram 机器人要能回复，需要先配置 AI 接口[/]")
    configure_ai_cli()
    if _ai_is_configured():
        return True
    CONSOLE.print("[red]AI 接口仍未配置，无法启动 Telegram 机器人[/]")
    return False


def _start_telegram_daemon() -> dict:
    daemon = _telegram_daemon_status()
    if daemon["running"]:
        return {"success": True, "message": f"Telegram 机器人已在后台运行（PID {daemon['pid']}）"}

    config = _read_telegram_config()
    if not _telegram_is_configured(config):
        return {
            "success": False,
            "message": "Telegram 机器人尚未配置，请先运行 gary telegram 或 /telegram",
        }
    if not _ai_is_configured():
        return {"success": False, "message": "AI 接口未配置，先运行 gary config 或 /config"}

    _ensure_gary_home()
    log_handle = TELEGRAM_LOG_PATH.open("a", encoding="utf-8")
    log_handle.write(
        f"\n[{datetime.now().isoformat(timespec='seconds')}] starting telegram daemon\n"
    )
    log_handle.flush()
    try:
        process = subprocess.Popen(
            [sys.executable, str(_HERE / "stm32_agent.py"), "telegram", "serve", "--daemonized"],
            cwd=str(_HERE),
            stdout=log_handle,
            stderr=log_handle,
            stdin=subprocess.DEVNULL,
            start_new_session=True,
            close_fds=True,
        )
    finally:
        log_handle.close()

    time.sleep(1.2)
    if _pid_is_alive(process.pid):
        return {
            "success": True,
            "message": f"Telegram 机器人已启动（PID {process.pid}）",
            "pid": process.pid,
        }
    return {"success": False, "message": f"后台启动失败，请查看日志: {TELEGRAM_LOG_PATH}"}


def _stop_telegram_daemon() -> dict:
    daemon = _telegram_daemon_status()
    pid = daemon.get("pid")
    if not daemon["running"] or not pid:
        return {"success": True, "message": "Telegram 机器人当前未运行"}

    try:
        os.kill(pid, signal.SIGTERM)
    except OSError as e:
        _clear_pid_file(pid)
        return {"success": False, "message": f"停止失败: {e}"}

    for _ in range(20):
        if not _pid_is_alive(pid):
            _clear_pid_file(pid)
            return {"success": True, "message": f"Telegram 机器人已停止（PID {pid}）"}
        time.sleep(0.2)

    return {"success": False, "message": f"停止超时，请手动检查 PID {pid}"}


def _reset_telegram_config() -> dict:
    _stop_telegram_daemon()
    if TELEGRAM_CONFIG_PATH.exists():
        TELEGRAM_CONFIG_PATH.unlink()
    return {"success": True, "message": f"已删除 Telegram 配置: {TELEGRAM_CONFIG_PATH}"}


class TelegramBotBridge:
    def __init__(self):
        self._stop_event = threading.Event()
        self._thread: Optional[threading.Thread] = None
        self._chat_agents: Dict[int, "STM32Agent"] = {}
        self._last_error = ""
        self._started_at: Optional[float] = None

    def is_running(self) -> bool:
        return bool(self._thread and self._thread.is_alive())

    def start(self) -> dict:
        config = _read_telegram_config()
        if not _telegram_is_configured(config):
            return {"success": False, "message": "Telegram 机器人未配置"}
        if not _ai_is_configured():
            return {"success": False, "message": "AI 接口未配置"}
        if self.is_running():
            return {"success": True, "message": "Telegram 机器人已在当前进程运行"}
        self._stop_event.clear()
        self._thread = threading.Thread(
            target=self._poll_loop, name="GaryTelegramBridge", daemon=True
        )
        self._thread.start()
        self._started_at = time.time()
        return {"success": True, "message": "Telegram 机器人已启动"}

    def stop(self):
        self._stop_event.set()
        if self._thread and self._thread.is_alive():
            self._thread.join(timeout=5.0)
        self._thread = None

    def status(self) -> dict:
        return {
            "running": self.is_running(),
            "started_at": self._started_at,
            "last_error": self._last_error,
            "chat_sessions": len(self._chat_agents),
        }

    def _get_agent(self, chat_id: int) -> "STM32Agent":
        agent = self._chat_agents.get(chat_id)
        if agent is None:
            agent = STM32Agent(interactive=False)
            self._chat_agents[chat_id] = agent
        agent.refresh_ai_client()
        return agent

    def _reset_chat(self, chat_id: int):
        self._chat_agents.pop(chat_id, None)

    def _poll_loop(self):
        while not self._stop_event.is_set():
            config = _read_telegram_config()
            token = str(config.get("bot_token", "")).strip()
            if not token:
                self._last_error = "Telegram Token 为空"
                time.sleep(2.0)
                continue

            latest_seen = int(config.get("last_update_id", 0))
            try:
                updates = (
                    _telegram_api_call(
                        token,
                        "getUpdates",
                        {
                            "offset": latest_seen + 1,
                            "timeout": 25,
                            "allowed_updates": ["message"],
                        },
                        timeout=35.0,
                    )
                    or []
                )
                self._last_error = ""
            except Exception as e:
                self._last_error = str(e)
                time.sleep(2.0)
                continue

            for update in updates:
                latest_seen = max(latest_seen, int(update.get("update_id", 0)))
                try:
                    self._handle_update(update)
                except Exception as e:
                    self._last_error = str(e)
                    CONSOLE.print(f"[red]Telegram 处理失败: {e}[/]")

            if latest_seen > int(config.get("last_update_id", 0)):
                config["last_update_id"] = latest_seen
                _write_telegram_config(config)

    def _handle_update(self, update: dict):
        message = update.get("message")
        if not message:
            return

        chat = message.get("chat") or {}
        from_user = message.get("from") or {}
        text = message.get("text", "")
        if not text:
            return

        config = _read_telegram_config()
        token = config["bot_token"]
        chat_id = int(chat.get("id"))
        user_id = int(from_user.get("id", 0))
        message_id = message.get("message_id")
        normalized = _normalize_telegram_incoming_text(text, str(config.get("bot_username", "")))
        if normalized is None:
            return
        update_id = int(update.get("update_id", 0))
        preview = normalized.replace("\n", " ").strip()[:80]
        _telegram_log(f"telegram update={update_id} chat={chat_id} user={user_id} text={preview!r}")

        reply = None
        if normalized.startswith("/"):
            if _telegram_is_authorized(config, chat_id, user_id):
                with _TelegramTypingPulse(token, chat_id):
                    handled, reply = self._handle_command(normalized, config, chat_id, user_id)
                    if handled:
                        if reply:
                            _telegram_send_text(
                                token, chat_id, reply, reply_to_message_id=message_id
                            )
                        return
            else:
                handled, reply = self._handle_command(normalized, config, chat_id, user_id)
                if handled:
                    if reply:
                        _telegram_send_text(token, chat_id, reply, reply_to_message_id=message_id)
                    return

        if not _telegram_is_authorized(config, chat_id, user_id):
            _telegram_log(f"telegram unauthorized chat={chat_id} user={user_id}")
            _telegram_send_text(
                token,
                chat_id,
                _telegram_unauthorized_text(chat_id, user_id),
                reply_to_message_id=message_id,
            )
            return

        agent = self._get_agent(chat_id)
        reporter = _TelegramPhaseReporter(token, chat_id, reply_to_message_id=message_id)

        def _on_text(preview_text: str):
            reporter.capture_preface(preview_text)

        def _on_tool(event: dict):
            phase = event.get("phase")
            name = event.get("name", "")
            if phase == "start":
                reporter.tool_start(name)
            elif phase == "finish":
                reporter.tool_finish(name, event.get("preview", ""))
            elif phase == "error":
                reporter.tool_error(name, event.get("error", ""))

        with _TelegramTypingPulse(token, chat_id):
            reporter.start()
            started_at = time.time()
            try:
                reply = agent.chat(
                    normalized,
                    stream_to_console=False,
                    text_callback=_on_text,
                    tool_callback=_on_tool,
                ).strip()
            finally:
                reporter.stop()
            reply = reporter.strip_preface_from_reply(reply)
            if not reply:
                _telegram_log(
                    f"telegram empty_final chat={chat_id} preface_sent={reporter.preface_sent} preface_len={len(reporter.preface_text)}"
                )
                reply = "工具已执行，但 AI 没有返回最终正文。请重试一次；若仍复现，我会继续排查。"
            _telegram_send_text(token, chat_id, reply, reply_to_message_id=message_id)
            _telegram_log(
                f"telegram reply chat={chat_id} elapsed={time.time() - started_at:.1f}s len={len(reply)}"
            )

    def _handle_command(
        self, command_text: str, config: dict, chat_id: int, user_id: int
    ) -> tuple[bool, str]:
        parts = command_text.split(None, 1)
        head = parts[0].lower()
        arg = parts[1].strip() if len(parts) > 1 else ""
        authorized = _telegram_is_authorized(config, chat_id, user_id)

        if head in ("/start", "/help"):
            if authorized:
                return True, _telegram_help_text()
            return True, _telegram_unauthorized_text(chat_id, user_id) + "\n\n发送授权后再试。"

        if not authorized:
            return True, _telegram_unauthorized_text(chat_id, user_id)

        if head in ("/clear", "/new"):
            self._reset_chat(chat_id)
            return True, "当前 Telegram 会话上下文已清空。"

        if head == "/status":
            return True, _format_hw_status_for_text(stm32_hardware_status())

        if head == "/connect":
            result = stm32_connect(arg or None)
            return True, result.get("message", json.dumps(result, ensure_ascii=False))

        if head == "/disconnect":
            result = stm32_disconnect()
            return True, result.get("message", json.dumps(result, ensure_ascii=False))

        if head == "/chip":
            if not arg:
                return True, f"当前芯片: {_current_chip or '(未设置)'}"
            result = stm32_set_chip(arg)
            return True, f"已切换芯片: {result.get('chip')} ({result.get('family')})"

        if head == "/projects":
            result = stm32_list_projects()
            return True, _format_projects_for_text(result.get("projects", []))

        if head == "/serial":
            tokens = arg.split()
            if tokens and tokens[0] == "list":
                ports = detect_serial_ports()
                return True, "可用串口: " + (", ".join(ports) if ports else "无")
            port = tokens[0] if tokens and tokens[0].startswith("/dev/") else None
            baud = None
            for token in tokens:
                if token.isdigit():
                    baud = int(token)
                    break
            result = stm32_serial_connect(port, baud)
            return True, result.get("message", json.dumps(result, ensure_ascii=False))

        if head == "/telegram":
            return True, "\n".join(_telegram_status_lines(include_commands=False))

        return False, ""


_telegram_bridge = TelegramBotBridge()
_telegram_cli_autostart_done = False


def _ensure_cli_telegram_daemon() -> Optional[dict]:
    global _telegram_cli_autostart_done
    if _telegram_cli_autostart_done:
        return None
    _telegram_cli_autostart_done = True

    configured = _telegram_is_configured()
    ai_ready = _ai_is_configured()
    if not configured and not ai_ready:
        return None
    if configured and not ai_ready:
        return {
            "success": False,
            "message": _cli_text(
                "已检测到 Telegram 配置，但 AI 接口未配置，未自动启动",
                "Telegram is configured, but AI is not configured, so it was not started automatically",
            ),
        }
    if not configured:
        return None

    daemon = _telegram_daemon_status()
    if daemon["running"]:
        return {
            "success": True,
            "message": f"{_cli_text('已在后台运行', 'already running in background')} (PID {daemon.get('pid')})",
        }

    result = _start_telegram_daemon()
    return {
        "success": bool(result.get("success")),
        "message": str(result.get("message", "")).strip(),
    }


def serve_telegram_forever(daemonized: bool = False) -> int:
    if not _telegram_is_configured():
        CONSOLE.print("[red]Telegram 机器人未配置[/]")
        return 1
    if not _ai_is_configured():
        CONSOLE.print("[red]AI 接口未配置，无法启动 Telegram 机器人[/]")
        return 1

    result = _telegram_bridge.start()
    if not result["success"]:
        CONSOLE.print(f"[red]{result['message']}[/]")
        return 1

    if daemonized:
        _write_pid_file(os.getpid())
        atexit.register(lambda: _clear_pid_file(os.getpid()))

    CONSOLE.print("[green]Telegram 机器人正在运行，Ctrl+C 停止[/]")
    try:
        while _telegram_bridge.is_running():
            time.sleep(1.0)
    except KeyboardInterrupt:
        CONSOLE.print("\n[yellow]收到停止信号，正在退出 Telegram 机器人...[/]")
    finally:
        _telegram_bridge.stop()
        if daemonized:
            _clear_pid_file(os.getpid())
    return 0


def _interactive_telegram_menu() -> bool:
    CONSOLE.print()
    CONSOLE.rule("[bold cyan]  Telegram 管理[/]")
    CONSOLE.print()
    _print_telegram_status(include_commands=False)
    CONSOLE.print("[bold cyan]  可执行操作：[/]")
    CONSOLE.print("    [yellow]1[/]. 查看状态")
    CONSOLE.print("    [yellow]2[/]. 重新配置机器人")
    CONSOLE.print("    [yellow]3[/]. 启动后台机器人")
    CONSOLE.print("    [yellow]4[/]. 停止后台机器人")
    CONSOLE.print("    [yellow]5[/]. 添加白名单")
    CONSOLE.print("    [yellow]6[/]. 删除白名单")
    CONSOLE.print("    [yellow]7[/]. 允许所有 chat")
    CONSOLE.print("    [yellow]8[/]. 切回白名单模式")
    CONSOLE.print("    [yellow]9[/]. 重置 Telegram 配置")
    CONSOLE.print()

    try:
        choice = input("  输入序号（回车取消）: ").strip()
    except (EOFError, KeyboardInterrupt):
        CONSOLE.print("\n[dim]已取消[/]\n")
        return True

    if not choice:
        CONSOLE.print("[dim]已取消[/]\n")
        return True

    if choice == "1":
        _print_telegram_status(include_commands=True)
        return True

    if choice == "2":
        if not _ensure_ai_for_telegram():
            return True
        result = configure_telegram_cli()
        if result.get("success"):
            daemon = _telegram_daemon_status()
            if daemon["running"]:
                stop_result = _stop_telegram_daemon()
                start_result = _start_telegram_daemon()
                CONSOLE.print(
                    f"[{'green' if stop_result['success'] else 'red'}]{stop_result['message']}[/]"
                )
                CONSOLE.print(
                    f"[{'green' if start_result['success'] else 'red'}]{start_result['message']}[/]"
                )
                CONSOLE.print()
        return True

    if choice == "3":
        return handle_telegram_command("start", source="builtin")

    if choice == "4":
        return handle_telegram_command("stop", source="builtin")

    if choice in ("5", "6"):
        action = "allow" if choice == "5" else "remove"
        label = "添加" if choice == "5" else "删除"
        try:
            raw = input(f"  输入要{label}的 chat_id 或 user:123（可多个，空格分隔）: ").strip()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[dim]已取消[/]\n")
            return True
        if not raw:
            CONSOLE.print("[dim]已取消[/]\n")
            return True
        return handle_telegram_command(f"{action} {raw}", source="builtin")

    if choice == "7":
        return handle_telegram_command("allow-all", source="builtin")

    if choice == "8":
        return handle_telegram_command("whitelist", source="builtin")

    if choice == "9":
        try:
            confirm = input("  确认重置并删除 Telegram 配置？输入 yes 确认: ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[dim]已取消[/]\n")
            return True
        if confirm != "yes":
            CONSOLE.print("[dim]已取消[/]\n")
            return True
        return handle_telegram_command("reset", source="builtin")

    CONSOLE.print(f"[yellow]未知选项: {choice}[/]\n")
    return True


def handle_telegram_command(args: str, source: str = "cli") -> bool:
    arg = (args or "").strip()
    config = _read_telegram_config()
    parts = arg.split(None, 1)
    subcmd = parts[0].lower() if parts and parts[0] else ""
    subarg = parts[1].strip() if len(parts) > 1 else ""

    if subcmd == "":
        if source == "builtin":
            return _interactive_telegram_menu()
        if not _telegram_is_configured(config):
            CONSOLE.print("[yellow]Telegram 机器人尚未配置，进入配置向导[/]")
            if not _ensure_ai_for_telegram():
                return True
            result = configure_telegram_cli()
            if not result.get("success"):
                return True
            start_result = _start_telegram_daemon()
            CONSOLE.print(
                f"[{'green' if start_result['success'] else 'red'}]{start_result['message']}[/]"
            )
            CONSOLE.print()
            return True
        _print_telegram_status(include_commands=True)
        return True

    if subcmd in ("status", "info", "list", "ls"):
        _print_telegram_status(include_commands=True)
        return True

    if subcmd in ("config", "setup"):
        if not _ensure_ai_for_telegram():
            return True
        configure_telegram_cli()
        return True

    if subcmd in ("start", "run"):
        if not _telegram_is_configured(config):
            CONSOLE.print("[yellow]Telegram 机器人尚未配置，进入配置向导[/]")
            if not _ensure_ai_for_telegram():
                return True
            result = configure_telegram_cli()
            if not result.get("success"):
                return True
        elif not _ensure_ai_for_telegram():
            return True
        result = _start_telegram_daemon()
        CONSOLE.print(f"[{'green' if result['success'] else 'red'}]{result['message']}[/]")
        CONSOLE.print()
        return True

    if subcmd == "serve":
        daemonized = "--daemonized" in _split_tokens(subarg)
        sys.exit(serve_telegram_forever(daemonized=daemonized))

    if subcmd == "stop":
        result = _stop_telegram_daemon()
        CONSOLE.print(f"[{'green' if result['success'] else 'red'}]{result['message']}[/]")
        CONSOLE.print()
        return True

    if subcmd == "restart":
        stop_result = _stop_telegram_daemon()
        start_result = _start_telegram_daemon()
        CONSOLE.print(
            f"[{'green' if stop_result['success'] else 'red'}]{stop_result['message']}[/]"
        )
        CONSOLE.print(
            f"[{'green' if start_result['success'] else 'red'}]{start_result['message']}[/]"
        )
        CONSOLE.print()
        return True

    if subcmd in ("allow", "add"):
        if not subarg:
            CONSOLE.print("[yellow]用法: gary telegram allow <chat_id|user:123> [...][/]\n")
            return True
        parsed = _parse_telegram_targets(subarg)
        if parsed["invalid"]:
            CONSOLE.print(f"[yellow]忽略非法项: {parsed['invalid']}[/]")
        saved = _telegram_set_permissions(
            add_chat_ids=parsed["chat_ids"],
            add_user_ids=parsed["user_ids"],
        )
        CONSOLE.print("[green]Telegram 白名单已更新[/]")
        CONSOLE.print(f"  chat_id: {saved['allowed_chat_ids']}")
        CONSOLE.print(f"  user_id: {saved['allowed_user_ids']}\n")
        return True

    if subcmd in ("remove", "delete", "del", "rm"):
        if not subarg:
            CONSOLE.print("[yellow]用法: gary telegram remove <chat_id|user:123> [...][/]\n")
            return True
        parsed = _parse_telegram_targets(subarg)
        if parsed["invalid"]:
            CONSOLE.print(f"[yellow]忽略非法项: {parsed['invalid']}[/]")
        saved = _telegram_set_permissions(
            remove_chat_ids=parsed["chat_ids"],
            remove_user_ids=parsed["user_ids"],
        )
        CONSOLE.print("[green]Telegram 白名单已更新[/]")
        CONSOLE.print(f"  chat_id: {saved['allowed_chat_ids']}")
        CONSOLE.print(f"  user_id: {saved['allowed_user_ids']}\n")
        return True

    if subcmd == "allow-all":
        saved = _telegram_set_permissions(allow_all_chats=True)
        CONSOLE.print("[green]已切换为允许所有 chat[/]")
        CONSOLE.print(f"  chat_id: {saved['allowed_chat_ids']}")
        CONSOLE.print(f"  user_id: {saved['allowed_user_ids']}\n")
        return True

    if subcmd == "whitelist":
        saved = _telegram_set_permissions(allow_all_chats=False)
        CONSOLE.print("[green]已切换为白名单模式[/]")
        CONSOLE.print(f"  chat_id: {saved['allowed_chat_ids']}")
        CONSOLE.print(f"  user_id: {saved['allowed_user_ids']}\n")
        return True

    if subcmd == "reset":
        result = _reset_telegram_config()
        CONSOLE.print(f"[{'green' if result['success'] else 'red'}]{result['message']}[/]")
        CONSOLE.print()
        return True

    CONSOLE.print(f"[yellow]未知 Telegram 命令: {subcmd}[/]")
    _print_telegram_status(include_commands=True)
    return True


def _shutdown_cli_runtime(stop_telegram: bool = True) -> dict:
    """退出 CLI 时统一清理资源。"""
    results = {
        "hardware": stm32_disconnect(),
        "bridge_stopped": False,
        "telegram": {"success": True, "message": "Telegram 保持运行"},
    }
    try:
        _telegram_bridge.stop()
        results["bridge_stopped"] = True
    except Exception as e:
        results["bridge_stopped"] = False
        results["bridge_error"] = str(e)
    if stop_telegram:
        try:
            results["telegram"] = _stop_telegram_daemon()
        except Exception as e:
            results["telegram"] = {"success": False, "message": f"停止 Telegram 机器人失败: {e}"}
    return results


# ─────────────────────────────────────────────────────────────
# UI 常量
# ─────────────────────────────────────────────────────────────
CONSOLE = Console()
THEME = "cyan"
MAX_CONTEXT_TOKENS = 128000
MAX_TOOL_RESULT_LEN = 8000

# ─────────────────────────────────────────────────────────────
# 寄存器地址表（按系列）
# ─────────────────────────────────────────────────────────────
_REG_F1 = {
    "RCC_CR": 0x40021000,
    "RCC_CFGR": 0x40021004,
    "RCC_APB1ENR": 0x4002101C,
    "RCC_APB2ENR": 0x40021018,
    "GPIOA_CRL": 0x40010800,
    "GPIOA_CRH": 0x40010804,
    "GPIOA_IDR": 0x40010808,
    "GPIOA_ODR": 0x4001080C,
    "GPIOB_CRL": 0x40010C00,
    "GPIOB_CRH": 0x40010C04,
    "GPIOB_IDR": 0x40010C08,
    "GPIOB_ODR": 0x40010C0C,
    "GPIOC_CRL": 0x40011000,
    "GPIOC_CRH": 0x40011004,
    "TIM1_CR1": 0x40012C00,
    "TIM1_CCER": 0x40012C20,
    "TIM2_CR1": 0x40000000,
    "TIM2_CCER": 0x40000020,
    "TIM3_CR1": 0x40000400,
    "TIM3_CCER": 0x40000420,
    "ADC1_SR": 0x40012400,
    "ADC1_CR2": 0x40012408,
    "I2C1_CR1": 0x40005400,
    "I2C1_SR1": 0x40005414,
    "I2C2_CR1": 0x40005800,
    "I2C2_SR1": 0x40005814,
    "USART1_SR": 0x40013800,
    "USART1_BRR": 0x40013808,
}
_REG_F4 = {
    "RCC_CR": 0x40023800,
    "RCC_CFGR": 0x40023808,
    "RCC_AHB1ENR": 0x40023830,
    "RCC_APB1ENR": 0x40023840,
    "RCC_APB2ENR": 0x40023844,
    "GPIOA_MODER": 0x40020000,
    "GPIOA_IDR": 0x40020010,
    "GPIOA_ODR": 0x40020014,
    "GPIOB_MODER": 0x40020400,
    "GPIOB_IDR": 0x40020410,
    "GPIOB_ODR": 0x40020414,
    "GPIOC_MODER": 0x40020800,
    "GPIOC_IDR": 0x40020810,
    "GPIOC_ODR": 0x40020814,
    "TIM2_CR1": 0x40000000,
    "TIM2_CCER": 0x40000020,
    "TIM3_CR1": 0x40000400,
    "TIM3_CCER": 0x40000420,
    "I2C1_CR1": 0x40005400,
    "I2C1_SR1": 0x40005414,
    "USART1_SR": 0x40011000,
    "USART1_BRR": 0x40011008,
}
_REG_F0F3 = {
    "RCC_CR": 0x40021000,
    "RCC_CFGR": 0x40021004,
    "RCC_AHBENR": 0x40021014,
    "RCC_APB2ENR": 0x40021018,
    "RCC_APB1ENR": 0x4002101C,
    "GPIOA_MODER": 0x48000000,
    "GPIOA_IDR": 0x48000010,
    "GPIOA_ODR": 0x48000014,
    "GPIOB_MODER": 0x48000400,
    "GPIOB_IDR": 0x48000410,
    "GPIOB_ODR": 0x48000414,
    "I2C1_CR1": 0x40005400,
    "USART1_BRR": 0x40013808,
    "USART1_CR1": 0x4001380C,
}
_REG_COMMON = {
    "SCB_CFSR": 0xE000ED28,
    "SCB_HFSR": 0xE000ED2C,
    "SCB_BFAR": 0xE000ED38,
    "NVIC_ISER0": 0xE000E100,
}


def _reg_map(family: str) -> dict:
    base = {"f1": _REG_F1, "f4": _REG_F4, "f0": _REG_F0F3, "f3": _REG_F0F3}
    regs = dict(base.get(family.lower(), _REG_F1))
    regs.update(_REG_COMMON)
    return regs


# ─────────────────────────────────────────────────────────────
# PyOCDBridge（替换 OpenOCD）
# ─────────────────────────────────────────────────────────────
class PyOCDBridge:
    """
    使用 pyocd Python 库直接控制 STM32（无需启动 openocd 进程）。
    支持所有 CMSIS-DAP / ST-Link / J-Link USB 探针。
    安装：pip install pyocd
    """

    def __init__(self):
        self._session = None
        self._target = None
        self.connected = False
        self.chip_info: dict = {}
        self._family = "f1"
        self._reg_map: dict = _reg_map("f1")

    # ---- 内部工具 ----
    def _chip_to_pyocd_target(self, chip: str) -> str:
        """将 STM32F103C8T6 → stm32f103c8（去掉封装+温度后缀）"""
        import re

        name = chip.lower().strip()
        # STM32 命名末尾：封装字母(T/U/H/Y) + 温度等级数字(3/6/7)，如 T6/U3/H7
        name = re.sub(r"[a-z]\d$", "", name)
        return name

    _pyocd_target_cache: Optional[tuple] = None  # (float, set)
    _CACHE_TTL = 60.0  # 秒

    @classmethod
    def _get_all_pyocd_targets(cls) -> set:
        """获取所有可用 pyocd 目标，60秒内复用缓存"""
        now = time.time()
        if cls._pyocd_target_cache is not None:
            ts, cached = cls._pyocd_target_cache
            if now - ts < cls._CACHE_TTL:
                return cached

        try:
            result = subprocess.run(
                [sys.executable, "-m", "pyocd", "list", "--targets"],
                capture_output=True,
                text=True,
                timeout=15,
            )
            known = set()
            for line in result.stdout.splitlines():
                parts = line.split()
                if parts and parts[0].startswith("stm32"):
                    known.add(parts[0].lower())
            cls._pyocd_target_cache = (now, known)
            return known
        except Exception:
            return set()

    def _resolve_best_target(self, target_name: str) -> str:
        """在 pyocd 所有可用目标（builtin + pack）中找最佳匹配（精确→前缀→截短）"""
        known = self._get_all_pyocd_targets()

        # 若子进程查询失败，降级到 TARGET 字典
        if not known:
            try:
                from pyocd.target import TARGET

                known = {k.lower() for k in TARGET}
            except ImportError:
                return target_name

        # 1. 精确匹配
        if target_name in known:
            return target_name

        # 2. 前缀匹配：target_name 是 known 中某目标的前缀（不太可能，但防御性保留）
        candidates = [k for k in known if k.startswith(target_name)]
        if candidates:
            best = max(candidates, key=lambda k: len(os.path.commonprefix([target_name, k])))
            CONSOLE.print(f"[yellow]  目标映射: {target_name} → {best}[/]")
            return best

        # 3. 截短搜索：逐位去尾，找相同系列最接近的目标
        for trim in range(1, 4):
            prefix = target_name[:-trim]
            if len(prefix) < 8:  # stm32fXX 最短前缀
                break
            candidates = [k for k in known if k.startswith(prefix)]
            if candidates:
                best = max(candidates, key=lambda k: len(os.path.commonprefix([target_name, k])))
                CONSOLE.print(f"[yellow]  目标近似匹配: {target_name} → {best}[/]")
                return best

        return target_name  # 未找到，原样返回让 pyocd 报错

    def _auto_install_pack(self, target_name: str) -> bool:
        """自动安装 pyocd CMSIS pack，返回是否成功"""
        CONSOLE.print(f"[yellow]  未找到目标 {target_name}，正在自动安装支持包...[/]")
        try:
            result = subprocess.run(
                [sys.executable, "-m", "pyocd", "pack", "install", target_name],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if result.returncode == 0:
                CONSOLE.print(f"[green]  支持包安装成功[/]")
                return True
            else:
                CONSOLE.print(f"[red]  支持包安装失败: {result.stderr.strip()}[/]")
                return False
        except Exception as e:
            CONSOLE.print(f"[red]  支持包安装出错: {e}[/]")
            return False

    def _detect_family(self, chip: str) -> str:
        chip_up = chip.upper()
        if "F0" in chip_up:
            return "f0"
        if "F3" in chip_up:
            return "f3"
        if "F4" in chip_up or "F7" in chip_up or "H7" in chip_up:
            return "f4"
        return "f1"

    def set_family(self, family: str):
        self._family = family
        self._reg_map = _reg_map(family)

    # ---- 连接 / 断开 ----
    def start(self, chip: str = DEFAULT_CHIP) -> bool:
        """连接第一个可用探针，成功返回 True"""
        self.stop()
        try:
            from pyocd.core.helpers import ConnectHelper
        except ImportError:
            CONSOLE.print("[red]pyocd 未安装，请运行: pip install pyocd[/]")
            return False

        # chip 为 None 时让 pyocd 自动检测目标
        explicit_chip = chip and chip.upper() != "AUTO"
        family = self._detect_family(chip) if explicit_chip else "f1"
        self.set_family(family)

        if explicit_chip:
            raw_target = self._chip_to_pyocd_target(chip)
            # 在已知目标库中找最佳匹配（精确→前缀→截短）
            target_name = self._resolve_best_target(raw_target)
        else:
            target_name = None

        probe_hint = f"目标: {target_name}" if target_name else "自动检测目标"
        CONSOLE.print(f"[dim]  连接探针（{probe_hint}）...[/]")

        def _do_connect(t_name):
            return ConnectHelper.session_with_chosen_probe(
                target_override=t_name,
                auto_unlock=True,
                connect_mode="halt",
                blocking=False,
                return_first=True,
                options={"frequency": 1000000},
            )

        try:
            self._session = _do_connect(target_name)
        except Exception as e:
            err_str = str(e)
            # 目标不认识 → 自动安装 pack 后重试一次
            if explicit_chip and ("not recognized" in err_str or "Target type" in err_str):
                raw_target = self._chip_to_pyocd_target(chip)
                if self._auto_install_pack(raw_target):
                    # pack 安装后重新解析（新 session 会扫描已安装 pack）
                    target_name = self._resolve_best_target(raw_target)
                    CONSOLE.print(f"[dim]  重新连接（{target_name}）...[/]")
                    try:
                        self._session = _do_connect(target_name)
                    except Exception as e2:
                        CONSOLE.print(f"[red]  连接失败: {e2}[/]")
                        self._session = None
                        self._target = None
                        return False
                else:
                    CONSOLE.print(f"[red]  连接失败: {e}[/]")
                    self._session = None
                    self._target = None
                    return False
            else:
                CONSOLE.print(f"[red]  连接失败: {e}[/]")
                self._session = None
                self._target = None
                return False

        try:
            if self._session is None:
                CONSOLE.print("[red]  未找到调试探针，请检查 USB 连接[/]")
                return False
            self._session.open()
            self._target = self._session.board.target

            # 读取 pyocd 实际识别到的目标型号
            detected = getattr(self._target, "target_type", None) or (target_name or "unknown")
            # 若是自动检测，用检测到的型号更新 chip 变量
            resolved_chip = chip.upper() if explicit_chip else detected.upper()
            resolved_family = self._detect_family(resolved_chip)
            self.set_family(resolved_family)

            self.chip_info = {
                "device": resolved_chip,
                "pyocd_target": detected,
                "family": resolved_family,
                "probe": self._session.board.description,
            }
            self.connected = True
            CONSOLE.print(
                f"[green]  已连接: {resolved_chip} " f"| 探针: {self._session.board.description}[/]"
            )
            # Warmup：halt→读CPUID，稳定 SWD 会话，保持 halt 状态以便烧录
            try:
                self._target.halt()
                time.sleep(0.1)
                self._target.read32(0xE000ED00)  # CPUID，只读安全寄存器
                time.sleep(0.05)
                # 不 resume，保持 halt——烧录前必须 halt，reconnect 后也不例外
            except Exception:
                pass
            return True
        except Exception as e:
            CONSOLE.print(f"[red]  连接后处理失败: {e}[/]")
            self._session = None
            self._target = None
            return False

    def stop(self):
        if self._session:
            try:
                self._target.resume()
            except Exception:
                pass
            try:
                self._session.close()
            except Exception:
                pass
            self._session = None
            self._target = None
        self.connected = False

    # ---- 烧录 ----
    def flash(self, bin_path: str) -> dict:
        if not self.connected:
            return {"ok": False, "msg": "探针未连接，请先 connect"}
        p = Path(bin_path)
        if not p.exists():
            return {"ok": False, "msg": f"文件不存在: {bin_path}"}

        try:
            from pyocd.flash.file_programmer import FileProgrammer
        except ImportError:
            return {"ok": False, "msg": "pyocd 未安装"}

        size = p.stat().st_size
        t0 = time.time()
        CONSOLE.print(f"[dim]  烧录 {size} 字节...[/]")
        # 烧录前 reset_and_halt，将 MCU 恢复到干净复位状态再写 flash
        # 仅 halt() 不够——若前一个固件开了 IWDG 或 I2C 卡死，flash 算法也会被影响
        try:
            self._target.reset_and_halt()
            time.sleep(0.1)
        except Exception:
            try:
                self._target.halt()
                time.sleep(0.05)
            except Exception:
                pass
        try:
            programmer = FileProgrammer(self._session)
            programmer.program(str(p), base_address=0x08000000)
        except Exception as e:
            return {"ok": False, "msg": f"烧录异常: {e}"}

        dt = time.time() - t0
        spd = size / dt / 1024 if dt > 0 else 0
        # 烧录后复位并运行（失败不致命，固件已写入）
        try:
            self._target.reset_and_halt()
            time.sleep(0.1)
            self._target.resume()
        except Exception as e:
            CONSOLE.print(f"[yellow]  复位警告（固件已烧录）: {e}[/]")
        return {"ok": True, "msg": f"烧录成功 {size}B / {dt:.1f}s ({spd:.1f} KB/s)"}

    # ---- 寄存器读取 ----
    def read_registers(self, names: Optional[list] = None) -> Optional[dict]:
        if not self.connected:
            return None
        try:
            self._target.halt()
            time.sleep(REGISTER_READ_DELAY)

            targets = names if names else list(self._reg_map.keys())
            regs = {}
            for name in targets:
                addr = self._reg_map.get(name)
                if addr is None:
                    continue
                try:
                    val = self._target.read32(addr)
                    regs[name] = f"0x{val:08X}"
                except Exception:
                    pass

            # 读 PC
            try:
                pc = self._target.read_core_register("pc")
                regs["PC"] = f"0x{pc:08X}"
            except Exception:
                pass

            self._target.resume()
            return regs
        except Exception as e:
            CONSOLE.print(f"[red]  寄存器读取异常: {e}[/]")
            return None

    def read_all_for_debug(self) -> Optional[dict]:
        return self.read_registers()

    def analyze_fault(self, regs: dict) -> str:
        cfsr_str = regs.get("SCB_CFSR", "0x00000000")
        try:
            cfsr = int(cfsr_str, 16)
        except ValueError:
            return "CFSR 格式错误"
        if cfsr == 0:
            return "无故障"
        checks = [
            (0x01, "IACCVIOL: 指令访问违规"),
            (0x02, "DACCVIOL: 数据访问违规"),
            (0x100, "IBUSERR: 指令总线错误"),
            (0x200, "PRECISERR: 精确总线错误（外设未使能时钟）"),
            (0x400, "IMPRECISERR: 非精确总线错误"),
            (0x10000, "UNDEFINSTR: 未定义指令"),
            (0x20000, "INVSTATE: 无效 EPSR 状态"),
            (0x1000000, "UNALIGNED: 非对齐访问"),
            (0x2000000, "DIVBYZERO: 除零"),
        ]
        faults = [desc for mask, desc in checks if cfsr & mask]
        return "; ".join(faults) if faults else f"未知故障 CFSR=0x{cfsr:08X}"

    def list_probes(self) -> list:
        """列出所有可用探针"""
        try:
            from pyocd.core.session import Session
            from pyocd.probe.aggregator import DebugProbeAggregator

            probes = DebugProbeAggregator.get_all_connected_probes()
            return [{"uid": p.unique_id, "description": p.product_name} for p in probes]
        except Exception:
            return []


# ─────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────
# 串口自动检测
# ─────────────────────────────────────────────────────────────
def detect_serial_ports(verbose: bool = False) -> list:
    """
    跨平台扫描可用串口，按 STM32 使用优先级排序。
    Windows : CH340/CP210x → COMx（数字小的优先）
    macOS   : tty.usbserial-* / tty.usbmodem* → cu.* 备选
    Linux   : ttyUSB* / ttyACM* → ttyAMA* → ttyS[4+]
    只返回当前用户有读写权限的端口。
    """
    import glob, platform, re

    plat = platform.system()  # 'Windows' / 'Darwin' / 'Linux'
    found = []  # 有序去重列表，USB优先

    def _add(port: str, usb: bool = False):
        """加入列表，usb=True 时插到所有非usb端口之前"""
        if port in found:
            return
        if usb:
            # 找到第一个非usb端口的位置，插入其前
            idx = next((i for i, p in enumerate(found) if p not in _usb_set), len(found))
            found.insert(idx, port)
            _usb_set.add(port)
        else:
            found.append(port)

    _usb_set: set = set()

    # ── 1. pyserial list_ports（所有平台最可靠的来源）────────────
    try:
        import serial.tools.list_ports as lp

        skip_kw = ("bluetooth", "virtual", "rfcomm", "modem")
        for info in lp.comports():
            port = info.device
            desc = (info.description or "").lower()
            hwid = (info.hwid or "n/a").lower()
            if any(k in desc for k in skip_kw):
                continue
            # 跳过 hwid='n/a' 的 ttyS：是内核注册的幽灵串口，没有实际硬件
            if hwid == "n/a" and re.search(r"ttyS\d+$", port):
                continue
            is_usb = "usb" in hwid or "ch34" in hwid or "cp21" in hwid or "ft23" in hwid
            _add(port, usb=is_usb)
    except Exception:
        pass

    # ── 2. 平台专属补充扫描（pyserial 有时漏掉设备）───────────────
    if plat == "Windows":
        # Windows: 枚举 COM1-COM256，跳过已找到的
        try:
            import winreg

            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"HARDWARE\DEVICEMAP\SERIALCOMM")
            i = 0
            while True:
                try:
                    _, port, _ = winreg.EnumValue(key, i)
                    _add(port)
                    i += 1
                except OSError:
                    break
        except Exception:
            pass

    elif plat == "Darwin":
        # macOS: tty.usbserial / tty.usbmodem 优先，cu.* 备选
        for pattern, is_usb in [
            ("/dev/tty.usbserial*", True),
            ("/dev/tty.usbmodem*", True),
            ("/dev/tty.SLAB*", True),  # CP210x
            ("/dev/tty.wchusbserial*", True),  # CH340
            ("/dev/cu.usbserial*", True),
            ("/dev/cu.usbmodem*", True),
            ("/dev/tty.*", False),
        ]:
            for p in sorted(glob.glob(pattern)):
                if os.access(p, os.R_OK | os.W_OK):
                    _add(p, usb=is_usb)

    else:  # Linux / FreeBSD / 其他 POSIX
        # USB 转串口（最高优先）
        for pattern in ["/dev/ttyUSB*", "/dev/ttyACM*"]:
            for p in sorted(glob.glob(pattern)):
                real = os.path.realpath(p)
                if os.access(real, os.R_OK | os.W_OK):
                    _add(real, usb=True)
        # by-id 符号链接（解析后去重，仍属 USB 优先）
        for p in sorted(glob.glob("/dev/serial/by-id/*")):
            real = os.path.realpath(p)
            if os.access(real, os.R_OK | os.W_OK) and real not in found:
                _add(real, usb=True)

        # SBC 硬件 UART 补充（ttyAMA 始终可靠；ttyS 用 sysfs 验证有无实际硬件）
        def _has_sysfs_device(port: str) -> bool:
            """Linux: /sys/class/tty/ttySx/device 存在 = 真实硬件 UART"""
            name = os.path.basename(port)
            return os.path.exists(f"/sys/class/tty/{name}/device")

        for p in sorted(glob.glob("/dev/ttyAMA*")):
            if os.access(p, os.R_OK | os.W_OK) and p not in found:
                _add(p, usb=False)

        for p in sorted(glob.glob("/dev/ttyS*")):
            if p in found or not os.access(p, os.R_OK | os.W_OK):
                continue
            if _has_sysfs_device(p):  # 有 sysfs device 条目 = 真实 UART
                _add(p, usb=False)

    if verbose:
        CONSOLE.print(f"[dim]  检测到串口: {found if found else '无'}[/]")
    return found


def auto_open_serial(baud: int = SERIAL_BAUD) -> tuple:
    """
    自动检测并尝试打开第一个可用串口，返回 (port, serial_obj) 或 (None, None)。
    """
    try:
        import serial as pyserial
    except ImportError:
        return None, None

    candidates = detect_serial_ports()
    for port in candidates:
        try:
            s = pyserial.Serial(port, baud, timeout=0.3)
            s.close()
            return port, None  # 可以打开，返回 port
        except Exception:
            continue
    return None, None


# ─────────────────────────────────────────────────────────────
# SerialMonitor（与 hardware.py 保持一致）
# ─────────────────────────────────────────────────────────────
class SerialMonitor:
    def __init__(self):
        self._serial = None
        self._port = None
        self._buffer = ""
        self._lock = threading.Lock()
        self._thread = None
        self._running = False

    def open(self, port: str = None, baud: int = SERIAL_BAUD) -> bool:
        try:
            import serial as pyserial
        except ImportError:
            CONSOLE.print("[yellow]  pyserial 未安装: pip install pyserial[/]")
            return False

        # 确定要尝试的端口列表
        if port:
            # 指定了端口：先试指定的，失败再自动扫描
            candidates = [port] + [p for p in detect_serial_ports() if p != port]
        else:
            # 未指定：完全自动检测
            candidates = detect_serial_ports()
            if not candidates:
                CONSOLE.print("[yellow]  串口: 未检测到任何可用串口[/]")
                return False

        for p in candidates:
            try:
                self._serial = pyserial.Serial(p, baud, timeout=0.5)
                self._serial.reset_input_buffer()
                self._running = True
                self._thread = threading.Thread(target=self._reader, daemon=True)
                self._thread.start()
                self._port = p
                CONSOLE.print(f"[green]  串口: {p} @ {baud}[/]")
                return True
            except Exception as e:
                if p == candidates[-1]:
                    # 最后一个也失败了，给出有用提示（跨平台）
                    import platform as _plt

                    if _plt.system() == "Linux" and not os.access(p, os.R_OK | os.W_OK):
                        try:
                            import grp as _grp

                            grp_name = _grp.getgrgid(os.stat(p).st_gid).gr_name
                        except Exception:
                            grp_name = "dialout"
                        CONSOLE.print(
                            f"[yellow]  串口: {p} 权限不足 → "
                            f"sudo usermod -aG {grp_name} $USER && newgrp {grp_name}[/]"
                        )
                    else:
                        CONSOLE.print(f"[yellow]  串口打开失败: {e}[/]")
                # 继续尝试下一个
                continue
        return False

    def _reader(self):
        try:
            import serial as _pyserial

            _SerialException = _pyserial.SerialException
        except ImportError:
            _SerialException = OSError

        consecutive_errors = 0
        while self._running and self._serial:
            try:
                data = self._serial.read(1024)
                if data:
                    consecutive_errors = 0
                    with self._lock:
                        self._buffer += data.decode("utf-8", errors="ignore")
                        if len(self._buffer) > 8192:
                            self._buffer = self._buffer[-8192:]
            except _SerialException:
                # 串口物理断开
                CONSOLE.print("[yellow]  ⚠ 串口断开[/]")
                self._running = False
                break
            except Exception:
                consecutive_errors += 1
                if consecutive_errors > 10:
                    CONSOLE.print("[yellow]  ⚠ 串口持续异常，停止读取[/]")
                    self._running = False
                    break
                time.sleep(0.1)

    def read_and_clear(self) -> str:
        with self._lock:
            out = self._buffer
            self._buffer = ""
            return out

    def clear(self):
        with self._lock:
            self._buffer = ""

    def wait_for(self, keyword: str, timeout: float = 5.0, clear_first: bool = True) -> str:
        if clear_first:
            self.clear()
        t0 = time.time()
        while time.time() - t0 < timeout:
            with self._lock:
                if keyword in self._buffer:
                    break
            time.sleep(0.1)
        time.sleep(0.3)
        return self.read_and_clear()

    def close(self):
        self._running = False
        if self._serial:
            try:
                self._serial.close()
            except Exception:
                pass
            self._serial = None


def _wait_serial_adaptive(
    serial,
    keyword: str,
    min_wait: float = 0.5,
    max_wait: float = 8.0,
) -> str:
    """
    自适应串口等待：
    - 先等 min_wait 秒（给 MCU 复位时间）
    - 之后每 200ms 采样一次，检测到 keyword 或有内容即停
    - 超过 max_wait 后强制返回
    """
    time.sleep(min_wait)
    t0 = time.time()
    accumulated = ""
    while time.time() - t0 < (max_wait - min_wait):
        chunk = serial.read_and_clear()
        if chunk:
            accumulated += chunk
            if keyword in accumulated:
                break
        time.sleep(0.2)
    time.sleep(0.3)
    accumulated += serial.read_and_clear()
    return accumulated


# ─────────────────────────────────────────────────────────────
# 全局硬件状态（工具函数直接访问）
# ─────────────────────────────────────────────────────────────
_compiler: Optional[Compiler] = None
_bridge: Optional[PyOCDBridge] = None
_serial: Optional[SerialMonitor] = None
_hw_connected = False
_serial_connected = False
_current_chip = DEFAULT_CHIP
_last_bin_path: Optional[str] = None
_last_code: Optional[str] = None
_compiler_mtime: float = 0.0  # compiler.py 上次加载时的 mtime

# 调试闭环计数器（每次新任务由 AI 调用 stm32_reset_debug_attempts 重置）
_debug_attempt = 0
MAX_DEBUG_ATTEMPTS = 8  # 提高上限，一个任务最多 8 轮（含修改迭代）


def _get_compiler() -> Compiler:
    """返回 Compiler 单例；若 compiler.py 磁盘文件已更新则自动热重载。"""
    global _compiler, _compiler_mtime
    import importlib

    compiler_path = _HERE / "compiler.py"
    try:
        mtime = compiler_path.stat().st_mtime
    except OSError:
        mtime = 0.0
    if mtime > _compiler_mtime:
        importlib.reload(_compiler_module)
        globals()["Compiler"] = _compiler_module.Compiler
        _compiler = None  # 旧实例作废，下方重新创建
        _compiler_mtime = mtime
    if _compiler is None:
        _compiler = _compiler_module.Compiler()
        _compiler.check(_current_chip)  # 探测 GCC/HAL，设置 has_gcc/has_hal
    return _compiler


def _get_bridge() -> PyOCDBridge:
    global _bridge
    if _bridge is None:
        _bridge = PyOCDBridge()
    return _bridge


def _get_serial() -> SerialMonitor:
    global _serial
    if _serial is None:
        _serial = SerialMonitor()
    return _serial


# ─────────────────────────────────────────────────────────────
# STM32 专属工具实现
# ─────────────────────────────────────────────────────────────


def stm32_generate_font(text: str, size: int = 16) -> dict:
    """
    将任意文字（含中文）渲染为 STM32 OLED 用的 C 点阵数组。
    固定使用「横向取模·高位在前（row-major, MSB=left）」格式——
    这是最直观、与 SSD1306 逐行刷新最匹配的格式，配套显示函数一并生成。
    返回 c_code（字模数组 + 完整显示函数），直接粘贴进 main.c 使用。
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return {"success": False, "message": "需要安装 Pillow: pip install Pillow"}

    import platform as _plat

    # 跨平台字体候选
    def _find_cjk_font() -> Optional[str]:
        """动态查找系统 CJK 字体路径"""
        # 优先用 fc-match（Linux/macOS）
        try:
            r = subprocess.run(
                ["fc-match", "--format=%{file}", ":lang=zh"],
                capture_output=True,
                text=True,
                timeout=5,
            )
            if r.returncode == 0 and r.stdout.strip():
                path = r.stdout.strip()
                if os.path.exists(path):
                    return path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 回退到硬编码候选列表
        candidates = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None
        # 原来的 for fp in font_candidates: ... 全部替换为：

    font_path = _find_cjk_font()
    if font_path is None:
        return {"success": False, "message": "未找到中文字体，请安装 fonts-noto-cjk"}
    try:
        font = ImageFont.truetype(font_path, size)
    except Exception as e:
        return {"success": False, "message": f"字体加载失败 ({font_path}): {e}"}

    def _render_char(char: str) -> list:
        """渲染单个字符到 size×size 位图，返回 0/1 列表（行优先）"""
        img = Image.new("L", (size, size), 0)
        draw = ImageDraw.Draw(img)
        try:
            bbox = font.getbbox(char)
            char_w = bbox[2] - bbox[0]
            char_h = bbox[3] - bbox[1]
            # 水平居中，垂直顶对齐（避免因字体 metrics 差异被裁底部）
            ox = (size - char_w) // 2 - bbox[0]
            oy = -bbox[1]  # 让字符顶部与图像顶部对齐
        except Exception:
            ox, oy = 0, 0
        draw.text((ox, oy), char, fill=255, font=font)
        return [1 if p > 127 else 0 for p in img.getdata()]

    def _to_row_msb(pixels: list) -> list:
        """横向取模·高位在前：每行从左到右，bit7=最左列"""
        data = []
        bytes_per_row = (size + 7) // 8
        for row in range(size):
            for b in range(bytes_per_row):
                byte = 0
                for bit in range(8):
                    col = b * 8 + bit
                    if col < size and pixels[row * size + col]:
                        byte |= 1 << (7 - bit)
                data.append(byte)
        return data

    def _ascii_preview(pixels: list) -> str:
        lines = []
        for row in range(size):
            lines.append("".join("█" if pixels[row * size + col] else "." for col in range(size)))
        return "\n".join(lines)

    chars_data = []
    previews = []
    char_list = []
    for char in text:
        pixels = _render_char(char)
        data = _to_row_msb(pixels)
        chars_data.append(data)
        previews.append(_ascii_preview(pixels))
        char_list.append(char)

    bytes_per_char = size * ((size + 7) // 8)
    fname = f"FONT_{size}x{size}"

    # ── 字模数组 ──────────────────────────────────────────────
    char_entries = []
    for i, (char, data) in enumerate(zip(char_list, chars_data)):
        preview_lines = previews[i].split("\n")
        preview_comment = "  /* " + "  ".join(preview_lines[:4]) + " ... */"
        hex_str = ", ".join(f"0x{b:02X}" for b in data)
        char_repr = char if ord(char) < 128 else f"{char}(U+{ord(char):04X})"
        char_entries.append(f"    /* [{i}] '{char_repr}' */\n    {{{hex_str}}}")

    array_code = (
        f"/* ═══ 字模数据：横向取模·高位在前 {size}x{size}px ═══\n"
        f"   格式：每行 {(size+7)//8} 字节，bit7=最左列，共 {bytes_per_char} 字节/字符\n"
        f"   字符表: {' '.join(repr(c) for c in char_list)} */\n"
        f"static const uint8_t {fname}[][{bytes_per_char}] = {{\n"
        + ",\n".join(char_entries)
        + "\n};\n"
    )

    # ── 配套显示函数（与字模格式严格匹配）──────────────────────
    display_func = f"""
/* ═══ 配套显示函数（必须与上面字模数据一起使用）═══ */
/* idx: 字符在 {fname} 中的下标（按字符表顺序） */
/* x,y: OLED 列(0-127)和页起始行(0-63)         */
void OLED_ShowFont{size}(uint8_t x, uint8_t y, uint8_t idx) {{
    const uint8_t *p = {fname}[idx];
    uint8_t bytes_per_row = {(size+7)//8};
    for (uint8_t row = 0; row < {size}; row++) {{
        OLED_SetCursor(x, y + row);   /* 设置到目标行 */
        for (uint8_t b = 0; b < bytes_per_row; b++) {{
            uint8_t byte = p[row * bytes_per_row + b];
            for (int8_t bit = 7; bit >= 0; bit--) {{
                uint8_t col = b * 8 + (7 - bit);
                if (col < {size}) {{
                    OLED_DrawPixel(x + col, y + row, (byte >> bit) & 1);
                }}
            }}
        }}
    }}
}}
/* 用法示例：显示字符表第0个字符在 (0,0) 位置
   OLED_ShowFont{size}(0, 0, 0);  // 显示 '{char_list[0] if char_list else "?"}' */
"""

    # ── ASCII 预览（调试用）────────────────────────────────────
    preview_block = "\n\n".join(f"/* '{c}':\n{p} */" for c, p in zip(char_list, previews))

    c_code = array_code + display_func
    return {
        "success": True,
        "c_code": c_code,
        "preview": preview_block,  # ASCII 预览，可用于肉眼验证字形
        "char_count": len(text),
        "bytes_per_char": bytes_per_char,
        "font_size": size,
        "mode": "row_msb",
        "char_order": char_list,
    }


def stm32_list_probes() -> dict:
    """列出所有可用的调试探针（ST-Link / CMSIS-DAP / J-Link）"""
    probes = _get_bridge().list_probes()
    if not probes:
        return {"success": True, "probes": [], "message": "未检测到任何探针，请检查 USB 连接"}
    return {"success": True, "probes": probes}


def stm32_connect(chip: str = None) -> dict:
    """连接 STM32 硬件（pyocd 探针 + 串口监控）"""
    global _hw_connected, _serial_connected, _current_chip
    # chip=None → 使用当前已知芯片或默认芯片，避免 pyocd 退化到 generic cortex_m
    target_chip = chip or _current_chip or DEFAULT_CHIP
    bridge = _get_bridge()
    serial = _get_serial()

    if bridge.start(target_chip):
        _hw_connected = True
        # 用 pyocd 实际识别到的型号（自动检测场景下会与传入值不同）
        _current_chip = bridge.chip_info.get("device", target_chip)
        _get_compiler().set_chip(_current_chip)
        _serial_connected = serial.open()
        return {
            "success": True,
            "chip": _current_chip,
            "probe": bridge.chip_info.get("probe", ""),
            "serial_connected": _serial_connected,
            "message": f"硬件已连接: {_current_chip}",
        }
    _hw_connected = False
    return {"success": False, "message": "连接失败，请检查探针 USB 连接和驱动"}


def stm32_serial_connect(port: str = None, baud: int = None) -> dict:
    """
    单独连接/重连 UART 串口（不影响 pyocd 探针连接）。
    用于更换串口设备或在 stm32_connect 之后补充连接串口。
    """
    global _serial_connected
    serial = _get_serial()
    # 若已连接先关闭
    serial.close()
    _serial_connected = False

    use_baud = baud or SERIAL_BAUD
    # port=None → 自动检测；否则先试指定端口，失败时扫描其他
    _serial_connected = serial.open(port or None, use_baud)
    actual_port = getattr(serial, "_port", port or "自动检测")
    if _serial_connected:
        return {
            "success": True,
            "port": actual_port,
            "baud": use_baud,
            "message": f"串口已连接: {actual_port} @ {use_baud}",
        }
    candidates = detect_serial_ports()
    return {
        "success": False,
        "port": port,
        "baud": use_baud,
        "message": f"串口打开失败，可用端口: {candidates if candidates else '无'}",
    }


def stm32_serial_disconnect() -> dict:
    """断开串口（保留 pyocd 探针连接）"""
    global _serial_connected
    _get_serial().close()
    _serial_connected = False
    return {"success": True, "message": "串口已断开"}


def stm32_disconnect() -> dict:
    """断开硬件连接（释放探针和串口）"""
    global _hw_connected, _serial_connected
    _get_bridge().stop()
    _get_serial().close()
    _hw_connected = False
    _serial_connected = False
    return {"success": True, "message": "已断开"}


def stm32_set_chip(chip: str) -> dict:
    """切换目标芯片型号（如 STM32F103C8T6 / STM32F407VET6）"""
    global _current_chip
    _current_chip = chip.strip().upper()
    ci = _get_compiler().set_chip(_current_chip)
    if _hw_connected:
        _get_bridge().set_family(ci.get("family", "f1"))
    return {"success": True, "chip": _current_chip, "family": ci.get("family", "f1")}


def stm32_hardware_status() -> dict:
    """获取当前硬件连接状态和工具链可用性"""
    ci = _get_compiler().check(_current_chip)
    return {
        "chip": _current_chip,
        "hw_connected": _hw_connected,
        "serial_connected": _serial_connected,
        "gcc_ok": ci.get("gcc", False),
        "gcc_version": ci.get("gcc_version", "未找到"),
        "hal_ok": ci.get("hal", False),
        "hal_lib_ok": ci.get("hal_lib", False),
        "workspace": str(WORKSPACE),
    }


def stm32_compile(code: str, chip: str = None) -> dict:
    """编译 STM32 C 代码（完整 main.c）"""
    global _last_bin_path, _last_code
    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())
    result = compiler.compile(code)
    if result["ok"]:
        _last_code = code
        _last_bin_path = result.get("bin_path")
        # 自动保存到 latest_workspace
        try:
            latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")
        except Exception as e:
            CONSOLE.print(f"[dim]  ⚠ 缓存保存失败: {e}[/]")
    payload = {
        "success": result["ok"],
        "message": (result.get("msg") or "")[:600],
        "bin_path": result.get("bin_path"),
        "bin_size": result.get("bin_size", 0),
    }
    if payload["success"]:
        _record_success_memory("compile_success", code=code, result=payload)
    return payload


def stm32_compile_rtos(code: str, chip: str = None) -> dict:
    """编译带 FreeRTOS 内核的完整 main.c 代码"""
    global _last_bin_path, _last_code
    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())
    result = compiler.compile_rtos(code)
    if result["ok"]:
        _last_code = code
        _last_bin_path = result.get("bin_path")
        try:
            latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
            latest.mkdir(parents=True, exist_ok=True)
            (latest / "main.c").write_text(code, encoding="utf-8")
        except Exception as e:
            CONSOLE.print(f"[dim]  ⚠ 缓存保存失败: {e}[/]")
    payload = {
        "success": result["ok"],
        "message": (result.get("msg") or "")[:600],
        "bin_path": result.get("bin_path"),
        "bin_size": result.get("bin_size", 0),
    }
    if payload["success"]:
        _record_success_memory("compile_success", code=code, result=payload)
    return payload


def stm32_recompile(mode: str = "auto") -> dict:
    """从 latest_workspace/main.c 直接重编译，无需 Gary 传递代码字符串。
    str_replace_edit 修改文件后调用此函数代替 read_file + stm32_compile。
    mode: "bare"(裸机) | "rtos"(FreeRTOS) | "auto"(自动检测，默认)
    """
    latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
    main_c = latest / "main.c"
    if not main_c.exists():
        return {"success": False, "message": "latest_workspace/main.c 不存在，请先编译一次完整代码"}
    code = main_c.read_text(encoding="utf-8")
    if mode == "auto":
        mode = (
            "rtos"
            if ("FreeRTOS.h" in code or "task.h" in code or "xTaskCreate" in code)
            else "bare"
        )
    if mode == "rtos":
        return stm32_compile_rtos(code)
    return stm32_compile(code)


def stm32_regen_bsp(chip: str = None) -> dict:
    """强制重新生成 startup.s / link.ld / FreeRTOSConfig.h。
    每次修改 compiler.py 或切换芯片后调用一次，确保构建文件是最新版本。
    自动检查 startup.s 是否包含 FPU 使能代码（Cortex-M4 必须）。
    """
    import importlib

    # 强制重载 compiler 模块，绕过进程内缓存
    importlib.reload(_compiler_module)
    globals()["Compiler"] = _compiler_module.Compiler
    global _compiler, _compiler_mtime
    _compiler = None
    try:
        compiler_path = _HERE / "compiler.py"
        _compiler_mtime = compiler_path.stat().st_mtime
    except OSError:
        _compiler_mtime = 0.0

    compiler = _get_compiler()
    if chip:
        compiler.set_chip(chip.strip().upper())

    # 调用 set_chip 触发 startup.s / link.ld 重新生成
    chip_name = chip.strip().upper() if chip else _current_chip
    ci = compiler.set_chip(chip_name)
    if ci is None:
        return {"success": False, "message": f"未知芯片: {chip_name}"}

    # 验证 startup.s 包含 FPU 使能
    from config import BUILD_DIR

    startup_path = BUILD_DIR / "startup.s"
    has_fpu_enable = False
    if startup_path.exists():
        content = startup_path.read_text()
        has_fpu_enable = "Enable FPU" in content or "CPACR" in content

    cpu = ci.get("cpu", "?")
    fpu = ci.get("fpu", False)
    ram_k = ci.get("ram_k", 0)
    flash_k = ci.get("flash_k", 0)

    warnings = []
    if fpu and not has_fpu_enable:
        warnings.append("⚠ startup.s 缺少 FPU 使能代码——请更新 compiler.py")
    if ram_k < 32 and ci.get("family") in ("f4",):
        warnings.append(f"⚠ RAM 仅 {ram_k}KB，FreeRTOS heap 可能不够")

    return {
        "success": True,
        "chip": chip_name,
        "cpu": cpu,
        "fpu_supported": fpu,
        "startup_fpu_enabled": has_fpu_enable,
        "flash_k": flash_k,
        "ram_k": ram_k,
        "files_regenerated": ["startup.s", "link.ld"],
        "warnings": warnings,
        "message": f"BSP 文件已重新生成 ({cpu}, {'有 FPU' if fpu else '无 FPU'}，{flash_k}K Flash / {ram_k}K RAM)"
        + (("；" + "；".join(warnings)) if warnings else ""),
    }


def stm32_analyze_fault_rtos() -> dict:
    """读取并分析 FreeRTOS 程序的 HardFault 寄存器。
    在普通 stm32_analyze_fault 基础上增加 FreeRTOS 专项诊断：
    - 识别 FPU 未使能导致的 PRECISERR
    - 识别 SysTick_Handler 冲突
    - 识别任务栈溢出引发的故障
    - 检查 startup.s 是否包含 FPU 初始化
    返回明确的根本原因和修复建议。
    """
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接"}

    bridge = _get_bridge()
    regs = bridge.read_registers(
        ["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC", "SCB_MMFAR", "SCB_CPACR_FIELD"]
    )
    if not regs:
        regs = bridge.read_registers(["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC"])
    if not regs:
        return {"success": False, "message": "寄存器读取失败，请确认硬件已连接并已暂停"}

    cfsr = regs.get("SCB_CFSR", "0x0")
    hfsr = regs.get("SCB_HFSR", "0x0")
    bfar = regs.get("SCB_BFAR", "0x0")
    pc = regs.get("PC", "0x0")

    try:
        cfsr_v = int(cfsr, 16)
        hfsr_v = int(hfsr, 16)
        bfar_v = int(bfar, 16)
        pc_v = int(pc, 16)
    except (ValueError, TypeError):
        cfsr_v = hfsr_v = bfar_v = pc_v = 0

    # 检查 startup.s 的 FPU 使能
    from config import BUILD_DIR

    startup_path = BUILD_DIR / "startup.s"
    has_fpu_in_startup = False
    if startup_path.exists():
        txt = startup_path.read_text()
        has_fpu_in_startup = "CPACR" in txt or "Enable FPU" in txt

    # 检查 CPACR 硬件状态（读 0xE000ED88 的 bit[23:20]）
    fpu_hw_enabled = False
    try:
        cpacr = bridge._target.read32(0xE000ED88)
        fpu_hw_enabled = ((cpacr >> 20) & 0xF) == 0xF
    except Exception:
        pass

    # ── 诊断逻辑 ──────────────────────────────────────────
    root_cause = "未知"
    fix = "请读取 SCB_CFSR 后手动分析"
    severity = "unknown"

    preciserr = bool(cfsr_v & (1 << 9))
    bfarvalid = bool(cfsr_v & (1 << 15))
    iaccviol = bool(cfsr_v & (1 << 0))
    undefinstr = bool(cfsr_v & (1 << 16))
    stkovf = bool(cfsr_v & (1 << 12))  # STKOVF: 栈下溢
    unstkovf = bool(cfsr_v & (1 << 11))  # UNSTKOVF: 不稳定栈下溢

    # 判断 BFAR 是否"合理"（STM32 合法地址范围）
    bfar_valid_range = (
        0x08000000 <= bfar_v < 0x08200000
        or 0x20000000 <= bfar_v < 0x20030000
        or 0x40000000 <= bfar_v < 0x60000000
        or 0xE0000000 <= bfar_v
    )
    bfar_garbage = bfarvalid and not bfar_valid_range

    if preciserr and bfar_garbage and not fpu_hw_enabled:
        root_cause = (
            "FPU 未使能：ARM_CM4F port.c 的 PendSV_Handler 执行 vpush/vpop 时触发 PRECISERR"
        )
        fix = (
            "修复方法（已在 compiler.py 修复，重新编译即可）：\n"
            "  1. 调用 stm32_regen_bsp() 重新生成 startup.s（含 CPACR 初始化）\n"
            "  2. 重新编译：stm32_compile_rtos(code)\n"
            "  不需要在代码里手动写 SCB->CPACR"
        )
        severity = "critical"
    elif undefinstr and not fpu_hw_enabled:
        root_cause = "FPU 指令在 FPU 禁用状态下执行（UNDEFINSTR）"
        fix = "同上，重新生成 startup.s 后重新编译"
        severity = "critical"
    elif stkovf or unstkovf:
        root_cause = "任务栈溢出（FreeRTOS 检测到栈越界）"
        fix = "增大任务的 stack_words 参数：普通任务 ≥128，FPU 任务 ≥256，含 printf 的任务 ≥384"
        severity = "high"
    elif preciserr and bfarvalid:
        root_cause = f"总线错误：精确地址 BFAR={bfar} 非法（可能是外设时钟未开启或访问了无效内存）"
        fix = "检查 PC 指向的函数，确认相关外设 RCC 时钟已 __HAL_RCC_xxx_CLK_ENABLE()"
        severity = "medium"
    elif hfsr_v & (1 << 1):
        root_cause = "向量表读取失败（VECTTBL）：中断向量地址无效"
        fix = "检查链接脚本 FLASH 地址和向量表对齐"
        severity = "critical"
    elif cfsr_v == 0 and hfsr_v == 0:
        root_cause = "无 HardFault（程序正常运行）"
        fix = "无需修复"
        severity = "none"

    # 补充 startup 检查信息
    startup_note = ""
    if not has_fpu_in_startup:
        startup_note = "⚠ 当前 startup.s 不含 FPU 使能代码，调用 stm32_regen_bsp() 后重新编译"
    elif not fpu_hw_enabled:
        startup_note = "⚠ startup.s 含 FPU 使能但硬件 CPACR 当前为 0（可能是旧固件未重新编译）"

    return {
        "success": True,
        "registers": regs,
        "fpu_hw_enabled": fpu_hw_enabled,
        "startup_has_fpu_enable": has_fpu_in_startup,
        "root_cause": root_cause,
        "severity": severity,
        "fix": fix,
        "startup_note": startup_note,
        "cfsr_bits": {
            "PRECISERR": preciserr,
            "BFARVALID": bfarvalid,
            "IACCVIOL": iaccviol,
            "UNDEFINSTR": undefinstr,
            "STKOVF": stkovf,
        },
    }


def stm32_rtos_check_code(code: str) -> dict:
    """FreeRTOS 代码静态检查 —— 编译前捕获常见 RTOS 编程错误。
    检查 SysTick 冲突、HAL_Delay 陷阱、缺少 hook 函数、栈大小、ISR 安全等。
    """
    import re

    errors = []
    warnings = []
    suggestions = []

    # 1. SysTick_Handler 冲突
    if re.search(r"\bvoid\s+SysTick_Handler\b", code):
        errors.append(
            "❌ 禁止自定义 SysTick_Handler —— FreeRTOS 已通过 xPortSysTickHandler 接管 SysTick。"
            "删除 SysTick_Handler，改用 vApplicationTickHook() 维持 HAL_IncTick()"
        )

    # 2. HAL_Delay 在 xTaskCreate 之前
    hal_delay_pos = code.find("HAL_Delay")
    xtask_pos = code.find("xTaskCreate")
    if hal_delay_pos >= 0 and xtask_pos >= 0 and hal_delay_pos < xtask_pos:
        # 排除注释中的情况（简单排除）
        line_with_delay = code[:hal_delay_pos].rfind("\n")
        delay_line = code[line_with_delay:hal_delay_pos]
        if "//" not in delay_line and "/*" not in delay_line:
            errors.append(
                "❌ HAL_Delay() 在 xTaskCreate() 之前调用 —— SysTick 会触发 FreeRTOS tick handler，"
                "访问未初始化的任务列表导致 HardFault。"
                "把含 HAL_Delay 的初始化移到任务函数内部"
            )

    # 3. 必需 hook 函数
    required_hooks = {
        "vApplicationTickHook": "void vApplicationTickHook(void) { HAL_IncTick(); }",
        "vApplicationMallocFailedHook": "void vApplicationMallocFailedHook(void) { while(1); }",
        "vApplicationStackOverflowHook": "void vApplicationStackOverflowHook(TaskHandle_t t, char *n) { while(1); }",
        "vApplicationIdleHook": "void vApplicationIdleHook(void) {}",
    }
    for hook, template in required_hooks.items():
        if hook not in code:
            errors.append(f"❌ 缺少 {hook} —— 请添加: {template}")

    # 4. 任务栈大小检查
    task_creates = re.findall(r'xTaskCreate\s*\(\s*(\w+)\s*,\s*"[^"]*"\s*,\s*(\d+)', code)
    for func_name, stack_str in task_creates:
        stack = int(stack_str)
        # 查找任务函数体
        func_pattern = rf"void\s+{re.escape(func_name)}\s*\("
        func_match = re.search(func_pattern, code)
        if func_match:
            # 提取函数体（简单：从匹配位置到后续 2000 字符）
            func_body = code[func_match.start() : func_match.start() + 2000]
            has_float = any(
                kw in func_body
                for kw in [
                    "float ",
                    "double ",
                    "sinf(",
                    "cosf(",
                    "sqrtf(",
                    "tanf(",
                    "arm_",
                    ".0f",
                    "fabsf(",
                    "powf(",
                    "logf(",
                    "expf(",
                ]
            )
            has_printf = any(kw in func_body for kw in ["snprintf(", "sprintf(", "printf("])
            if has_printf and stack < 384:
                warnings.append(f"⚠ 任务 {func_name} 使用 snprintf 但栈仅 {stack} words，建议 ≥384")
            elif has_float and stack < 256:
                warnings.append(f"⚠ 任务 {func_name} 使用浮点运算但栈仅 {stack} words，建议 ≥256")
            elif stack < 128:
                warnings.append(f"⚠ 任务 {func_name} 栈仅 {stack} words，建议 ≥128")

    # 5. ISR 安全检查 —— 在 IRQHandler 中使用非 FromISR 的 API
    irq_funcs = re.findall(r"void\s+(\w+_IRQHandler)\s*\(void\)\s*\{", code)
    for irq_name in irq_funcs:
        # 提取 IRQ 函数体
        irq_start = code.find(f"void {irq_name}")
        if irq_start >= 0:
            irq_body = code[irq_start : irq_start + 1500]
            # 检查危险 API（非 FromISR 版本）
            unsafe_apis = [
                "xQueueSend(",
                "xQueueReceive(",
                "xSemaphoreTake(",
                "xSemaphoreGive(",
                "vTaskDelay(",
                "vTaskDelayUntil(",
                "xTaskCreate(",
                "vTaskDelete(",
                "printf(",
            ]
            for api in unsafe_apis:
                if api in irq_body and api.replace("(", "FromISR(") not in irq_body:
                    safe_api = api.replace("(", "FromISR(")
                    warnings.append(f"⚠ {irq_name} 中使用了 {api} —— ISR 中必须用 {safe_api}")

    # 6. 缺少头文件
    if "strlen(" in code or "memset(" in code or "memcpy(" in code:
        if "#include <string.h>" not in code and "#include<string.h>" not in code:
            warnings.append("⚠ 使用了 strlen/memset/memcpy 但未 #include <string.h>")
    if "sinf(" in code or "cosf(" in code or "sqrtf(" in code:
        if "#include <math.h>" not in code and "#include<math.h>" not in code:
            warnings.append("⚠ 使用了 sinf/cosf/sqrtf 但未 #include <math.h>")
    if "snprintf(" in code:
        if "#include <stdio.h>" not in code and "#include<stdio.h>" not in code:
            warnings.append("⚠ 使用了 snprintf 但未 #include <stdio.h>")

    # 7. 建议
    if "xSemaphoreCreateBinary" in code and "vTaskNotifyGive" not in code:
        suggestions.append(
            "💡 考虑用任务通知 (vTaskNotifyGive/ulTaskNotifyTake) 替代二值信号量，速度更快且更省内存"
        )
    if (
        "xEventGroupCreate" not in code
        and re.findall(r"xSemaphoreCreateBinary", code).__len__() >= 3
    ):
        suggestions.append("💡 多个二值信号量可能适合用事件组 (xEventGroupCreate) 替代")
    if "vTaskDelayUntil" not in code and "vTaskDelay" in code:
        suggestions.append(
            "💡 需要精确周期执行时用 vTaskDelayUntil 替代 vTaskDelay（避免累积漂移）"
        )

    return {
        "success": True,
        "errors": errors,
        "warnings": warnings,
        "suggestions": suggestions,
        "error_count": len(errors),
        "warning_count": len(warnings),
        "message": (
            f"检查完成: {len(errors)} 个错误, {len(warnings)} 个警告, {len(suggestions)} 个建议"
            + (
                "\n" + "\n".join(errors + warnings + suggestions)
                if errors or warnings or suggestions
                else "\n✅ 代码通过所有 RTOS 检查"
            )
        ),
    }


def stm32_rtos_task_stats() -> dict:
    """通过 pyocd 读取 FreeRTOS 运行时任务统计信息。
    从 ELF 符号表解析全局变量地址，读取任务数、堆使用、当前任务等。
    需要先编译（有 ELF 文件）并连接硬件。
    """
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接"}

    from config import BUILD_DIR

    elf_path = BUILD_DIR / "firmware.elf"
    if not elf_path.exists():
        return {"success": False, "message": "ELF 文件不存在，请先编译"}

    bridge = _get_bridge()

    # 通过 nm 解析 ELF 符号表获取变量地址
    import subprocess

    symbols_to_find = {
        "uxCurrentNumberOfTasks": None,  # uint32_t 任务数
        "xFreeBytesRemaining": None,  # size_t heap_4 剩余
        "xMinimumEverFreeBytesRemaining": None,  # size_t 历史最低
        "pxCurrentTCB": None,  # TCB* 当前任务
    }

    try:
        r = subprocess.run(
            ["arm-none-eabi-nm", "-g", str(elf_path)], capture_output=True, text=True, timeout=5
        )
        if r.returncode != 0:
            return {"success": False, "message": "nm 解析 ELF 失败"}

        for line in r.stdout.split("\n"):
            parts = line.strip().split()
            if len(parts) >= 3:
                addr_str, _typ, name = parts[0], parts[1], parts[2]
                if name in symbols_to_find:
                    symbols_to_find[name] = int(addr_str, 16)
    except Exception as e:
        return {"success": False, "message": f"符号解析异常: {e}"}

    result = {"success": True}

    # 读取各变量值
    try:
        addr = symbols_to_find["uxCurrentNumberOfTasks"]
        if addr:
            result["task_count"] = bridge._target.read32(addr)

        addr = symbols_to_find["xFreeBytesRemaining"]
        if addr:
            result["heap_free_bytes"] = bridge._target.read32(addr)

        addr = symbols_to_find["xMinimumEverFreeBytesRemaining"]
        if addr:
            result["heap_min_ever_free"] = bridge._target.read32(addr)

        addr = symbols_to_find["pxCurrentTCB"]
        if addr:
            tcb_ptr = bridge._target.read32(addr)
            result["current_tcb_addr"] = f"0x{tcb_ptr:08X}"

            # 尝试读取当前任务名（TCB 偏移量 52 处是 pcTaskName，长度 configMAX_TASK_NAME_LEN=16）
            try:
                name_offset = 52  # 标准 FreeRTOS TCB pcTaskName 偏移
                name_bytes = bytearray()
                for i in range(16):
                    b = bridge._target.read8(tcb_ptr + name_offset + i)
                    if b == 0:
                        break
                    name_bytes.append(b)
                result["current_task_name"] = name_bytes.decode("ascii", errors="replace")
            except Exception:
                pass

    except Exception as e:
        result["read_error"] = str(e)

    # 构造摘要消息
    parts = []
    if "task_count" in result:
        parts.append(f"任务数: {result['task_count']}")
    if "heap_free_bytes" in result:
        parts.append(f"堆剩余: {result['heap_free_bytes']}B")
    if "heap_min_ever_free" in result:
        parts.append(f"堆历史最低: {result['heap_min_ever_free']}B")
    if "current_task_name" in result:
        parts.append(f"当前任务: {result['current_task_name']}")
    result["message"] = " | ".join(parts) if parts else "读取完成（部分符号未找到）"

    return result


def stm32_rtos_suggest_config(
    task_count: int, use_fpu: bool = False, use_printf: bool = False, ram_k: int = 0
) -> dict:
    """根据用户的任务需求，计算推荐的 FreeRTOS 配置参数。
    估算堆大小、栈大小、优先级数，并检查 RAM 是否足够。
    """
    compiler = _get_compiler()
    ci = compiler._chip_info
    if ci is None:
        return {"success": False, "message": "芯片未设置，请先 stm32_set_chip"}

    actual_ram_k = ram_k if ram_k > 0 else ci.get("ram_k", 64)
    has_fpu = ci.get("fpu", False)

    # 栈大小推荐
    if use_printf:
        recommended_stack = 384
        stack_reason = "含 snprintf/printf，需要 384+ words"
    elif use_fpu or has_fpu:
        recommended_stack = 256
        stack_reason = "FPU 上下文保存需要 256+ words"
    else:
        recommended_stack = 128
        stack_reason = "普通任务最小 128 words"

    # 堆大小估算
    tcb_size = 92  # TCB 约 92 字节
    stack_bytes = recommended_stack * 4  # words → bytes
    per_task = tcb_size + stack_bytes
    idle_task = per_task  # Idle 任务
    timer_task = tcb_size + (recommended_stack * 2 * 4)  # Timer 任务栈更大

    total_task_mem = per_task * task_count + idle_task + timer_task
    # 额外开销：队列/信号量/事件组 约 20%
    recommended_heap = int(total_task_mem * 1.3)
    # 对齐到 256B
    recommended_heap = ((recommended_heap + 255) // 256) * 256

    max_heap = actual_ram_k * 1024 // 2  # 最多用一半 RAM
    if recommended_heap > max_heap:
        recommended_heap = max_heap

    # RAM 使用估算
    static_overhead = 4096  # .data + .bss + ISR 栈 约 4KB
    total_ram_usage = static_overhead + recommended_heap
    ram_pct = total_ram_usage * 100 // (actual_ram_k * 1024)

    warnings = []
    if ram_pct > 85:
        warnings.append(f"⚠ RAM 使用率预估 {ram_pct}%，建议减少任务数或栈大小")
    if task_count > 7:
        warnings.append("⚠ 任务数较多，注意优先级分配避免优先级反转")
    if actual_ram_k < 16 and task_count > 2:
        warnings.append(f"⚠ RAM 仅 {actual_ram_k}KB，建议最多 2 个任务")

    config = {
        "success": True,
        "recommended_stack_words": recommended_stack,
        "stack_reason": stack_reason,
        "recommended_heap_bytes": recommended_heap,
        "recommended_priorities": min(task_count + 2, 7),
        "estimated_ram_usage_bytes": total_ram_usage,
        "estimated_ram_percent": ram_pct,
        "ram_total_kb": actual_ram_k,
        "per_task_overhead_bytes": per_task,
        "warnings": warnings,
        "message": (
            f"推荐配置: 栈={recommended_stack}words, 堆={recommended_heap}B, "
            f"优先级={min(task_count + 2, 7)}, RAM预估使用{ram_pct}%"
            + (("  " + "; ".join(warnings)) if warnings else "")
        ),
    }
    return config


def stm32_rtos_plan_project(
    description: str, peripherals: list = None, task_hints: list = None
) -> dict:
    """FreeRTOS 项目规划工具 —— 复杂 RTOS 项目的架构规划。

    根据用户需求描述，生成结构化的项目规划：
    - 任务分解（任务名、职责、栈大小、优先级）
    - 通信拓扑（任务间用什么机制通信：Queue/Semaphore/Notification/EventGroup）
    - 中断处理策略（哪些中断需要处理，如何通知任务）
    - 外设分配（哪个任务负责哪些外设）
    - 时序约束（哪些操作有实时性要求）
    - 资源估算（总堆/栈/RAM 使用）

    AI 在规划复杂 RTOS 项目时必须先调用此工具，待用户确认后再写代码。
    """
    import re

    compiler = _get_compiler()
    ci = compiler._chip_info
    if ci is None:
        return {"success": False, "message": "芯片未设置，请先 stm32_set_chip"}

    has_fpu = ci.get("fpu", False)
    ram_k = ci.get("ram_k", 64)
    flash_k = ci.get("flash_k", 256)
    cpu = ci.get("cpu", "cortex-m3")

    # ── 分析需求 ──────────────────────────────────────────
    desc_lower = description.lower()

    # 检测是否涉及浮点
    uses_float = any(
        kw in desc_lower
        for kw in [
            "浮点",
            "float",
            "sin",
            "cos",
            "pid",
            "温度计算",
            "dsp",
            "adc采样",
            "滤波",
            "fft",
            "角度",
            "加速度",
            "陀螺仪",
        ]
    )

    # 检测是否需要 printf/snprintf
    uses_printf = any(
        kw in desc_lower for kw in ["printf", "snprintf", "格式化", "打印浮点", "调试输出"]
    )

    # 检测外设
    detected_peripherals = []
    periph_map = {
        "uart": ["串口", "uart", "usart", "通信", "打印", "日志"],
        "i2c": ["i2c", "oled", "传感器", "bmp", "sht", "mpu", "加速度", "陀螺仪", "屏幕"],
        "spi": ["spi", "sd卡", "flash", "w25q", "tft", "lcd"],
        "adc": ["adc", "模拟", "采样", "电压", "温度", "光照"],
        "tim_pwm": ["pwm", "舵机", "电机", "蜂鸣器", "呼吸灯", "调光"],
        "tim_basic": ["定时器", "timer", "周期", "计时"],
        "gpio_out": ["led", "继电器", "数码管", "指示灯", "开关输出"],
        "gpio_in": ["按键", "按钮", "开关输入", "限位", "光电"],
        "exti": ["外部中断", "中断触发", "边沿检测"],
        "dma": ["dma", "高速传输"],
    }
    for periph, keywords in periph_map.items():
        if any(kw in desc_lower for kw in keywords):
            detected_peripherals.append(periph)

    if peripherals:
        for p in peripherals:
            if p.lower() not in detected_peripherals:
                detected_peripherals.append(p.lower())

    # ── 任务规划 ──────────────────────────────────────────
    planned_tasks = []

    # 根据需求自动推荐任务结构
    task_templates = {
        "sensor": {
            "triggers": [
                "传感器",
                "adc",
                "采样",
                "温度",
                "湿度",
                "加速度",
                "陀螺仪",
                "光照",
                "压力",
            ],
            "name": "SensorTask",
            "purpose": "传感器数据采集与处理",
            "priority": 3,
            "stack": 256 if uses_float else 128,
            "comm_out": "Queue（发送处理后的数据）",
        },
        "display": {
            "triggers": ["显示", "oled", "lcd", "tft", "数码管", "屏幕"],
            "name": "DisplayTask",
            "purpose": "显示刷新（从队列接收数据更新显示）",
            "priority": 1,
            "stack": 256,
            "comm_in": "Queue（接收要显示的数据）",
        },
        "control": {
            "triggers": ["控制", "电机", "舵机", "pid", "调节", "反馈"],
            "name": "ControlTask",
            "purpose": "控制算法执行（PID 等）",
            "priority": 4,
            "stack": 384 if uses_float else 256,
            "comm_in": "Queue（接收传感器数据）",
            "comm_out": "Queue/直接GPIO（输出控制信号）",
        },
        "comm": {
            "triggers": ["通信", "上位机", "蓝牙", "wifi", "发送", "协议"],
            "name": "CommTask",
            "purpose": "通信处理（串口/蓝牙数据收发）",
            "priority": 2,
            "stack": 384 if uses_printf else 256,
            "comm_in": "Queue（待发送的数据）",
        },
        "led": {
            "triggers": ["led", "指示灯", "呼吸灯", "闪烁"],
            "name": "LEDTask",
            "purpose": "LED 状态指示",
            "priority": 1,
            "stack": 128,
        },
        "button": {
            "triggers": ["按键", "按钮", "输入", "用户交互"],
            "name": "ButtonTask",
            "purpose": "按键扫描与事件分发",
            "priority": 2,
            "stack": 128,
            "comm_out": "TaskNotify/EventGroup（按键事件通知其他任务）",
        },
        "alarm": {
            "triggers": ["报警", "蜂鸣器", "警报", "阈值"],
            "name": "AlarmTask",
            "purpose": "报警判断与执行",
            "priority": 3,
            "stack": 128,
            "comm_in": "TaskNotify（由传感器任务触发）",
        },
        "log": {
            "triggers": ["日志", "记录", "sd卡", "存储"],
            "name": "LogTask",
            "purpose": "数据记录与存储",
            "priority": 1,
            "stack": 384,
            "comm_in": "Queue（待记录的数据）",
        },
    }

    if task_hints:
        # 用户提供了任务提示，直接使用
        for hint in task_hints:
            if isinstance(hint, dict):
                planned_tasks.append(hint)
            elif isinstance(hint, str):
                planned_tasks.append(
                    {
                        "name": hint,
                        "purpose": f"用户指定任务: {hint}",
                        "priority": 2,
                        "stack": 256 if uses_float else 128,
                    }
                )
    else:
        # 自动推荐
        for tmpl_name, tmpl in task_templates.items():
            if any(kw in desc_lower for kw in tmpl["triggers"]):
                task = {k: v for k, v in tmpl.items() if k != "triggers"}
                planned_tasks.append(task)

    # 如果没有匹配到任何任务，给一个默认的
    if not planned_tasks:
        planned_tasks.append(
            {
                "name": "MainTask",
                "purpose": "主要业务逻辑",
                "priority": 2,
                "stack": 256 if uses_float else 128,
            }
        )

    # ── 中断规划 ──────────────────────────────────────────
    interrupt_plan = []
    if "exti" in detected_peripherals or "中断" in desc_lower:
        interrupt_plan.append(
            {
                "irq": "EXTIx_IRQHandler",
                "strategy": "vTaskNotifyGiveFromISR → 唤醒处理任务",
                "note": "ISR 中仅做通知，数据处理在任务中完成",
            }
        )
    if "uart" in detected_peripherals:
        interrupt_plan.append(
            {
                "irq": "USARTx_IRQHandler",
                "strategy": "接收中断 → xQueueSendFromISR → CommTask 处理",
                "note": "使用 HAL_UART_Receive_IT 触发中断回调",
            }
        )
    if "tim_basic" in detected_peripherals or "tim_pwm" in detected_peripherals:
        interrupt_plan.append(
            {
                "irq": "TIMx_IRQHandler",
                "strategy": "定时器中断回调 → 设置标志或通知任务",
                "note": "周期性采样可用 vTaskDelayUntil 替代定时器中断",
            }
        )

    # ── 通信拓扑推荐 ──────────────────────────────────────
    comm_topology = []
    task_names = [t["name"] for t in planned_tasks]

    if "SensorTask" in task_names and "DisplayTask" in task_names:
        comm_topology.append(
            {
                "from": "SensorTask",
                "to": "DisplayTask",
                "mechanism": "xQueueSend/xQueueReceive",
                "data": "传感器数据结构体",
            }
        )
    if "SensorTask" in task_names and "ControlTask" in task_names:
        comm_topology.append(
            {
                "from": "SensorTask",
                "to": "ControlTask",
                "mechanism": "xQueueSend/xQueueReceive",
                "data": "传感器原始值",
            }
        )
    if "SensorTask" in task_names and "AlarmTask" in task_names:
        comm_topology.append(
            {
                "from": "SensorTask",
                "to": "AlarmTask",
                "mechanism": "xTaskNotifyGive（阈值触发时通知）",
                "data": "无（AlarmTask 自行读取共享数据）",
            }
        )
    if "ButtonTask" in task_names:
        for t in task_names:
            if t != "ButtonTask" and t != "LEDTask":
                comm_topology.append(
                    {
                        "from": "ButtonTask",
                        "to": t,
                        "mechanism": "xEventGroupSetBits",
                        "data": "按键事件位",
                    }
                )
                break  # 只连接一个示例

    if "CommTask" in task_names:
        data_sources = [t for t in task_names if t in ("SensorTask", "ControlTask", "LogTask")]
        for src in data_sources[:1]:
            comm_topology.append(
                {
                    "from": src,
                    "to": "CommTask",
                    "mechanism": "xQueueSend",
                    "data": "待发送的数据包",
                }
            )

    # ── 资源估算 ──────────────────────────────────────────
    total_stack_bytes = sum(t.get("stack", 128) * 4 for t in planned_tasks)
    tcb_overhead = len(planned_tasks) * 92
    idle_timer_overhead = 92 + 256 * 4 + 92 + 512 * 4  # Idle + Timer 任务
    queue_overhead = len(comm_topology) * 120  # 每个队列约 120B
    recommended_heap = int(
        (total_stack_bytes + tcb_overhead + idle_timer_overhead + queue_overhead) * 1.3
    )
    recommended_heap = ((recommended_heap + 255) // 256) * 256

    total_ram_est = recommended_heap + 4096  # heap + static
    ram_pct = total_ram_est * 100 // (ram_k * 1024)

    resource_check = {
        "total_tasks": len(planned_tasks) + 2,  # +Idle +Timer
        "total_stack_bytes": total_stack_bytes,
        "recommended_heap": recommended_heap,
        "estimated_ram_usage": total_ram_est,
        "ram_percent": ram_pct,
        "flash_k": flash_k,
        "ram_k": ram_k,
    }

    warnings = []
    if ram_pct > 85:
        warnings.append(
            f"⚠ RAM 使用率预估 {ram_pct}%（{total_ram_est}B / {ram_k*1024}B），考虑减少任务数或栈大小"
        )
    if ram_k < 16 and len(planned_tasks) > 2:
        warnings.append(f"⚠ RAM 仅 {ram_k}KB，{len(planned_tasks)} 个任务可能不够用")
    if not has_fpu and uses_float:
        warnings.append("⚠ 当前芯片无 FPU，浮点运算将使用软件模拟（较慢）")

    # ── 构造规划文档 ──────────────────────────────────────
    plan_text = f"📋 FreeRTOS 项目规划\n"
    plan_text += f"芯片: {ci.get('define','')} ({cpu}, {flash_k}K Flash, {ram_k}K RAM, {'FPU' if has_fpu else '无FPU'})\n"
    plan_text += f"需求: {description}\n\n"

    plan_text += "━━━ 任务规划 ━━━\n"
    for i, t in enumerate(planned_tasks, 1):
        plan_text += (
            f"  {i}. {t['name']} (优先级={t.get('priority',2)}, "
            f"栈={t.get('stack',128)}words)\n"
            f"     职责: {t.get('purpose','')}\n"
        )
        if "comm_in" in t:
            plan_text += f"     输入: {t['comm_in']}\n"
        if "comm_out" in t:
            plan_text += f"     输出: {t['comm_out']}\n"
    plan_text += f"  + Idle任务 + Timer任务 (系统自动创建)\n\n"

    if comm_topology:
        plan_text += "━━━ 通信拓扑 ━━━\n"
        for c in comm_topology:
            plan_text += f"  {c['from']} → {c['to']}: {c['mechanism']} ({c.get('data','')})\n"
        plan_text += "\n"

    if interrupt_plan:
        plan_text += "━━━ 中断策略 ━━━\n"
        for ip in interrupt_plan:
            plan_text += f"  {ip['irq']}: {ip['strategy']}\n"
            if ip.get("note"):
                plan_text += f"    💡 {ip['note']}\n"
        plan_text += "\n"

    plan_text += "━━━ 资源估算 ━━━\n"
    plan_text += f"  任务总数: {resource_check['total_tasks']} (含 Idle + Timer)\n"
    plan_text += f"  栈总量: {total_stack_bytes}B\n"
    plan_text += f"  推荐堆: {recommended_heap}B\n"
    plan_text += f"  RAM 预估: {total_ram_est}B / {ram_k*1024}B ({ram_pct}%)\n"

    if detected_peripherals:
        plan_text += f"\n━━━ 使用外设 ━━━\n"
        plan_text += f"  {', '.join(detected_peripherals)}\n"

    if warnings:
        plan_text += f"\n━━━ 警告 ━━━\n"
        for w in warnings:
            plan_text += f"  {w}\n"

    return {
        "success": True,
        "plan": {
            "tasks": planned_tasks,
            "communication": comm_topology,
            "interrupts": interrupt_plan,
            "peripherals": detected_peripherals,
            "resources": resource_check,
        },
        "warnings": warnings,
        "uses_fpu": uses_float and has_fpu,
        "uses_printf": uses_printf,
        "message": plan_text,
    }


def stm32_flash(bin_path: str = None) -> dict:
    """烧录固件到 STM32（需要先 connect + compile）"""
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接，请先调用 stm32_connect"}
    path = bin_path or _last_bin_path
    if not path or not Path(path).exists():
        return {"success": False, "message": f"固件文件不存在: {path}"}
    _get_serial().clear()
    r = _get_bridge().flash(path)
    return {"success": r["ok"], "message": r["msg"]}


def stm32_read_registers(regs: list = None) -> dict:
    """读取 STM32 硬件寄存器（RCC、GPIO、TIM、UART 等）"""
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接"}
    result = _get_bridge().read_registers(regs) if regs else _get_bridge().read_all_for_debug()
    if result is not None:
        return {"success": True, "registers": result}
    return {"success": False, "message": "寄存器读取失败"}


def stm32_analyze_fault() -> dict:
    """读取并分析 HardFault 寄存器（SCB_CFSR / SCB_HFSR）"""
    if not _hw_connected:
        return {"success": False, "message": "硬件未连接"}
    regs = _get_bridge().read_registers(["SCB_CFSR", "SCB_HFSR", "SCB_BFAR", "PC"])
    if not regs:
        return {"success": False, "message": "寄存器读取失败"}
    analysis = _get_bridge().analyze_fault(regs)
    return {"success": True, "registers": regs, "analysis": analysis}


def stm32_serial_read(timeout: float = 3.0, wait_for: str = None) -> dict:
    """读取 UART 串口输出（调试日志）"""
    if not _serial_connected:
        return {"success": False, "message": "串口未连接"}
    serial = _get_serial()
    if wait_for:
        output = serial.wait_for(wait_for, timeout=timeout)
    else:
        time.sleep(min(timeout, 2.0))
        output = serial.read_and_clear()
    return {"success": True, "output": output, "has_output": bool(output.strip())}


def stm32_reset_debug_attempts() -> dict:
    """重置调试轮次计数器。开始一个全新需求时调用，确保计数从 1 开始。"""
    global _debug_attempt
    _debug_attempt = 0
    return {"success": True, "message": "计数器已重置"}


def stm32_auto_flash_cycle(code: str, request: str = "") -> dict:
    """
    完整开发闭环（自动计轮次，最多 MAX_DEBUG_ATTEMPTS 轮）：
      编译 → 烧录（若已连接硬件）→ 等待启动 → 读串口 → 读寄存器
    返回每步结果 + 当前轮次 + 是否应放弃。
    """
    global _last_code, _last_bin_path, _debug_attempt
    _debug_attempt += 1
    attempt = _debug_attempt
    steps = []

    # 超出最大轮次 → 直接告知 AI 放弃
    if attempt > MAX_DEBUG_ATTEMPTS:
        return {
            "success": False,
            "give_up": True,
            "attempt": attempt,
            "message": f"已达到最大调试轮次 ({MAX_DEBUG_ATTEMPTS})，自动调试无法解决，请检查硬件接线或手动排查",
            "steps": [],
        }

    remaining = MAX_DEBUG_ATTEMPTS - attempt
    CONSOLE.print(f"[dim]  第 {attempt}/{MAX_DEBUG_ATTEMPTS} 轮[/]")

    # 1. 编译
    comp = stm32_compile(code)
    steps.append({"step": "compile", "success": comp["success"], "msg": comp["message"][:300]})
    if not comp["success"]:
        return {
            "success": False,
            "attempt": attempt,
            "remaining": remaining,
            "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
            "steps": steps,
            "error": "编译失败，请根据错误信息修改代码",
            "compile_errors": comp["message"],
        }

    # 2. 烧录（失败时延迟重试，最多2次）
    if _hw_connected and comp.get("bin_path"):
        fr = stm32_flash(comp["bin_path"])
        if not fr["success"]:
            for retry in range(2):
                wait_sec = 1.5 * (retry + 1)
                CONSOLE.print(
                    f"[yellow]  烧录失败（{fr['message'][:60]}），{wait_sec:.0f}s 后重连重试...[/]"
                )
                time.sleep(wait_sec)
                stm32_connect(_current_chip)
                fr = stm32_flash(comp["bin_path"])
                if fr["success"]:
                    break
        steps.append({"step": "flash", "success": fr["success"], "msg": fr["message"]})
        if not fr["success"]:
            return {
                "success": False,
                "attempt": attempt,
                "remaining": remaining,
                "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
                "steps": steps,
                "error": f"烧录失败（重试2次后仍失败）: {fr['message']}",
            }
        # 3. 串口监控
        uart_out = ""
        sensor_errors = []
        if _serial_connected:
            CONSOLE.print("[dim]  等待启动...[/]")
            uart_out = _wait_serial_adaptive(
                _get_serial(),
                keyword="Gary:BOOT",
                min_wait=0.5,
                max_wait=POST_FLASH_DELAY + 4.0,
            )
            boot_ok = "Gary:BOOT" in uart_out
            # 检测传感器错误关键词
            sensor_errors = [
                line.strip()
                for line in uart_out.splitlines()
                if "ERR:" in line or "Error" in line or "not found" in line.lower()
            ]
            # 打印串口输出到终端供用户查看
            if uart_out.strip():
                CONSOLE.print(f"[dim]  串口输出:[/]\n[cyan]{uart_out.strip()[:400]}[/]")
            else:
                CONSOLE.print("[yellow]  串口无输出（程序未启动或卡死）[/]")
            steps.append(
                {
                    "step": "uart",
                    "output": uart_out[:500],
                    "boot_ok": boot_ok,
                    "sensor_errors": sensor_errors,
                }
            )
        else:
            time.sleep(POST_FLASH_DELAY + 1.0)
            boot_ok = True  # 无串口跳过验证

        # 4. 读寄存器（打印到终端）
        regs = _get_bridge().read_all_for_debug()
        has_fault = False
        if regs:
            KEY_SET = (
                "SCB_CFSR",
                "PC",
                "RCC_APB2ENR",
                "RCC_APB1ENR",
                "GPIOA_CRL",
                "GPIOA_CRH",
                "GPIOA_ODR",
                "GPIOA_IDR",
                "GPIOB_CRL",
                "GPIOB_CRH",
                "GPIOB_ODR",
                "GPIOB_IDR",
                "GPIOC_CRL",
                "GPIOC_CRH",
                "GPIOC_ODR",
                "GPIOC_IDR",
                "GPIOA_MODER",
                "GPIOB_MODER",
                "GPIOC_MODER",
                "TIM2_CR1",
                "I2C1_CR1",
                "I2C1_SR1",
                "I2C2_CR1",
                "I2C2_SR1",
            )
            key_regs = {k: v for k, v in regs.items() if k in KEY_SET}
            has_fault = regs.get("SCB_CFSR", "0x00000000") not in ("0x00000000", "0x0")
            # 打印寄存器到终端
            CONSOLE.print("[dim]  关键寄存器:[/]")
            for k, v in key_regs.items():
                color = "red" if (k == "SCB_CFSR" and has_fault) else "dim"
                CONSOLE.print(f"[{color}]    {k} = {v}[/]")
            steps.append(
                {
                    "step": "registers",
                    "key_regs": key_regs,
                    "has_hardfault": has_fault,
                    "fault_analysis": _get_bridge().analyze_fault(regs) if has_fault else "",
                }
            )
        else:
            CONSOLE.print("[yellow]  寄存器读取失败（探针连接问题）[/]")

        # ── 5. 硬件缺失检测（I2C NACK/ARLO = 外设未接）──
        hw_missing = []
        if regs:
            for i2c_name in ("I2C1", "I2C2"):
                sr1_str = regs.get(f"{i2c_name}_SR1", "0x00000000")
                try:
                    sr1 = int(sr1_str, 16)
                except ValueError:
                    continue
                if sr1 & 0x0400:  # bit10 = AF (Acknowledge Failure / NACK)
                    hw_missing.append(f"{i2c_name}: AF(NACK)——设备未应答，很可能未接或地址错误")
                if sr1 & 0x0200:  # bit9 = ARLO (Arbitration Lost)
                    hw_missing.append(f"{i2c_name}: ARLO(仲裁丢失)——总线无设备响应或接线错误")
        # 串口 ERR 也算硬件问题
        if sensor_errors:
            hw_missing.extend(sensor_errors)

        if hw_missing:
            CONSOLE.print(f"[red bold]  ⚠ 检测到硬件缺失:[/]")
            for m in hw_missing:
                CONSOLE.print(f"[red]    • {m}[/]")

        runtime_ok = boot_ok and not has_fault and not hw_missing

        if runtime_ok:
            if request:
                _stm32_save_project(code, comp, request)
            _record_success_memory(
                "runtime_success",
                code=code,
                result=comp,
                request=request,
                steps=steps,
            )
            return {
                "success": True,
                "attempt": attempt,
                "steps": steps,
                "bin_size": comp.get("bin_size", 0),
            }
        elif hw_missing:
            # 硬件缺失 → 代码没问题，不要再修代码了，直接告知用户
            return {
                "success": False,
                "attempt": attempt,
                "remaining": remaining,
                "give_up": True,
                "hw_missing": hw_missing,
                "steps": steps,
                "error": (
                    "⚠ 硬件未接或接线错误（不是代码问题，停止修改代码）：\n"
                    + "\n".join(f"  • {m}" for m in hw_missing)
                    + "\n请告知用户检查硬件连接后重试。"
                ),
            }
        else:
            err_msg = "HardFault 或程序未正常启动，请根据 steps 中的寄存器和串口信息修复"
            if not boot_ok and not uart_out.strip():
                err_msg = "串口无任何输出——程序在打印 Gary: 之前就卡死了（常见原因：I2C 等待超时/传感器未接/死循环）"
            elif not boot_ok and uart_out.strip():
                err_msg = f"程序有输出但未打印 Gary: 启动标志，串口内容: {uart_out.strip()[:200]}"
            return {
                "success": False,
                "attempt": attempt,
                "remaining": remaining,
                "give_up": attempt >= MAX_DEBUG_ATTEMPTS,
                "steps": steps,
                "error": err_msg,
            }

    # 无硬件 → 仅编译
    if request:
        _stm32_save_project(code, comp, request)
    return {
        "success": True,
        "attempt": attempt,
        "steps": steps,
        "note": "硬件未连接，已完成编译",
        "bin_size": comp.get("bin_size", 0),
        "bin_path": comp.get("bin_path"),
    }


def stm32_save_code(code: str, request: str = "untitled") -> dict:
    """保存代码到项目目录"""
    comp_result = {"bin_path": None, "bin_size": 0}
    path = _stm32_save_project(code, comp_result, request)
    return {"success": True, "path": str(path), "message": f"已保存: {path}"}


def _stm32_save_project(code: str, comp: dict, request: str) -> Path:
    """内部：保存项目文件"""
    global _last_code
    ts = time.strftime("%Y%m%d_%H%M%S")
    safe = "".join(c if c.isalnum() or c in "_- " else "" for c in request[:30]).strip()
    d = PROJECTS_DIR / f"{ts}_{safe}"
    d.mkdir(parents=True, exist_ok=True)
    (d / "main.c").write_text(code, encoding="utf-8")
    if comp.get("bin_path") and Path(comp["bin_path"]).exists():
        shutil.copy2(comp["bin_path"], d / "firmware.bin")
    (d / "config.json").write_text(
        json.dumps(
            {
                "chip": _current_chip,
                "request": request,
                "bin_size": comp.get("bin_size", 0),
                "timestamp": ts,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    _last_code = code
    # 同步更新 latest_workspace，保证 stm32_recompile 始终能找到最新代码
    try:
        latest = Path.home() / ".stm32_agent" / "workspace" / "projects" / "latest_workspace"
        latest.mkdir(parents=True, exist_ok=True)
        (latest / "main.c").write_text(code, encoding="utf-8")
    except Exception:
        pass
    CONSOLE.print(f"[dim]  已保存: {d}[/]")
    return d


def stm32_list_projects() -> dict:
    """列出最近 15 个历史项目"""
    if not PROJECTS_DIR.exists():
        return {"success": True, "projects": [], "message": "暂无项目"}
    projects = []
    for p in sorted(PROJECTS_DIR.iterdir(), reverse=True)[:15]:
        cf = p / "config.json"
        if cf.exists():
            try:
                c = json.loads(cf.read_text(encoding="utf-8"))
                projects.append(
                    {
                        "name": p.name,
                        "chip": c.get("chip", "?"),
                        "request": c.get("request", ""),
                        "timestamp": c.get("timestamp", ""),
                    }
                )
            except Exception:
                pass
    return {"success": True, "projects": projects}


def stm32_read_project(project_name: str) -> dict:
    """读取指定项目的 main.c 代码"""
    p = PROJECTS_DIR / project_name / "main.c"
    if not p.exists():
        return {"success": False, "message": f"项目不存在: {project_name}"}
    code = p.read_text(encoding="utf-8")
    return {"success": True, "code": code, "path": str(p), "lines": len(code.splitlines())}


# ─────────────────────────────────────────────────────────────
# 通用文件/命令工具（来自 claude_terminal，简化版）
# ─────────────────────────────────────────────────────────────


def read_file(file_path: str) -> dict:
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        content = p.read_text(encoding="utf-8", errors="ignore")
        lines = content.splitlines()
        numbered = "\n".join(f"{i+1:4d} | {l}" for i, l in enumerate(lines[:800]))
        return {
            "success": True,
            "numbered_view": numbered,
            "raw_content": content[:40000],
            "total_lines": len(lines),
        }
    except Exception as e:
        return {"error": str(e)}


def create_or_overwrite_file(file_path: str, content: str) -> dict:
    try:
        p = Path(file_path).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return {"success": True, "path": str(p), "lines": len(content.splitlines())}
    except Exception as e:
        return {"error": str(e)}


def str_replace_edit(file_path: str, old_str: str, new_str: str) -> dict:
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        content = p.read_text(encoding="utf-8", errors="ignore")
        count = content.count(old_str)
        if count == 0:
            return {"error": "未找到 old_str，请检查空格/换行是否完全一致"}
        if count > 1:
            return {"error": f"找到 {count} 个匹配，请增加上下文使其唯一"}
        new_content = content.replace(old_str, new_str, 1)
        p.write_text(new_content, encoding="utf-8")
        return {"success": True, "message": "替换成功", "path": str(p)}
    except Exception as e:
        return {"error": str(e)}


def list_directory(path: str = ".") -> dict:
    try:
        p = Path(path).expanduser().resolve()
        items = [{"name": x.name, "type": "dir" if x.is_dir() else "file"} for x in p.iterdir()]
        return {
            "success": True,
            "path": str(p),
            "items": sorted(items, key=lambda x: (x["type"], x["name"])),
        }
    except Exception as e:
        return {"error": str(e)}


def execute_command(command: str) -> dict:
    if any(f in command for f in ["rm -rf /", ":(){ :|:& };:"]):
        return {"error": "命令被安全策略拒绝"}
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=60)
        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:3000],
            "stderr": result.stderr[:1000],
            "returncode": result.returncode,
        }
    except subprocess.TimeoutExpired:
        return {"error": "命令超时（60s）"}
    except Exception as e:
        return {"error": str(e)}


def search_files(query: str, path: str = ".", file_type: str = None) -> dict:
    try:
        results = []
        for fp in Path(path).expanduser().resolve().rglob("*"):
            if not fp.is_file():
                continue
            if query.lower() not in fp.name.lower():
                continue
            if file_type and fp.suffix != file_type:
                continue
            results.append(str(fp))
            if len(results) >= 20:
                break
        return {"success": True, "files": results}
    except Exception as e:
        return {"error": str(e)}


def web_search(query: str) -> dict:
    try:
        import requests

        r = requests.get(
            "http://127.0.0.1:8080/search", params={"q": query, "format": "json"}, timeout=8
        )
        data = r.json()
        results = [
            {"title": x.get("title"), "url": x.get("url"), "snippet": x.get("content", "")[:200]}
            for x in data.get("results", [])[:5]
        ]
        return {"success": True, "results": results}
    except Exception as e:
        return {"error": f"搜索失败（需要本地 SearXNG）: {e}"}


# ─────────────────────────────────────────────────────────────
# 扩展工具（来自 claude_terminal / tool_schemas）
# ─────────────────────────────────────────────────────────────


def append_file_content(file_path: str, content: str) -> dict:
    """向文件末尾追加内容"""
    try:
        p = Path(file_path).expanduser().resolve()
        mode = "a" if p.exists() else "w"
        prefix = ""
        if mode == "a" and p.stat().st_size > 0:
            with open(p, "rb") as f:
                f.seek(-1, 2)
                if f.read(1) != b"\n":
                    prefix = "\n"
        with open(p, mode, encoding="utf-8") as f:
            f.write(prefix + content)
        return {"success": True, "path": str(p), "message": "内容已追加"}
    except Exception as e:
        return {"error": str(e)}


def grep_search(
    pattern: str, path: str = ".", include_extension: str = None, recursive: bool = True
) -> dict:
    """使用正则搜索文件内容（递归）"""
    try:
        search_path = Path(path).expanduser().resolve()
        results = []
        count = 0
        max_results = 20
        glob_pattern = "**/*" if recursive else "*"
        for fp in search_path.glob(glob_pattern):
            if not fp.is_file():
                continue
            if include_extension and fp.suffix != include_extension:
                continue
            if fp.stat().st_size > 1024 * 1024:
                continue
            try:
                with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                    file_content = f.read()
                matches = list(re.finditer(pattern, file_content, re.MULTILINE))
                if matches:
                    file_matches = []
                    for m in matches[:5]:
                        line_num = file_content.count("\n", 0, m.start()) + 1
                        line_start = file_content.rfind("\n", 0, m.start()) + 1
                        line_end = file_content.find("\n", m.end())
                        if line_end == -1:
                            line_end = len(file_content)
                        line_content = file_content[line_start:line_end].strip()
                        file_matches.append(f"Line {line_num}: {line_content[:100]}")
                    results.append(
                        f"File: {fp.relative_to(search_path)}\n" + "\n".join(file_matches)
                    )
                    count += 1
                    if count >= max_results:
                        break
            except Exception:
                continue
        return {
            "success": True,
            "matches_found": count,
            "results": "\n\n".join(results) if results else "No matches found",
        }
    except Exception as e:
        return {"error": str(e)}


def execute_batch_commands(commands: list, stop_on_error: bool = True) -> dict:
    """批量顺序执行多条 Shell 命令，默认遇错停止"""
    results = []
    overall_success = True
    for cmd in commands:
        res = execute_command(cmd)
        results.append({"command": cmd, "result": res})
        if not res.get("success", False):
            overall_success = False
            if stop_on_error:
                break
    return {"success": overall_success, "executed_count": len(results), "results": results}


def fetch_url(url: str) -> dict:
    """抓取 URL 页面并返回纯文本内容"""
    try:
        import requests

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return {"error": "beautifulsoup4 未安装: pip install beautifulsoup4"}
        headers = {"User-Agent": "Mozilla/5.0 (compatible; STM32Agent/1.0)"}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        return {"success": True, "url": url, "content": text[:5000], "truncated": len(text) > 5000}
    except Exception as e:
        return {"error": f"获取失败: {e}"}


def get_current_time() -> dict:
    """获取当前系统时间、星期和时区"""
    try:
        now = datetime.now()
        return {
            "success": True,
            "current_time": now.strftime("%Y-%m-%d %H:%M:%S"),
            "weekday": now.strftime("%A"),
            "timezone": str(now.astimezone().tzinfo),
        }
    except Exception as e:
        return {"error": str(e)}


def ask_human(question: str) -> dict:
    """向用户提问并等待输入"""
    try:
        CONSOLE.print(f"\n[cyan][❓ AI Question]: {question}[/]")
        answer = input(" > ")
        return {"success": True, "answer": answer}
    except Exception as e:
        return {"error": str(e)}


def git_status() -> dict:
    """执行 git status 查看修改状态"""
    return execute_command("git status")


def git_diff() -> dict:
    """执行 git diff 查看代码变更"""
    return execute_command("git diff")


def git_commit(message: str) -> dict:
    """执行 git commit -m <message>"""
    return execute_command(f"git commit -m {shlex.quote(message)}")


def edit_file_lines(
    file_path: str, operation: str, start_line: int, end_line: int = None, new_content: str = None
) -> dict:
    """基于行号编辑文件（replace/insert/delete）"""
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        with open(p, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
        total = len(lines)
        if start_line < 1 or start_line > total:
            return {"error": f"start_line {start_line} 超出范围 [1, {total}]"}
        if end_line is None:
            end_line = start_line
        if end_line < start_line or end_line > total:
            return {"error": f"end_line {end_line} 无效"}
        si, ei = start_line - 1, end_line
        if operation == "replace":
            if new_content is None:
                return {"error": "replace 需要 new_content"}
            if not new_content.endswith("\n"):
                new_content += "\n"
            new_lines = lines[:si] + new_content.splitlines(keepends=True) + lines[ei:]
        elif operation == "insert":
            if new_content is None:
                return {"error": "insert 需要 new_content"}
            if not new_content.endswith("\n"):
                new_content += "\n"
            new_lines = lines[:si] + new_content.splitlines(keepends=True) + lines[si:]
        elif operation == "delete":
            new_lines = lines[:si] + lines[ei:]
        else:
            return {"error": f"未知操作: {operation}"}
        with open(p, "w", encoding="utf-8") as f:
            f.writelines(new_lines)
        return {
            "success": True,
            "path": str(p),
            "operation": operation,
            "new_total_lines": len(new_lines),
        }
    except Exception as e:
        return {"error": str(e)}


def insert_content_by_regex(file_path: str, regex_pattern: str, content: str) -> dict:
    """在文件第一个正则匹配位置之后插入内容"""
    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        with open(p, "r", encoding="utf-8", errors="ignore") as f:
            file_content = f.read()
        m = re.search(regex_pattern, file_content, re.MULTILINE)
        if not m:
            return {"error": f"正则 '{regex_pattern}' 未匹配到内容"}
        new_content = file_content[: m.end()] + content + file_content[m.end() :]
        with open(p, "w", encoding="utf-8") as f:
            f.write(new_content)
        return {
            "success": True,
            "path": str(p),
            "match_found": m.group(0)[:50],
            "message": "内容已插入",
        }
    except Exception as e:
        return {"error": str(e)}


def check_python_code(file_path: str) -> dict:
    """检查 Python 文件语法和风格（flake8 / ast）"""
    import ast

    try:
        p = Path(file_path).expanduser().resolve()
        if not p.exists():
            return {"error": f"文件不存在: {file_path}"}
        try:
            with open(p, "r", encoding="utf-8") as f:
                ast.parse(f.read())
        except SyntaxError as e:
            return {
                "success": False,
                "error_type": "SyntaxError",
                "line": e.lineno,
                "message": str(e),
            }
        lint_result = ""
        try:
            result = subprocess.run(
                f"flake8 {shlex.quote(str(p))}",
                shell=True,
                capture_output=True,
                text=True,
                timeout=10,
            )
            if result.returncode != 0 and result.stdout:
                lint_result = f"Flake8:\n{result.stdout}"
        except Exception:
            pass
        return {
            "success": True,
            "message": "语法检查通过",
            "linter_output": lint_result or "无问题",
        }
    except Exception as e:
        return {"error": str(e)}


def run_python_code(code: str) -> dict:
    """执行 Python 代码片段（临时文件沙箱）"""
    import tempfile

    try:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".py", delete=False, encoding="utf-8"
        ) as tmp:
            tmp.write(code)
            tmp_path = tmp.name
        result = subprocess.run(
            [sys.executable, tmp_path], capture_output=True, text=True, timeout=30
        )
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:3000],
            "stderr": result.stderr[:1000],
            "returncode": result.returncode,
        }
    except Exception as e:
        return {"error": str(e)}


# ── Word 文档工具 ──────────────────────────────────────────────


def _get_docx_module():
    """懒加载 python-docx，未安装时返回 None"""
    try:
        import docx

        return docx
    except ImportError:
        return None


def read_docx(file_path: str) -> dict:
    """读取 Word 文档(.docx)的文本内容"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return {"success": True, "content": text, "total_paragraphs": len(doc.paragraphs)}
    except Exception as e:
        return {"error": str(e)}


def replace_docx_text(
    file_path: str, old_text: str, new_text: str, use_regex: bool = False
) -> dict:
    """替换 Word 文档中的文本（支持正则）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        count = 0
        for para in doc.paragraphs:
            if use_regex:
                if re.search(old_text, para.text):
                    replaced = re.sub(old_text, new_text, para.text)
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = replaced
                    else:
                        para.add_run(replaced)
                    count += 1
            else:
                if old_text in para.text:
                    replaced_in_run = False
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
                            replaced_in_run = True
                    if not replaced_in_run:
                        para.text = para.text.replace(old_text, new_text)
                        count += 1
        doc.save(file_path)
        return {"success": True, "replaced_count": count, "message": f"已替换 {count} 处"}
    except Exception as e:
        return {"error": str(e)}


def append_docx_content(
    file_path: str, content: str, after_paragraph_index: int = None, style: str = None
) -> dict:
    """向 Word 文档追加内容（支持指定位置插入）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        paragraphs_text = [t for t in content.split("\n") if t.strip()]
        if after_paragraph_index is None:
            for p_text in paragraphs_text:
                p = doc.add_paragraph(p_text)
                if style:
                    try:
                        p.style = style
                    except Exception:
                        pass
        else:
            n = len(doc.paragraphs)
            if after_paragraph_index < 0 or after_paragraph_index >= n:
                return {"error": f"索引 {after_paragraph_index} 超出范围 (0-{n-1})"}
            if after_paragraph_index == n - 1:
                for p_text in paragraphs_text:
                    p = doc.add_paragraph(p_text)
                    if style:
                        try:
                            p.style = style
                        except Exception:
                            pass
            else:
                next_para = doc.paragraphs[after_paragraph_index + 1]
                base_style = doc.paragraphs[after_paragraph_index].style
                for p_text in paragraphs_text:
                    new_p = next_para.insert_paragraph_before(p_text)
                    if style:
                        try:
                            new_p.style = style
                        except Exception:
                            pass
                    else:
                        new_p.style = base_style
        doc.save(file_path)
        return {"success": True, "message": "内容已追加"}
    except Exception as e:
        return {"error": str(e)}


def inspect_docx_structure(file_path: str, max_paragraphs: int = 50) -> dict:
    """查看 Word 文档段落结构（用于定位插入点）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        structure = []
        for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            if not preview.strip():
                preview = "[空段落]"
            structure.append(f"[{i}] {preview}")
        return {
            "success": True,
            "total_paragraphs": len(doc.paragraphs),
            "structure": "\n".join(structure),
        }
    except Exception as e:
        return {"error": str(e)}


def insert_docx_content_after_heading(
    file_path: str, heading_text: str, content: str, style: str = None
) -> dict:
    """在 Word 文档指定标题后插入内容（大小写不敏感）"""
    docx_mod = _get_docx_module()
    if docx_mod is None:
        return {"error": "python-docx 未安装: pip install python-docx"}
    try:
        doc = docx_mod.Document(file_path)
        target_para = None
        for para in doc.paragraphs:
            if heading_text.lower() in para.text.lower():
                target_para = para
                break
        if not target_para:
            return {"error": f"未找到标题: {heading_text}"}
        index = doc.paragraphs.index(target_para)
        return append_docx_content(file_path, content, index, style)
    except Exception as e:
        return {"error": str(e)}


# ── 电脑控制工具 ──────────────────────────────────────────────


def computer_screenshot() -> dict:
    """截取当前桌面截图并保存为 PNG"""
    try:
        import pyautogui

        ts = int(time.time())
        path = os.path.abspath(f"screenshot_{ts}.png")
        pyautogui.screenshot(path)
        return {"success": True, "image_path": path, "message": f"截图已保存: {path}"}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_mouse_move(x: int, y: int) -> dict:
    """移动鼠标到指定坐标"""
    try:
        import pyautogui

        pyautogui.moveTo(x, y)
        return {"success": True, "action": "move", "x": x, "y": y}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_mouse_click(button: str = "left") -> dict:
    """鼠标点击（left/right/double）"""
    try:
        import pyautogui

        if button == "double":
            pyautogui.doubleClick()
        else:
            pyautogui.click(button=button)
        return {"success": True, "button": button}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


def computer_keyboard_type(text: str) -> dict:
    """向焦点窗口输入文本"""
    try:
        import pyautogui

        pyautogui.write(text)
        return {"success": True, "typed": text}
    except ImportError:
        return {"error": "pyautogui 未安装: pip install pyautogui"}
    except Exception as e:
        return {"error": str(e)}


# ─────────────────────────────────────────────────────────────
# 工具注册表
# ─────────────────────────────────────────────────────────────
TOOLS_MAP: Dict[str, Any] = {
    # STM32 专属
    "stm32_list_probes": stm32_list_probes,
    "stm32_connect": stm32_connect,
    "stm32_disconnect": stm32_disconnect,
    "stm32_serial_connect": stm32_serial_connect,
    "stm32_serial_disconnect": stm32_serial_disconnect,
    "stm32_set_chip": stm32_set_chip,
    "stm32_hardware_status": stm32_hardware_status,
    "stm32_compile": stm32_compile,
    "stm32_compile_rtos": stm32_compile_rtos,
    "stm32_recompile": stm32_recompile,
    "stm32_regen_bsp": stm32_regen_bsp,
    "stm32_analyze_fault_rtos": stm32_analyze_fault_rtos,
    "stm32_rtos_check_code": stm32_rtos_check_code,
    "stm32_rtos_task_stats": stm32_rtos_task_stats,
    "stm32_rtos_suggest_config": stm32_rtos_suggest_config,
    "stm32_rtos_plan_project": stm32_rtos_plan_project,
    "stm32_flash": stm32_flash,
    "stm32_read_registers": stm32_read_registers,
    "stm32_analyze_fault": stm32_analyze_fault,
    "stm32_serial_read": stm32_serial_read,
    "stm32_auto_flash_cycle": stm32_auto_flash_cycle,
    "stm32_reset_debug_attempts": stm32_reset_debug_attempts,
    "stm32_generate_font": stm32_generate_font,
    "stm32_save_code": stm32_save_code,
    "stm32_list_projects": stm32_list_projects,
    "stm32_read_project": stm32_read_project,
    "gary_save_member_memory": gary_save_member_memory,
    # 通用
    "read_file": read_file,
    "create_or_overwrite_file": create_or_overwrite_file,
    "str_replace_edit": str_replace_edit,
    "list_directory": list_directory,
    "execute_command": execute_command,
    "search_files": search_files,
    "web_search": web_search,
    # 扩展通用工具
    "append_file_content": append_file_content,
    "grep_search": grep_search,
    "execute_batch_commands": execute_batch_commands,
    "fetch_url": fetch_url,
    "get_current_time": get_current_time,
    "ask_human": ask_human,
    "git_status": git_status,
    "git_diff": git_diff,
    "git_commit": git_commit,
    "edit_file_lines": edit_file_lines,
    "insert_content_by_regex": insert_content_by_regex,
    "check_python_code": check_python_code,
    "run_python_code": run_python_code,
    # Word 文档工具
    "read_docx": read_docx,
    "replace_docx_text": replace_docx_text,
    "append_docx_content": append_docx_content,
    "inspect_docx_structure": inspect_docx_structure,
    "insert_docx_content_after_heading": insert_docx_content_after_heading,
    # 电脑控制工具
    "computer_screenshot": computer_screenshot,
    "computer_mouse_move": computer_mouse_move,
    "computer_mouse_click": computer_mouse_click,
    "computer_keyboard_type": computer_keyboard_type,
}
TOOLS_MAP.update(EXTRA_TOOLS_MAP)
TOOLS_MAP.update(SKILL_TOOLS_MAP)

# ─────────────────────────────────────────────────────────────
# Tool Schemas（供 AI 调用）
# ─────────────────────────────────────────────────────────────
TOOL_SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "stm32_list_probes",
            "description": "列出所有已连接的调试探针（ST-Link、CMSIS-DAP、J-Link）。连接前先调用此函数确认探针存在。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_connect",
            "description": "连接 STM32 硬件（pyocd 探针 + UART 串口）。烧录前必须先调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chip": {
                        "type": "string",
                        "description": "芯片型号，如 STM32F103C8T6（可选，不填用当前设置）",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_disconnect",
            "description": "断开探针和串口连接。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_serial_connect",
            "description": (
                "单独连接/重连 UART 串口监控（不影响 pyocd 探针）。"
                "串口用于接收 Debug_Print 日志和 Gary:BOOT 启动标记，是 AI 判断程序运行状态的关键。"
                "stm32_connect 会自动尝试用默认端口连接，若失败或需要更换端口时调用此函数。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "port": {
                        "type": "string",
                        "description": "串口设备路径，如 /dev/ttyUSB0、/dev/ttyAMA0（不填用 config.py 默认值）",
                    },
                    "baud": {"type": "integer", "description": "波特率，默认 115200"},
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_serial_disconnect",
            "description": "断开串口（保留探针连接）。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_set_chip",
            "description": "切换目标芯片型号，同时更新寄存器地址表。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chip": {
                        "type": "string",
                        "description": "芯片完整型号，如 STM32F103C8T6 / STM32F407VET6",
                    },
                },
                "required": ["chip"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_hardware_status",
            "description": "查询当前硬件状态：芯片型号、探针/串口连接状态、GCC 版本、HAL 库是否就绪。开始工作前建议先调用。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_compile",
            "description": "使用 arm-none-eabi-gcc + HAL 库编译完整的 main.c 代码，返回编译结果和 bin 路径。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {
                        "type": "string",
                        "description": "完整的 main.c 代码（含所有 #include 和函数定义）",
                    },
                    "chip": {"type": "string", "description": "可选：临时指定芯片型号"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_compile_rtos",
            "description": "使用 arm-none-eabi-gcc + HAL + FreeRTOS Kernel 编译带 RTOS 的完整 main.c 代码。"
            "仅在用户需要 FreeRTOS 多任务时使用，裸机项目使用 stm32_compile。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {
                        "type": "string",
                        "description": "完整的 main.c 代码（含 FreeRTOS 头文件和任务定义）",
                    },
                    "chip": {"type": "string", "description": "可选：临时指定芯片型号"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_recompile",
            "description": (
                "直接从 latest_workspace/main.c 重新编译，无需传递代码字符串。"
                "str_replace_edit 修改文件后使用此工具代替 read_file + stm32_compile，"
                "节省 token，避免代码传输中的幻觉。"
                "mode='auto' 自动检测裸机或 RTOS；mode='bare' 强制裸机；mode='rtos' 强制 FreeRTOS。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "mode": {
                        "type": "string",
                        "enum": ["auto", "bare", "rtos"],
                        "description": "编译模式：auto(自动检测)、bare(裸机)、rtos(FreeRTOS)。默认 auto。",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_regen_bsp",
            "description": (
                "强制重新生成 BSP 文件（startup.s、link.ld、FreeRTOSConfig.h）并验证 FPU 使能代码。"
                "在切换芯片型号后、FreeRTOS 编译前、或怀疑 startup.s 不含 CPACR 初始化时调用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "chip": {
                        "type": "string",
                        "description": "可选：指定芯片型号（如 STM32F411CE）",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_analyze_fault_rtos",
            "description": (
                "FreeRTOS 专用 HardFault 分析。读取 SCB_CFSR/HFSR/BFAR/PC，"
                "检查 startup.s FPU 使能状态，给出 FPU 未使能、栈溢出、非法地址等具体诊断。"
                "FreeRTOS 程序 HardFault 时优先调用此工具而非 stm32_analyze_fault。"
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_rtos_check_code",
            "description": (
                "FreeRTOS 代码静态检查 —— 编译前自动检测常见 RTOS 编程错误。"
                "检查项：SysTick_Handler 冲突、HAL_Delay 陷阱、缺少 hook 函数、"
                "任务栈大小不足、ISR 中使用非 FromISR API、缺少头文件等。"
                "编写 RTOS 代码后、调用 stm32_compile_rtos 前使用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "完整的 main.c 代码"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_rtos_task_stats",
            "description": (
                "读取 FreeRTOS 运行时任务统计：任务数、堆剩余/历史最低、当前任务名。"
                "通过 ELF 符号表定位内存地址，用 pyocd 读取。"
                "需要已编译（有 ELF）且硬件已连接。用于性能分析和堆使用诊断。"
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_rtos_suggest_config",
            "description": (
                "根据任务需求计算推荐的 FreeRTOS 配置：栈大小、堆大小、优先级数。"
                "评估 RAM 使用率并给出警告。在规划 RTOS 程序前调用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task_count": {
                        "type": "integer",
                        "description": "计划创建的任务数量（不含 Idle 和 Timer）",
                    },
                    "use_fpu": {
                        "type": "boolean",
                        "description": "任务是否使用浮点运算（默认 false）",
                    },
                    "use_printf": {
                        "type": "boolean",
                        "description": "任务是否使用 snprintf（默认 false）",
                    },
                    "ram_k": {
                        "type": "integer",
                        "description": "可选：指定 RAM 大小(KB)，不填用当前芯片参数",
                    },
                },
                "required": ["task_count"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_rtos_plan_project",
            "description": (
                "FreeRTOS 项目架构规划 —— 复杂 RTOS 项目的必备第一步。"
                "根据需求描述自动生成：任务分解、通信拓扑、中断策略、外设分配、资源估算。"
                "满足以下任一条件时必须调用：①任务数≥3 ②涉及中断+任务通信 ③涉及多个外设协同 ④涉及控制算法。"
                "规划结果需向用户展示并确认后再写代码。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "description": {
                        "type": "string",
                        "description": "用户需求的完整描述（中文），包含功能要求、外设需求、性能要求等",
                    },
                    "peripherals": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": '可选：明确指定的外设列表，如 ["uart", "i2c", "adc", "pwm"]',
                    },
                    "task_hints": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": '可选：用户指定的任务名称提示，如 ["SensorTask", "MotorTask"]',
                    },
                },
                "required": ["description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_flash",
            "description": "通过 pyocd 将已编译的固件烧录到 STM32。需要先 connect 和 compile。",
            "parameters": {
                "type": "object",
                "properties": {
                    "bin_path": {
                        "type": "string",
                        "description": "可选：bin 文件路径，不填则用上次编译结果",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_read_registers",
            "description": "读取 STM32 关键寄存器（RCC、GPIO ODR/IDR、TIM、UART、I2C 等）。只在 stm32_auto_flash_cycle 未返回寄存器数据时补充调用一次，禁止循环重复调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "regs": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": '可选：指定寄存器名称列表，如 ["RCC_APB2ENR", "GPIOA_CRL"]。不填读取所有调试寄存器。',
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_analyze_fault",
            "description": "读取并分析 HardFault 状态寄存器（SCB_CFSR/HFSR/BFAR + PC），定位故障原因。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_serial_read",
            "description": "读取 UART 串口输出（调试日志、错误信息）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "timeout": {"type": "number", "description": "读取超时（秒），默认 3.0"},
                    "wait_for": {
                        "type": "string",
                        "description": "可选：等待直到出现此字符串（如 Gary:）",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_auto_flash_cycle",
            "description": (
                "完整开发闭环（推荐）：编译 → 烧录 → 读串口 → 读寄存器，一步到位。"
                "代码生成后直接调用此函数，获取完整的验证结果。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "完整的 main.c 代码"},
                    "request": {"type": "string", "description": "需求描述（用于保存项目，可选）"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_reset_debug_attempts",
            "description": (
                "重置调试轮次计数器。必须调用的场景：(1)全新需求 (2)用户要求修改功能/显示内容/引脚/逻辑。"
                "不调用的场景：仅在修复上一轮的编译错误/烧录失败/运行异常时继续重试。"
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_generate_font",
            "description": (
                "将任意文字（含中文）渲染为真实字模 C 数组，使用系统字体精确生成，"
                "固定格式：横向取模·高位在前（row-major MSB=left），并附带配套的 OLED_ShowFontN() 显示函数。"
                "返回的 c_code 包含字模数组 + 显示函数，直接粘贴进 main.c 使用，无需修改。"
                "显示中文/特殊字符时必须先调用此工具，禁止手写字模数据。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "要生成字模的文字，如 '你好世界'"},
                    "size": {
                        "type": "integer",
                        "description": "字体大小（像素），默认 16，常用 8/12/16/24/32",
                        "default": 16,
                    },
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_save_code",
            "description": "将代码保存到项目目录（不编译，仅保存源码）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "main.c 代码内容"},
                    "request": {"type": "string", "description": "项目描述（作为目录名）"},
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_list_projects",
            "description": "列出最近 15 个历史项目（名称、芯片、需求描述、时间）。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stm32_read_project",
            "description": "读取指定历史项目的 main.c 源码，用于查看或修改。",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_name": {
                        "type": "string",
                        "description": "项目目录名（从 stm32_list_projects 获取）",
                    },
                },
                "required": ["project_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "gary_save_member_memory",
            "description": (
                "把高价值、可复用的经验写入 member.md。"
                "适用于：成功模板、关键初始化顺序、硬件易错点、寄存器判定经验、RTOS/裸机专项坑。"
                "禁止写冗长日志，必须提炼成短而可执行的结论。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "经验标题，简短具体，如“F103 裸机 UART 先打 Gary:BOOT 再启 I2C”",
                    },
                    "experience": {
                        "type": "string",
                        "description": "经验正文，2-6 行为宜，写成可复用的做法/结论",
                    },
                    "tags": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": '可选标签，如 ["baremetal", "uart", "i2c", "boot_marker"]',
                    },
                    "importance": {
                        "type": "string",
                        "enum": ["medium", "high", "critical"],
                        "description": "经验重要度，默认 high。",
                    },
                },
                "required": ["title", "experience"],
            },
        },
    },
    # 通用工具
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "读取文件内容（带行号）。",
            "parameters": {
                "type": "object",
                "properties": {"file_path": {"type": "string"}},
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_or_overwrite_file",
            "description": "创建或完全覆盖一个文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string"},
                    "content": {"type": "string"},
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "str_replace_edit",
            "description": "精确替换文件中的字符串（old_str 必须在文件中唯一，包含 3-5 行上下文）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string"},
                    "old_str": {
                        "type": "string",
                        "description": "要替换的原文（必须唯一，包含足够上下文）",
                    },
                    "new_str": {"type": "string", "description": "替换后的新文"},
                },
                "required": ["file_path", "old_str", "new_str"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_directory",
            "description": "列出目录内容。",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string"}},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_command",
            "description": "执行 Shell 命令（如查看日志、安装包、运行脚本等）。",
            "parameters": {
                "type": "object",
                "properties": {"command": {"type": "string"}},
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_files",
            "description": "按文件名关键字搜索文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string"},
                    "path": {"type": "string", "description": "搜索目录（默认当前）"},
                    "file_type": {"type": "string", "description": "扩展名过滤，如 .c / .h"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "网络搜索（需要本地 SearXNG 实例）。",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string"}},
                "required": ["query"],
            },
        },
    },
    # ── 扩展通用工具 ──────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "append_file_content",
            "description": "向文件末尾追加内容（代码/文本文件）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "目标文件路径"},
                    "content": {"type": "string", "description": "要追加的内容"},
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "grep_search",
            "description": "在文件中递归搜索正则表达式模式，返回匹配位置和内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "正则表达式模式"},
                    "path": {"type": "string", "description": "搜索目录（默认 .）"},
                    "include_extension": {"type": "string", "description": "按扩展名过滤，如 .py"},
                    "recursive": {"type": "boolean", "description": "是否递归（默认 True）"},
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_batch_commands",
            "description": "批量顺序执行多条 Shell 命令，默认遇错停止。",
            "parameters": {
                "type": "object",
                "properties": {
                    "commands": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "命令列表",
                    },
                    "stop_on_error": {
                        "type": "boolean",
                        "description": "遇错是否停止（默认 True）",
                    },
                },
                "required": ["commands"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "抓取 URL 页面并返回纯文本内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "要获取的 URL"},
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "获取当前系统时间、星期和时区。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_human",
            "description": "向用户提问，等待用户在终端输入回答。",
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string", "description": "要问用户的问题"},
                },
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_status",
            "description": "执行 git status，查看当前仓库的文件修改状态。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_diff",
            "description": "执行 git diff，查看实际代码变更内容。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_commit",
            "description": "执行 git commit -m <message> 提交变更。",
            "parameters": {
                "type": "object",
                "properties": {
                    "message": {"type": "string", "description": "提交信息"},
                },
                "required": ["message"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_file_lines",
            "description": (
                "基于行号编辑文件（优先使用 str_replace_edit）。\n"
                "操作类型：replace（替换行范围）、insert（在行前插入）、delete（删除行范围）。\n"
                "仅在无法用 str_replace_edit 时使用（如在空白处插入新代码）。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "文件路径"},
                    "operation": {
                        "type": "string",
                        "enum": ["replace", "insert", "delete"],
                        "description": "操作类型",
                    },
                    "start_line": {"type": "integer", "description": "起始行（1-indexed）"},
                    "end_line": {
                        "type": "integer",
                        "description": "结束行（可选，默认等于 start_line）",
                    },
                    "new_content": {
                        "type": "string",
                        "description": "新内容（replace/insert 时必填）",
                    },
                },
                "required": ["file_path", "operation", "start_line"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "insert_content_by_regex",
            "description": "在文件中第一个正则匹配位置之后插入内容，适合向类/函数后添加新方法。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "文件路径"},
                    "regex_pattern": {
                        "type": "string",
                        "description": "用于定位插入点的正则表达式",
                    },
                    "content": {"type": "string", "description": "要插入的内容"},
                },
                "required": ["file_path", "regex_pattern", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "check_python_code",
            "description": "检查 Python 文件的语法错误和代码风格（flake8/ast）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Python 文件路径（.py）"},
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_python_code",
            "description": "执行 Python 代码片段（临时文件沙箱），用于验证逻辑或测试库。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "要执行的 Python 代码"},
                },
                "required": ["code"],
            },
        },
    },
    # ── Word 文档工具 ──────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "读取 Word 文档(.docx)的文本内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "replace_docx_text",
            "description": "替换 Word 文档中的文本（尽量保留格式，支持正则）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "old_text": {"type": "string", "description": "要查找的文本或正则模式"},
                    "new_text": {"type": "string", "description": "替换为的新文本"},
                    "use_regex": {
                        "type": "boolean",
                        "description": "是否启用正则匹配（默认 False）",
                    },
                },
                "required": ["file_path", "old_text", "new_text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "append_docx_content",
            "description": "向 Word 文档追加内容，可追加到末尾或指定段落索引之后。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "content": {"type": "string", "description": "要追加的内容（\\n 分隔多段落）"},
                    "after_paragraph_index": {
                        "type": "integer",
                        "description": "插入位置段落索引（不填则追加到末尾）",
                    },
                    "style": {
                        "type": "string",
                        "description": "Word 样式名，如 'Heading 1'、'Normal'（可选）",
                    },
                },
                "required": ["file_path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "inspect_docx_structure",
            "description": "查看 Word 文档段落结构（索引和内容预览），用于确定插入位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "max_paragraphs": {
                        "type": "integer",
                        "description": "最多显示段落数（默认 50）",
                    },
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "insert_docx_content_after_heading",
            "description": "在 Word 文档指定标题段落之后插入内容（大小写不敏感匹配）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "heading_text": {
                        "type": "string",
                        "description": "目标标题文本（大小写不敏感）",
                    },
                    "content": {"type": "string", "description": "要插入的内容"},
                    "style": {"type": "string", "description": "Word 样式（可选）"},
                },
                "required": ["file_path", "heading_text", "content"],
            },
        },
    },
    # ── 电脑控制工具 ──────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "computer_screenshot",
            "description": "截取当前桌面截图，返回保存路径。",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_mouse_move",
            "description": "移动鼠标到指定坐标（x, y）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "x": {"type": "integer", "description": "X 坐标"},
                    "y": {"type": "integer", "description": "Y 坐标"},
                },
                "required": ["x", "y"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_mouse_click",
            "description": "在当前鼠标位置点击（left/right/double）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "button": {
                        "type": "string",
                        "enum": ["left", "right", "double"],
                        "description": "点击类型（默认 left）",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_keyboard_type",
            "description": "向当前焦点窗口输入文本。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "要输入的文本"},
                },
                "required": ["text"],
            },
        },
    },
]
TOOL_SCHEMAS.extend(EXTRA_TOOL_SCHEMAS)
TOOL_SCHEMAS.extend(SKILL_TOOL_SCHEMAS)
_skill_mgr = init_skills(TOOLS_MAP, TOOL_SCHEMAS)

# ─────────────────────────────────────────────────────────────
# STM32 系统提示词
STM32_BASE_SYSTEM_PROMPT = """你是 Gary Dev Agent，专为 STM32 嵌入式开发设计的 AI 助手，深度集成了编译、烧录、调试工具链。

## 核心能力
1. **代码生成**：根据自然语言需求生成完整可编译的 STM32 HAL C 代码
2. **编译验证**：调用 arm-none-eabi-gcc 编译，立即发现并修复错误
3. **固件烧录**：通过 pyocd 将固件烧录到 STM32（支持 ST-Link / CMSIS-DAP / J-Link）
4. **硬件调试**：读取外设寄存器，分析 HardFault，监控 UART 日志
5. **代码修改**：对话式增量修改，保留已有逻辑

## 标准工作流

### 串口监控（AI 判断程序运行状态的唯一来源）
- 串口 = STM32 UART TX → USB-TTL 适配器 → 主机 `/dev/ttyUSBx` 或 `/dev/ttyAMAx`
- `stm32_hardware_status` 返回 `serial_connected: false` 时，**必须提醒用户连接串口**
- 用户可用 `/serial /dev/ttyUSB0` 连接，或告诉 AI 调用 `stm32_serial_connect(port=...)`
- 无串口时 AI 无法看到 `Gary:BOOT`、`Debug_Print` 输出和运行时错误，调试能力严重受限
- 烧录成功但无串口时，在回复末尾加一句：`⚠️ 串口未连接，无法监控运行状态`

### 全新代码生成 / 功能修改
1. 调用 `stm32_reset_debug_attempts` — **以下情况必须调用**：全新需求、修改功能/引脚/内容/逻辑；仅在修复上轮编译/烧录/运行错误时跳过
2. 调用 `stm32_hardware_status` — 了解当前芯片和工具链状态，**检查 serial_connected**
3. 生成完整 main.c（见代码规范）
4. - 代码直接作为参数传入，不需要在对话里额外展示
   - **禁止**只把代码输出在文本里而不调用此工具
5. 读取工具返回值中的关键字段：
   - `success: true` → 读 `steps` 中 `step=registers` 的 `key_regs`，通过寄存器值向用户说明验证结果
   - `give_up: true` → **立即停止**，告知用户已达上限，建议手动排查硬件
   - `hw_missing` 字段存在 → **这是硬件未接/接线错误，不是代码 bug！立即停止修改代码**，将 hw_missing 列表完整告知用户，说明哪个总线/设备有问题，让用户检查接线后重试
   - `success: false, give_up: false` → 根据 `steps` 中的错误修复代码，再次调用（**不要重置计数器**）
   - 若 key_regs 为空，最多补充调用**一次** `stm32_read_registers`，之后无论结果如何直接向用户汇报
6. 寄存器解读规则（**必须向用户说明验证结果**）：
   - GPIO 输出验证：`GPIOA_ODR` 的 bit N 为 1 → PA[N] 已拉高；bit N 为 0 → 已拉低
     例：`GPIOA_ODR=0x00000001` → PA0=HIGH ✓；`GPIOA_ODR=0x00000000` → PA0=LOW ✗
   - GPIO 模式验证（F1）：`GPIOA_CRL` 每 4 bit 控制一个引脚，bit[1:0]=11→输出，bit[3:2]=00→推挽
   - GPIO 模式验证（F4/F7/H7）：`GPIOA_MODER` 每 2 bit 控制一个引脚，01→输出，00→输入
   - RCC 时钟验证：`RCC_APB2ENR` bit2=1→GPIOA 时钟已开；bit3=1→GPIOB 时钟已开
   - 有 HardFault：`SCB_CFSR != 0` → 调用 `stm32_analyze_fault` 分析
7. 修复方向：
   - `compile_errors` 非空 → 修复编译错误
   - `has_hardfault: true` → 调用 `stm32_analyze_fault`，根据 CFSR 修复
   - `boot_ok: false` → 程序未启动，检查 SysTick_Handler 和 UART 初始化
   - 寄存器值不符预期（如 ODR bit 未置位）→ 检查 RCC 时钟是否开启、GPIO 模式配置是否正确

### 增量修改（最重要！）
用户对上一次代码提出修改要求时（如"改成共阳"、"加一个按键"）：
1. 先通过对话上下文或 `stm32_read_project` 获取**上一次完整代码**
2. **只修改用户要求的部分**，其余逻辑原封不动
3. 例：上次是跑马灯共阴 → 用户说"改共阳" → 只改电平逻辑，不重写整个程序
4. 若用户需求与上次完全无关，才从头生成

### 修改历史项目
1. `stm32_list_projects` → `stm32_read_project(name)` 读取源码
2. `str_replace_edit` 精确替换（old_str 必须在文件中唯一，含3-5行上下文）

## STM32 代码规范（严格遵守）

### 必须包含
- 完整 `#include`（stm32xxx_hal.h 及各外设头文件）
- `SystemClock_Config()` — **只用 HSI 内部时钟，禁止 HSE**；根据 chip 型号正确配置 PLL 倍频/分频/Flash 等待周期/APB 分频
- `SysTick_Handler` — **必须定义，否则 HAL_Delay 永久阻塞：**
  ```c
  void SysTick_Handler(void) { HAL_IncTick(); }
  ```

### main() 函数结构（**严格按此顺序，不可调换**）
```c
int main(void) {
    HAL_Init();
    SystemClock_Config();
    // 1. 最先初始化 UART（仅配置 GPIO 和 USART，不涉及外部设备）
    MX_USART1_UART_Init();
    // 2. 紧接着打印启动标记——此时其他外设都还没初始化
    Debug_Print("Gary:BOOT\\r\\n");
    // 3. 然后初始化其他外设（I2C、SPI、TIM、OLED 等）
    MX_I2C1_Init();  // OLED
    MX_I2C2_Init();  // 传感器
    // 4. 检测外部传感器是否在线（必须有超时，不可阻塞）
    if (HAL_I2C_IsDeviceReady(&hi2c2, SENSOR_ADDR<<1, 3, 200) != HAL_OK) {
        Debug_Print("ERR: Sensor not found\\r\\n");
        // 有 OLED 时在屏幕显示错误
    }
    // 5. 主循环
    while (1) { ... }
}
```
**关键**：`Debug_Print("Gary:BOOT")` 必须紧跟 UART 初始化，在 I2C/SPI/TIM 等一切初始化**之前**。
若 I2C 初始化卡死（传感器未接导致总线锁死），至少串口已经打印了启动标志，AI 能正确判断"程序已启动但外设有问题"。
- 轻量调试函数（**不得用 sprintf**，手写整数转字符串）：
  ```c
  void Debug_Print(const char* s) {
      HAL_UART_Transmit(&huartX, (uint8_t*)s, strlen(s), 100);
  }
  void Debug_PrintInt(const char* prefix, int val) {
      // 手写：除法取位 + '0' 偏移，或查表
      char buf[16]; int i = 0, neg = 0;
      if (val < 0) { neg = 1; val = -val; }
      if (val == 0) { buf[i++] = '0'; }
      else { while (val) { buf[i++] = '0' + val % 10; val /= 10; } }
      if (neg) buf[i++] = '-';
      // 反转后发送
      HAL_UART_Transmit(&huartX, (uint8_t*)prefix, strlen(prefix), 100);
      for (int j = i-1; j >= 0; j--) HAL_UART_Transmit(&huartX, (uint8_t*)&buf[j], 1, 100);
      HAL_UART_Transmit(&huartX, (uint8_t*)"\\r\\n", 2, 100);
  }
  ```
- 每个关键外设（I2C、SPI、ADC 等）初始化后检查返回值：
  ```c
  if (HAL_I2C_Init(&hi2c1) != HAL_OK) { Debug_Print("ERR: I2C Init Fail\\r\\n"); }
  ```
- **I2C 传感器必须检测设备是否在线**（不能假设已连接）：
  ```c
  if (HAL_I2C_IsDeviceReady(&hi2c2, SENSOR_ADDR << 1, 3, 100) != HAL_OK) {
      Debug_Print("ERR: Sensor not found\\r\\n");
      // 若有 OLED，显示错误信息并停留：
      OLED_ShowString(0, 0, "Sensor Error");
      while (1) { HAL_Delay(500); }
  }
  ```
  传感器地址用 7-bit 值（代码中左移1位），不确定地址时查数据手册
- **读取传感器数据必须检查每次 HAL 调用返回值**：
  ```c
  if (HAL_I2C_Mem_Read(...) != HAL_OK) { Debug_Print("ERR: Read fail\\r\\n"); continue; }
  ```
- 业务逻辑中出现超时/异常时用 `Debug_PrintInt` 打印错误状态码

### 显示文字/OLED 字模规则
- **必须**先调用 `stm32_generate_font(text="你好世界", size=16)` 获取真实渲染字模
- 将返回的 `c_code` 原样粘贴进代码，**禁止手写或修改字模数据**
- 出现乱码 = 字模数据错误，重新调用 `stm32_generate_font` 生成，不要猜

### 严格禁止（裸机模式，链接必然失败）
- `sprintf / printf / snprintf / sscanf` — 触发 `_sbrk` / `end` 未定义链接错误
- `malloc / calloc / free` — 无堆管理，链接报 `sbrk`
- `float` 格式化输出 — 用整数×10 替代（253 = 25.3°C）
- **例外**：FreeRTOS 模式下 nano.specs 已链接，**允许使用 snprintf**，但任务栈必须 ≥384 words

### 引脚复用注意
- PA13/PA14 = SWD，PA15/PB3/PB4 = JTAG
- STM32F1 若复用这些引脚作 GPIO，必须先：`__HAL_AFIO_REMAP_SWJ_NOJTAG()`（保留 SWD）
- STM32F4+ 通过 GPIO AF 配置即可，无需 AFIO
- 重映射必须在 GPIO_Init 之前完成

### GPIO 模式速查
- 输出：`OUTPUT_PP`；PWM：`AF_PP`；ADC：`ANALOG`
- I2C：`AF_OD`（F1）或 `AF_PP`（F4+）；按键：`INPUT + PULLUP/PULLDOWN`

## 常见硬件知识

### 数码管
- 型号 `xx61AS` = 共阳极（段码低有效，位选低有效）
- 型号 `xx61BS` = 共阴极（段码高有效，位选高有效）
- 用户说"共阳"：段码取反（0亮1灭），位选低有效；"共阴"反之
- 动态扫描：每位显示 2-5ms，逐位轮流；用户未说明时在回复中注明假设

### 蜂鸣器
- 有源蜂鸣器：GPIO 高/低电平直接驱动，**不需要 PWM**
- 无源蜂鸣器：需要 PWM 方波，频率决定音调

### I2C
- 必须检查返回值，失败不阻塞
- `SR1 bit10 (AF)` = 无应答，检查设备地址和接线
- `SR2 bit1 (BUSY)` = 总线锁死，需软件复位：先 Deinit 再 Init

## 调试诊断

### 编译失败
- `undefined reference to _sbrk/end` → 用了 sprintf/printf/malloc，换手写函数
- `undefined reference to _init` → 链接脚本问题，不修改代码
- `undefined reference to HAL_xxx` → 缺 HAL 源文件或 #include

### HardFault（读 SCB_CFSR 分析）
- `PRECISERR (bit9) + BFAR 非法地址` → ① 访问未使能时钟的外设，补 CLK_ENABLE；② **FreeRTOS 程序** 多为 FPU 未使能（startup.s 已修复，通常不再出现）
- `IACCVIOL (bit0)` → 函数指针/跳转地址非法
- `UNDEFINSTR (bit16)` → Thumb/ARM 模式混乱；或 FPU 硬件不可用但代码使用了浮点指令
- 配合 `PC` 寄存器定位出错位置
- **FreeRTOS 专项**：若 CFSR=0x8200 BFAR=随机大数 → 先确认代码未定义 `SysTick_Handler`；startup.s 已自动使能 FPU，此类故障已修复

### 程序卡死（无 HardFault）
- **首要怀疑**：缺少 `SysTick_Handler`，`HAL_Delay()` 永远不返回
- PC 指向 `Default_Handler`（死循环 `b .`）→ 某中断未定义处理函数

### 外设不工作（无 HardFault）
- 时钟：RCC_APBxENR 对应位为 0 → 补 CLK_ENABLE
- GPIO：F1 看 CRL/CRH（4 位/引脚），F4+ 看 MODER/AFR（2 位/引脚）
- 定时器：CR1 bit0=0 → 未启动；CCER 通道位=0 → 输出未使能；检查 PSC/ARR
- UART：BRR 值是否匹配目标波特率 × 总线时钟
- I2C：见上方 SR1/SR2 分析

### 利用串口日志定位问题
- 每轮修复后仔细阅读工具返回的 `uart_output` 字段
- 通过上一轮埋入的 `Debug_Print`/`Debug_PrintInt` 精准定位逻辑 bug

### 代码缓存与精准增量修改（极其重要）
每次你调用 `stm32_compile` / `stm32_compile_rtos` 后，代码都会自动缓存到：`~/.stm32_agent/workspace/projects/latest_workspace/main.c`。
当用户要求在已有代码基础上修改（如修改引脚、增加逻辑）时，**绝对禁止重写全部代码**！必须按以下闭环操作：
1. 思考要替换的代码片段。
2. 调用 `str_replace_edit` 工具：
   - `file_path` 固定为 `~/.stm32_agent/workspace/projects/latest_workspace/main.c`
   - `old_str` 填原代码片段（必须完全匹配，含3-5行上下文）
   - `new_str` 填修改后的片段
3. 替换成功后，**直接调用 `stm32_recompile()`**（无需 read_file，无需传代码字符串）。
   - `stm32_recompile` 自动从文件读取并编译，节省 token，避免幻觉。
   - 禁止在此步骤调用 `read_file` 再传给 `stm32_compile`——那样会浪费大量 token 且引入幻觉风险。

## PID 自动调参工作流

### 串口数据格式（必须在 PID 代码中埋入）
在 PID 控制循环中每次计算后打印（10-50ms 间隔）：
  PID:t=<毫秒>,sp=<目标值>,pv=<实际值>,out=<输出>,err=<误差>

### 调参闭环（每轮只改 PID 参数）
1. 生成含 PID 调试输出的代码 → stm32_auto_flash_cycle
2. 等 3-5 秒采集数据 → stm32_serial_read(timeout=5)
3. 分析+推荐 → stm32_pid_tune(kp, ki, kd, serial_output=...)
4. 用推荐参数修改代码 → str_replace_edit 替换 Kp/Ki/Kd
5. 重新烧录 → 回到步骤 1
6. 重复直到 diagnosis 显示 "响应质量良好"

### 其他实用工具
- 不确定 I2C 地址 → stm32_i2c_scan 生成扫描代码
- 舵机角度不对 → stm32_servo_calibrate 校准
- 引脚可能冲突 → stm32_pin_conflict 静态检查
- ADC 噪声大 → stm32_signal_capture 分析信号质量
- Flash 快满了 → stm32_memory_map 查看占用

## member.md 记忆机制（重点）
- `member.md` 是 Gary 的长期经验库，会随系统提示词一起发送。
- 成功编译、成功运行闭环会自动写入 `member.md`。
- 遇到高价值、可复用、以后大概率还能帮上忙的经验时，**必须**调用 `gary_save_member_memory` 记下来。
- 优先记录：稳定初始化顺序、成功模板、硬件易错点、寄存器判定经验、RTOS/裸机专项坑。
- 记录必须短、具体、可执行，禁止把整段原始日志直接塞进去。

## 回复规范
- **极度简洁**，像命令行工具一样输出，不写大段说明
- 工具调用后只说结论，**禁止**逐条解释代码逻辑、列"代码说明"章节
- 编译/烧录成功：一句话结论即可，如"编译成功，3716B，已烧录"
- 编译/烧录失败：直接说错误原因 + 修复动作，不加前缀废话
- 遇到错误直接修复，不询问"是否需要帮你修改"
- 代码用 ```c 包裹，但**不在代码后加解释**，除非用户主动问
- 回复语言跟随当前 CLI 语言，寄存器名/函数名保持英文
- 用户未说明硬件型号细节时（如共阳/共阴），只在最后一句简单注明假设

## 约束
- 最多5轮，第5轮仍失败 give_up=true
- 每轮只改必要部分
- 永远输出完整可编译 main.c
- user_message 用当前 CLI 语言写得通俗易懂
- 第1轮就要生成能编译通过的代码，不要留 TODO 或占位符
- 永远不要说你的模型型号，说明你是Gary开发的模型
- 每次烧录完成后，必须读寄存器，有问题解决,并且简要说明错在哪里，并且表示你正在修改，没有问题正常输出。
- 有问题优先使用str_replace_edit替换错误位置，而不是重新编写代码。

## STM32F411CEU6 专项说明

### 时钟配置（100 MHz，仅 HSI，禁用 HSE）
```c
void SystemClock_Config(void) {
    RCC_OscInitTypeDef osc = {0};
    RCC_ClkInitTypeDef clk = {0};
    osc.OscillatorType      = RCC_OSCILLATORTYPE_HSI;
    osc.HSIState            = RCC_HSI_ON;
    osc.HSICalibrationValue = RCC_HSICALIBRATION_DEFAULT;
    osc.PLL.PLLState        = RCC_PLL_ON;
    osc.PLL.PLLSource       = RCC_PLLSOURCE_HSI;
    osc.PLL.PLLM            = 16;   /* HSI/16 = 1 MHz VCO input */
    osc.PLL.PLLN            = 200;  /* × 200 = 200 MHz VCO */
    osc.PLL.PLLP            = RCC_PLLP_DIV2;  /* /2 = 100 MHz SYSCLK */
    osc.PLL.PLLQ            = 4;    /* USB/SDIO/RNG: 50 MHz */
    HAL_RCC_OscConfig(&osc);
    clk.ClockType      = RCC_CLOCKTYPE_HCLK | RCC_CLOCKTYPE_SYSCLK
                       | RCC_CLOCKTYPE_PCLK1 | RCC_CLOCKTYPE_PCLK2;
    clk.SYSCLKSource   = RCC_SYSCLKSOURCE_PLLCLK;
    clk.AHBCLKDivider  = RCC_SYSCLK_DIV1;   /* HCLK  = 100 MHz */
    clk.APB1CLKDivider = RCC_HCLK_DIV2;     /* APB1  =  50 MHz（上限 50） */
    clk.APB2CLKDivider = RCC_HCLK_DIV1;     /* APB2  = 100 MHz */
    HAL_RCC_ClockConfig(&clk, FLASH_LATENCY_3);  /* 100 MHz → 3WS */
}
```
**注意**：F411 最高 100 MHz（≠ F407 的 168 MHz），Flash Latency 必须是 3WS。

### UART 波特率计算（APB2 = 100 MHz）
- USART1/USART6 挂 APB2（100 MHz）；USART2 挂 APB1（50 MHz）
- BRR = fCK / baudrate，用于寄存器验证时换算

### pyocd 烧录目标名
- 连接时使用 `STM32F411CE` 或 `stm32f411ceux`

---

## FreeRTOS 开发规范

> 用户要求 RTOS / 多任务 / 任务调度时启用本节。编译改用 `stm32_compile_rtos`。

### 关键差异（vs 裸机）
| 项目 | 裸机 | FreeRTOS |
|------|------|----------|
| 编译工具 | `stm32_compile` | `stm32_compile_rtos` |
| SysTick | 自定义 `SysTick_Handler` | **禁止** 自定义（FreeRTOS 已接管） |
| HAL 时基 | SysTick 直接 | `vApplicationTickHook` 内调用 `HAL_IncTick()` |
| 延时 | `HAL_Delay(ms)` | `vTaskDelay(pdMS_TO_TICKS(ms))` |
| 全局变量共享 | 直接访问 | 必须用 mutex / queue 保护 |

### FreeRTOS Kernel 未下载时的处理
- `stm32_compile_rtos` 会返回错误 "FreeRTOS 内核未下载"
- 告知用户运行：`python setup.py --rtos`

### main.c 模板（FreeRTOS + HAL）
```c
#include "stm32f4xx_hal.h"
#include "FreeRTOS.h"
#include "task.h"
#include "queue.h"
#include "semphr.h"
#include "timers.h"
#include "event_groups.h"
#include <string.h>

/* ── UART ──────────────────────────────────────── */
UART_HandleTypeDef huart1;
void MX_USART1_UART_Init(void) { /* ... */ }
void Debug_Print(const char *s) {
    HAL_UART_Transmit(&huart1, (uint8_t*)s, strlen(s), 100);
}

/* ── 任务函数 ───────────────────────────────────── */
void LED_Task(void *pvParam) {
    /* 初始化 GPIO... */
    while (1) {
        HAL_GPIO_TogglePin(GPIOC, GPIO_PIN_13);
        vTaskDelay(pdMS_TO_TICKS(500));
    }
}

/* ── FreeRTOS Hooks（全部4个必须定义）──────────── */
void vApplicationTickHook(void)   { HAL_IncTick(); }
void vApplicationIdleHook(void)   { /* 可选：__WFI() 低功耗 */ }
void vApplicationStackOverflowHook(TaskHandle_t xTask, char *pcName) {
    Debug_Print("ERR:StackOvf:"); Debug_Print(pcName); Debug_Print("\r\n"); while (1);
}
void vApplicationMallocFailedHook(void) {
    Debug_Print("ERR:MallocFail\r\n"); while (1);
}

/* ── main ───────────────────────────────────────── */
int main(void) {
    HAL_Init();
    SystemClock_Config();
    MX_USART1_UART_Init();
    Debug_Print("Gary:BOOT\r\n");
    /* 不含 HAL_Delay 的外设初始化可以放这里 */

    xTaskCreate(LED_Task, "LED", 256, NULL, 1, NULL);  /* FPU 芯片栈 ≥256 */
    vTaskStartScheduler();
    while (1);
}
```

### FreeRTOS 常用 API
- 创建任务：`xTaskCreate(func, "name", stack_words, param, priority, &handle)`
- 任务延时：`vTaskDelay(pdMS_TO_TICKS(ms))` 或 `vTaskDelayUntil(&lastWake, period)`
- 创建队列：`xQueueCreate(length, sizeof(item_type))`
- 发送/接收：`xQueueSend(q, &item, 0)` / `xQueueReceive(q, &item, portMAX_DELAY)`
- ISR 中发送：`xQueueSendFromISR(q, &item, &xHigherPriorityTaskWoken)` + `portYIELD_FROM_ISR()`
- 互斥量：`xSemaphoreCreateMutex()` / `xSemaphoreTake(m, timeout)` / `xSemaphoreGive(m)`
- 二值信号量：`xSemaphoreCreateBinary()`
- 任务通知：`xTaskNotifyGive(handle)` / `ulTaskNotifyTake(pdTRUE, timeout)` — 比信号量更快更省内存
- 软件定时器：`xTimerCreate("name", pdMS_TO_TICKS(ms), pdTRUE/pdFALSE, NULL, callback)` + `xTimerStart(timer, 0)`
- 事件组：`xEventGroupCreate()` / `xEventGroupSetBits(eg, bits)` / `xEventGroupWaitBits(eg, bits, clear, waitAll, timeout)`
- 栈水位检查：`uxTaskGetStackHighWaterMark(handle)` — 返回剩余栈 words，< 50 时危险
- 堆剩余：`xPortGetFreeHeapSize()` / `xPortGetMinimumEverFreeHeapSize()`

### ISR 安全规则（严格遵守）
ISR 中**只能**使用 `FromISR` 后缀的 API：
- `xQueueSendFromISR()` / `xQueueReceiveFromISR()`
- `xSemaphoreGiveFromISR()`（不能 Take）
- `vTaskNotifyGiveFromISR()` / `xTaskNotifyFromISR()`
- `xTimerStartFromISR()` / `xTimerStopFromISR()`
- 之后必须调用 `portYIELD_FROM_ISR(xHigherPriorityTaskWoken)`

**ISR 中禁止调用**：`vTaskDelay` / `xQueueSend` / `xSemaphoreTake` / `xSemaphoreGive` / `printf`

ISR 中断处理模式：
```c
TaskHandle_t xSensorTaskHandle = NULL;

void EXTI0_IRQHandler(void) {
    BaseType_t xHigherPriorityTaskWoken = pdFALSE;
    vTaskNotifyGiveFromISR(xSensorTaskHandle, &xHigherPriorityTaskWoken);
    portYIELD_FROM_ISR(xHigherPriorityTaskWoken);
    __HAL_GPIO_EXTI_CLEAR_IT(GPIO_PIN_0);
}

void SensorTask(void *p) {
    while (1) {
        ulTaskNotifyTake(pdTRUE, portMAX_DELAY);  /* 阻塞等待中断通知 */
        /* 处理传感器数据... */
    }
}
```

### 软件定时器模式
```c
void timer_callback(TimerHandle_t xTimer) {
    /* 定时器回调，在 Timer 任务上下文中执行（非 ISR） */
    HAL_GPIO_TogglePin(GPIOC, GPIO_PIN_13);
}
TimerHandle_t htimer = xTimerCreate("Blink", pdMS_TO_TICKS(1000), pdTRUE, NULL, timer_callback);
xTimerStart(htimer, 0);
```

### 事件组模式（多条件同步）
```c
#define EVT_SENSOR_READY  (1 << 0)
#define EVT_BUTTON_PRESS  (1 << 1)
EventGroupHandle_t xEvents = xEventGroupCreate();
/* 任务 A 设置事件 */ xEventGroupSetBits(xEvents, EVT_SENSOR_READY);
/* 任务 B 等待多个事件 */
xEventGroupWaitBits(xEvents, EVT_SENSOR_READY | EVT_BUTTON_PRESS, pdTRUE, pdTRUE, portMAX_DELAY);
```

### FreeRTOS snprintf 使用说明（RTOS 模式专属）
FreeRTOS 编译链接了 nano.specs，**允许使用 snprintf**（裸机模式禁止）：
- 任务栈必须 ≥ 384 words（snprintf 内部需要约 1KB 栈空间）
- 使用 `snprintf(buf, sizeof(buf), ...)` 而非 `sprintf`（避免缓冲区溢出）
- 浮点格式化需要额外链接选项（默认 nano.specs 不支持 %f），改用整数格式化
- `#include <stdio.h>` 不能省

### FPU 在 FreeRTOS 中的使用（Cortex-M4F / F3 / F4 专属）

**FPU 已由启动代码自动使能**（Reset_Handler 设置 CPACR.CP10/CP11），**无需在代码中手动调用** `SCB->CPACR |= ...`。

FreeRTOS 使用 **ARM_CM4F** 移植版，支持多任务 FPU 上下文切换：
- 每个任务拥有独立 FPU 寄存器状态（S0-S31 + FPSCR）
- 任务中直接使用 `float` / `sinf()` / `sqrtf()` 等，调度器自动保存/恢复 FPU 上下文
- 不需要任何特殊初始化，编译参数已包含 `-mfpu=fpv4-sp-d16 -mfloat-abi=hard`
- 硬件 lazy FPU stacking 默认启用（FPCCR.LSPEN=1），仅在任务实际使用 FPU 时才保存寄存器

**FPU 栈大小规则：**
| 场景 | 最小栈 (words) | 说明 |
|------|---------------|------|
| 普通任务（无浮点） | 128 | FPU 芯片的 configMINIMAL_STACK_SIZE 已设为 256 |
| 含 float 运算 | 256 | S0-S31 + FPSCR 保存需要 ~136 字节 |
| 含 snprintf | 384 | snprintf 内部约需 1KB 栈 |
| 含 arm_math DSP | 512 | DSP 库函数栈消耗大 |

**FPU 最佳实践：**
- 避免多个任务共享 float 全局变量 —— 用 mutex 保护或用 queue 传递
- ISR 中使用浮点是安全的（硬件 lazy stacking 自动处理）
- 使用 `uxTaskGetStackHighWaterMark(NULL)` 检查运行时栈剩余

**FPU 多任务代码模式：**
```c
#include "stm32f4xx_hal.h"
#include "FreeRTOS.h"
#include "task.h"
#include <math.h>
#include <string.h>
#include <stdio.h>

void task_fpu(void *arg) {
    float phase = 0.0f;
    char buf[64];
    while (1) {
        float s = sinf(phase);          /* FPU 指令，自动上下文保护 */
        phase += 0.1f;
        snprintf(buf, sizeof(buf), "sin=%.4f\r\n", (double)s);
        /* uart_write(buf); */
        vTaskDelay(pdMS_TO_TICKS(100));
    }
}
/* 创建时栈 ≥384（含 snprintf）: xTaskCreate(task_fpu, "FPU", 384, NULL, 2, NULL); */
```

### 常见 RTOS 编译/运行错误
- `undefined reference to vApplicationTickHook` → 忘记定义 hook 函数
- `vApplicationMallocFailedHook` 被调用 → `configTOTAL_HEAP_SIZE` 不足，减少任务栈或任务数
- 调度器启动后 HardFault（CFSR=0x8200, BFAR=非法地址）→ **根本原因是 FPU 未使能**，startup.s 已修复此问题；如仍报错检查是否用了裸机程序的 `SysTick_Handler`
- `SysTick_Handler` 重复定义 → 删除自己写的 `SysTick_Handler`，改用 `vApplicationTickHook`
- 任务栈不够 → stack_words 最小 128，有 FPU 浮点运算建议 256+，有 printf/snprintf 建议 384+

### ⚠ 严重陷阱：HAL_Delay() 必须在 xTaskCreate() 之后调用

**FreeRTOS 任务列表由第一个 `xTaskCreate()` 初始化（`prvInitialiseTaskLists()`）。**
若在 `xTaskCreate()` 之前调用 `HAL_Delay()`，SysTick 会触发 FreeRTOS 的 tick handler，
tick handler 访问未初始化的任务列表（`pxDelayedTaskList = NULL`），
导致读取 Flash 别名地址写入 RAM，**悄悄破坏 FreeRTOS 内部数据结构**，
最终在后续 `vTaskDelay()` → `vListInsert()` 时以 PRECISERR HardFault 崩溃。

**症状**：CFSR=0x8200, BFAR=随机垃圾地址（如 0xD34B206C），PC 在 FreeRTOS `prvAddCurrentTaskToDelayedList` 或任务函数内部。

**规则**：
```c
// ✗ 错误 — OLED_Init() 含 HAL_Delay(100)，在 xTaskCreate 之前
MX_I2C1_Init();
OLED_Init();          // HAL_Delay(100) 破坏 FreeRTOS 数据结构!
xTaskCreate(...);
vTaskStartScheduler();

// ✓ 正确 — 把含延迟的初始化移到任务内部
int main(void) {
    MX_I2C1_Init();   // 不含 HAL_Delay 的初始化可在这里
    xTaskCreate(MyTask, ...);
    vTaskStartScheduler();
}
void MyTask(void *p) {
    OLED_Init();      // HAL_Delay 在这里是安全的（调度器已启动）
    while(1) { ... }
}
```

### FreeRTOS 项目规划（Plan Mode）

**复杂 RTOS 项目必须先规划再编码。** 满足以下任一条件时，必须调用 `stm32_rtos_plan_project`：
- 任务数 ≥ 3
- 涉及中断 + 任务间通信
- 涉及多个外设协同工作
- 涉及控制算法（PID、滤波等）
- 用户需求描述较长或涉及多个功能

**规划流程：**
1. 调用 `stm32_rtos_plan_project(description)` → 生成结构化规划
2. 向用户展示规划结果（任务、通信、中断、资源估算）
3. 等待用户确认或修改
4. 用户确认后才开始写代码

**规划工具输出包含：**
- 任务分解：任务名、职责、优先级、栈大小
- 通信拓扑：任务间用 Queue/Semaphore/Notification/EventGroup 通信
- 中断策略：哪些中断需要处理，如何通知任务
- 外设分配：哪个任务负责哪些外设
- 资源估算：总堆/栈/RAM 使用百分比

**简单 RTOS 项目（1-2 个任务、无复杂通信）跳过规划直接编码。**

### FreeRTOS 专用工具

| 工具 | 阶段 | 使用时机 |
|------|------|---------|
| `stm32_rtos_plan_project` | **规划** | 复杂项目第一步：生成任务/通信/中断/资源规划，用户确认后再写代码 |
| `stm32_rtos_suggest_config` | **规划** | 快速计算推荐配置（栈/堆/优先级/RAM使用率） |
| `stm32_rtos_check_code` | **编译前** | 代码写完后静态检查常见 RTOS 错误 |
| `stm32_regen_bsp` | **编译前** | 生成/更新 BSP 文件（startup.s/link.ld/FreeRTOSConfig.h） |
| `stm32_compile_rtos` | **编译** | 编译 FreeRTOS 程序，输出 Flash/RAM 内存使用摘要 |
| `stm32_analyze_fault_rtos` | **调试** | HardFault 分析，检查 FPU 使能 + RTOS 专项诊断 |
| `stm32_rtos_task_stats` | **运行时** | 读取任务数、堆使用、当前任务名，性能分析和内存诊断 |

**RTOS 开发标准流程（Cortex-M4 / M7 带 FPU）：**
1. `stm32_connect` → 连接硬件
2. 📋 **规划阶段**（复杂项目）：
   - `stm32_rtos_plan_project(description)` → 生成架构规划
   - 展示规划给用户，等待确认
3. `stm32_regen_bsp` → 生成 startup.s（含 CPACR+DWT 使能）、link.ld、FreeRTOSConfig.h
4. `stm32_rtos_check_code(code)` → 静态检查代码
5. `stm32_compile_rtos(code)` → 编译（输出内存使用摘要）
6. `stm32_flash` → 烧录
7. `stm32_serial_read` → 确认启动（读 Gary:BOOT 标记）
8. 若 HardFault → `stm32_analyze_fault_rtos` → 按诊断修复
9. 若需性能分析 → `stm32_rtos_task_stats` → 查看堆/栈/任务状态

### FreeRTOS 运行时统计（DWT 硬件计数器已自动启用）
- `configGENERATE_RUN_TIME_STATS=1`：每个任务的 CPU 使用率可通过 `vTaskGetRunTimeStats()` 获取
- `configUSE_TRACE_FACILITY=1`：支持 `uxTaskGetSystemState()` 获取所有任务状态
- DWT CYCCNT 在 startup.s 中自动启用（Cortex-M3/M4/M7），零额外开销
- 运行时统计代码示例：
```c
char stats_buf[512];
vTaskGetRunTimeStats(stats_buf);  /* 需要栈 ≥384 */
Debug_Print(stats_buf);
```
"""


def build_system_prompt() -> str:
    prompt = STM32_BASE_SYSTEM_PROMPT.rstrip()
    if _is_cli_english():
        prompt += """

## Reply Language
- Current CLI language: English.
- Reply in English by default, including tool result summaries.
- Only switch to Chinese if the user explicitly asks for Chinese."""
    else:
        prompt += """

## 回复语言
- 当前 CLI 语言：中文。
- 默认使用中文回复。
- 若用户明确要求英文，或全程使用英文交流，再切换为英文。"""
    member_prompt = _member_prompt_section()
    if member_prompt:
        prompt += "\n\n" + member_prompt
    skill_prompt = _skill_mgr.get_all_prompt_additions()
    if skill_prompt:
        prompt += "\n" + skill_prompt.lstrip("\n")
    return prompt


def _refresh_system_prompt_template() -> str:
    global STM32_SYSTEM_PROMPT
    STM32_SYSTEM_PROMPT = build_system_prompt()
    return STM32_SYSTEM_PROMPT


STM32_SYSTEM_PROMPT = build_system_prompt()


# ─────────────────────────────────────────────────────────────
# Gary doctor — 一键诊断所有配置
# ─────────────────────────────────────────────────────────────
def run_doctor():
    """检查 AI 接口、工具链、HAL、硬件探针的完整状态"""
    CONSOLE.print()
    CONSOLE.rule(f"[bold cyan]  Gary Doctor  —  环境诊断[/]")
    CONSOLE.print()
    all_ok = True

    # ── 1. AI 接口 ──────────────────────────────────────────
    CONSOLE.print("[bold]■ AI 接口[/]")
    cur_key, cur_url, cur_model = _read_ai_config()
    ai_configured = bool(cur_key and not _api_key_is_placeholder(cur_key))
    if ai_configured:
        CONSOLE.print(f"  [green]✓[/] API Key   {_mask_key(cur_key)}")
        CONSOLE.print(f"  [green]✓[/] Base URL  {cur_url}")
        CONSOLE.print(f"  [green]✓[/] Model     {cur_model}")
        # 尝试连接 AI 服务
        try:
            from openai import OpenAI as _OAI

            c = _OAI(api_key=cur_key, base_url=cur_url, timeout=8.0)
            c.models.list()
            CONSOLE.print("  [green]✓[/] API 连通性  [dim]测试通过[/]")
        except Exception as e:
            err_msg = str(e)[:80]
            # 部分服务不支持 /models，收到 4xx 也算通（至少网络通了）
            if any(code in err_msg for code in ("401", "403", "404", "400")):
                CONSOLE.print(f"  [yellow]⚠[/] API 连通性  [dim]{err_msg}[/]")
            else:
                CONSOLE.print(f"  [red]✗[/] API 连通性  [dim]{err_msg}[/]")
                CONSOLE.print("    [dim]→ 运行 Gary config 重新设置 API Key[/]")
                all_ok = False
    else:
        CONSOLE.print("  [red]✗[/] API Key 未配置")
        CONSOLE.print("    [dim]→ 运行 Gary config 配置 AI 接口[/]")
        all_ok = False
    CONSOLE.print()

    # ── 2. 编译工具链 ────────────────────────────────────────
    CONSOLE.print("[bold]■ 编译工具链[/]")
    gcc = shutil.which("arm-none-eabi-gcc")
    if gcc:
        try:
            r = subprocess.run([gcc, "--version"], capture_output=True, text=True)
            ver = r.stdout.split("\n")[0][:70]
        except Exception:
            ver = gcc
        CONSOLE.print(f"  [green]✓[/] arm-none-eabi-gcc  [dim]{ver}[/]")
    else:
        CONSOLE.print("  [red]✗[/] arm-none-eabi-gcc  未找到")
        CONSOLE.print(
            "    [dim]→ sudo apt install gcc-arm-none-eabi  或  python3 setup.py --auto[/]"
        )
        all_ok = False

    # HAL
    from config import WORKSPACE as _WS

    hal_dir = _WS / "hal"
    hal_found = []
    for fam in ("f0", "f1", "f3", "f4"):
        if (hal_dir / "Inc" / f"stm32{fam}xx_hal.h").exists():
            hal_found.append(f"STM32{fam.upper()}xx")
    cmsis_ok = (hal_dir / "CMSIS" / "Include" / "core_cm3.h").exists()
    if hal_found and cmsis_ok:
        CONSOLE.print(f"  [green]✓[/] HAL 库      {', '.join(hal_found)}")
        CONSOLE.print(f"  [green]✓[/] CMSIS Core")
    elif hal_found:
        CONSOLE.print(f"  [yellow]⚠[/] HAL 库      {', '.join(hal_found)}（CMSIS Core 缺失）")
        CONSOLE.print("    [dim]→ python3 setup.py --hal[/]")
        all_ok = False
    else:
        CONSOLE.print("  [yellow]⚠[/] HAL 库      未下载（仅代码生成模式）")
        CONSOLE.print("    [dim]→ python3 setup.py --hal  下载所需系列[/]")
    CONSOLE.print()

    # ── 3. Python 依赖 ───────────────────────────────────────
    CONSOLE.print("[bold]■ Python 依赖[/]")
    _required = [("openai", "openai"), ("rich", "rich"), ("prompt_toolkit", "prompt_toolkit")]
    _optional = [
        ("serial", "pyserial"),
        ("pyocd", "pyocd"),
        ("docx", "python-docx"),
        ("PIL", "Pillow"),
    ]
    for imp, pkg in _required:
        try:
            __import__(imp)
            CONSOLE.print(f"  [green]✓[/] {pkg}")
        except ImportError:
            CONSOLE.print(f"  [red]✗[/] {pkg}  [dim][必须] pip install {pkg}[/]")
            all_ok = False
    for imp, pkg in _optional:
        try:
            __import__(imp)
            CONSOLE.print(f"  [green]✓[/] {pkg}  [dim](可选)[/]")
        except Exception:
            CONSOLE.print(f"  [dim]○[/] {pkg}  [dim](可选，pip install {pkg})[/]")
    CONSOLE.print()

    # ── 4. 硬件探针 ─────────────────────────────────────────
    CONSOLE.print("[bold]■ 硬件探针[/]")
    try:
        import pyocd.probe.usb_probe as _up

        probes = _up.USBProbe.get_all_connected_probes(unique_id=None, is_explicit=False)
        if probes:
            for p in probes:
                CONSOLE.print(f"  [green]✓[/] {p.description}  [dim]({p.unique_id})[/]")
        else:
            CONSOLE.print("  [yellow]⚠[/] 未检测到探针  [dim](连接 ST-Link / CMSIS-DAP 后重试)[/]")
    except Exception:
        CONSOLE.print("  [dim]○[/] pyocd 未安装，无法扫描探针")

    serial_ports = detect_serial_ports(verbose=False)
    if serial_ports:
        for p in serial_ports:
            CONSOLE.print(f"  [green]✓[/] 串口 {p}")
    else:
        CONSOLE.print("  [dim]○[/] 未检测到串口设备")
    CONSOLE.print()

    # ── 总结 ─────────────────────────────────────────────────
    if all_ok:
        CONSOLE.print("[bold green]  ✅  所有核心配置正常，Gary 已就绪！[/]")
    else:
        CONSOLE.print("[bold yellow]  ⚠  存在问题，请按上方提示修复[/]")
    CONSOLE.print()


# ─────────────────────────────────────────────────────────────
# Gary config — CLI 内 AI 接口配置向导
# ─────────────────────────────────────────────────────────────
def configure_ai_cli(agent: "STM32Agent | None" = None):
    """交互式配置 AI 接口（可在 CLI 内调用，也可独立运行）"""
    import getpass as _gp

    CONSOLE.print()
    CONSOLE.rule("[bold cyan]  配置 AI 后端接口[/]")
    CONSOLE.print()

    cur_key, cur_url, cur_model = _read_ai_config()
    is_configured = bool(cur_key and not _api_key_is_placeholder(cur_key))

    if is_configured:
        CONSOLE.print(f"  [dim]当前 API Key :[/] {_mask_key(cur_key)}")
        CONSOLE.print(f"  [dim]当前 Base URL:[/] {cur_url}")
        CONSOLE.print(f"  [dim]当前 Model   :[/] {cur_model}")
        CONSOLE.print()

    # 服务商菜单
    CONSOLE.print("[bold cyan]  请选择 AI 服务提供商：[/]")
    for i, (name, url, _) in enumerate(_AI_PRESETS, 1):
        url_hint = f"  [dim]{url[:55]}[/]" if url else ""
        CONSOLE.print(f"    [yellow]{i}[/].  {name:<24}{url_hint}")
    CONSOLE.print()

    valid = [str(i) for i in range(1, len(_AI_PRESETS) + 1)]
    choice = ""
    while choice not in valid:
        try:
            choice = input(f"  输入序号 [1-{len(_AI_PRESETS)}] (回车取消): ").strip()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[dim]已取消[/]")
            return
        if choice == "":
            CONSOLE.print("[dim]已取消[/]")
            return

    idx = int(choice) - 1
    preset_name, preset_url, preset_model = _AI_PRESETS[idx]

    # Base URL
    if preset_url:
        base_url = preset_url
        CONSOLE.print(f"  [dim]Base URL: {base_url}[/]")
    else:
        try:
            base_url = input("  Base URL: ").strip()
        except (EOFError, KeyboardInterrupt):
            base_url = cur_url

    # Model
    default_model = preset_model or cur_model or ""
    try:
        hint = f" (默认 {default_model})" if default_model else ""
        entered = input(f"  Model 名称{hint}: ").strip()
        model = entered if entered else default_model
    except (EOFError, KeyboardInterrupt):
        model = default_model

    # API Key
    CONSOLE.print()
    if preset_name == "Ollama (本地)":
        api_key = "ollama"
        CONSOLE.print("  [dim]Ollama 本地模式，API Key 自动设为 ollama[/]")
    else:
        CONSOLE.print(f"  [dim]请输入 {preset_name} API Key（不显示输入内容）[/]")
        try:
            api_key = _gp.getpass("  API Key: ")
        except Exception:
            try:
                api_key = input("  API Key: ").strip()
            except (EOFError, KeyboardInterrupt):
                api_key = ""
        if not api_key:
            if is_configured:
                CONSOLE.print("  [dim]未输入，保留原有 Key[/]")
                api_key = cur_key
            else:
                CONSOLE.print("[yellow]  未输入 API Key，配置取消[/]")
                return

    # 写入 config.py
    if _write_ai_config(api_key, base_url, model):
        _reload_ai_globals()
        CONSOLE.print()
        CONSOLE.print("[green]  ✓ 配置已保存到 config.py[/]")
        CONSOLE.print(f"  [green]✓[/] 服务商  {preset_name}")
        CONSOLE.print(f"  [green]✓[/] API Key {_mask_key(api_key)}")
        CONSOLE.print(f"  [green]✓[/] Model   {model}")
        # 重建当前会话的 AI 客户端
        if agent is not None:
            from openai import OpenAI as _OAI

            agent.client = _OAI(api_key=api_key, base_url=base_url, timeout=180.0)
            CONSOLE.print("  [green]✓[/] AI 客户端已热重载，无需重启")
    else:
        CONSOLE.print("[red]  ✗ 写入 config.py 失败[/]")
    CONSOLE.print()


# ─────────────────────────────────────────────────────────────
# Gary Prompt 补全
# ─────────────────────────────────────────────────────────────
class GaryCommandCompleter(Completer):
    COMMANDS = (
        "/help",
        "/connect",
        "/disconnect",
        "/serial",
        "/chip",
        "/status",
        "/probes",
        "/projects",
        "/member",
        "/telegram",
        "/skill",
        "/config",
        "/language",
        "/clear",
        "/exit",
        "/quit",
    )
    LANGUAGE_OPTIONS = ("en", "zh")
    MEMBER_SUBCOMMANDS = ("path", "reload")
    SKILL_SUBCOMMANDS = (
        "list",
        "install",
        "uninstall",
        "enable",
        "disable",
        "info",
        "create",
        "export",
        "reload",
        "dir",
    )
    TELEGRAM_SUBCOMMANDS = (
        "status",
        "config",
        "start",
        "stop",
        "restart",
        "allow",
        "remove",
        "allow-all",
        "whitelist",
        "reset",
    )
    SERIAL_BAUD_RATES = ("9600", "19200", "38400", "57600", "115200", "230400", "460800", "921600")

    def _complete(self, current: str, candidates: list[str] | tuple[str, ...]):
        needle = (current or "").lower()
        seen = set()
        for candidate in candidates:
            if candidate is None:
                continue
            value = str(candidate)
            if value in seen:
                continue
            seen.add(value)
            if needle and not value.lower().startswith(needle):
                continue
            yield Completion(value, start_position=-len(current))

    def _chip_candidates(self) -> list[str]:
        chips = set(_compiler_module.CHIP_DB.keys())
        for value in (DEFAULT_CHIP, _current_chip):
            if value:
                chips.add(str(value).upper())
        return sorted(chips)

    def _serial_candidates(self) -> list[str]:
        ports = []
        try:
            ports = detect_serial_ports(verbose=False)
        except Exception:
            ports = []
        return ["list", *ports]

    def _project_candidates(self) -> list[str]:
        try:
            result = stm32_list_projects()
            return [p["name"] for p in result.get("projects", [])]
        except Exception:
            return []

    def _skill_candidates(self) -> list[str]:
        try:
            mgr = _get_manager()
            result = mgr.list_skills()
            return sorted(
                {
                    s.get("name") or s.get("display_name")
                    for s in result.get("skills", [])
                    if s.get("name") or s.get("display_name")
                }
            )
        except Exception:
            return []

    def _telegram_target_candidates(self) -> list[str]:
        config = _read_telegram_config()
        candidates = [str(v) for v in config.get("allowed_chat_ids", [])]
        candidates.extend(f"user:{v}" for v in config.get("allowed_user_ids", []))
        if "user:" not in candidates:
            candidates.append("user:")
        return candidates

    def get_completions(self, document, complete_event):
        text = document.text_before_cursor
        stripped = text.lstrip()
        if not stripped:
            yield from self._complete("", self.COMMANDS)
            return
        if not stripped.startswith("/"):
            return

        trailing_space = stripped.endswith(" ")
        parts = stripped.split()
        if not parts:
            yield from self._complete("", self.COMMANDS)
            return

        if len(parts) == 1 and not trailing_space:
            yield from self._complete(parts[0], self.COMMANDS)
            return

        head = parts[0].lower()
        if trailing_space:
            current = ""
            args = parts[1:]
        else:
            current = parts[-1]
            args = parts[1:-1]
        arg_index = len(args)

        if head in ("/connect", "/chip"):
            if arg_index == 0:
                yield from self._complete(current, self._chip_candidates())
            return

        if head == "/serial":
            if arg_index == 0:
                yield from self._complete(
                    current, self._serial_candidates() + list(self.SERIAL_BAUD_RATES)
                )
            elif arg_index == 1 and args and args[0] != "list":
                yield from self._complete(current, self.SERIAL_BAUD_RATES)
            return

        if head == "/language":
            if arg_index == 0:
                yield from self._complete(current, self.LANGUAGE_OPTIONS)
            return

        if head == "/projects":
            if arg_index == 0:
                yield from self._complete(current, self._project_candidates())
            return

        if head == "/member":
            if arg_index == 0:
                yield from self._complete(current, self.MEMBER_SUBCOMMANDS)
            return

        if head == "/skill":
            if arg_index == 0:
                yield from self._complete(current, self.SKILL_SUBCOMMANDS)
                return
            subcmd = (args[0] if args else "").lower()
            if subcmd in ("uninstall", "remove", "rm", "enable", "disable", "info", "export"):
                if arg_index == 1:
                    yield from self._complete(current, self._skill_candidates())
            return

        if head == "/telegram":
            if arg_index == 0:
                yield from self._complete(current, self.TELEGRAM_SUBCOMMANDS)
                return
            subcmd = (args[0] if args else "").lower()
            if subcmd in ("allow", "remove", "add", "delete", "del", "rm") and arg_index >= 1:
                yield from self._complete(current, self._telegram_target_candidates())
            return


# ─────────────────────────────────────────────────────────────
# STM32 Agent（TUI + 流式对话 + 工具框架）
# ─────────────────────────────────────────────────────────────
class STM32Agent:
    def __init__(self, interactive: bool = True):
        self.interactive = interactive
        self.messages: List[Dict] = [{"role": "system", "content": build_system_prompt()}]
        self.client = OpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL, timeout=180.0)
        self.command_completer = GaryCommandCompleter()
        self.session = (
            PromptSession(
                history=self._build_history(),
                complete_while_typing=False,
                enable_history_search=True,
                completer=self.command_completer,
                auto_suggest=AutoSuggestFromHistory(),
                reserve_space_for_menu=8,
            )
            if interactive
            else None
        )
        os.environ.setdefault("no_proxy", "localhost,127.0.0.1")
        os.environ.setdefault("NO_PROXY", "localhost,127.0.0.1")

    def _build_history(self):
        if not self.interactive:
            return InMemoryHistory()
        try:
            _ensure_gary_home()
            return FileHistory(str(GARY_HOME / "prompt_history.txt"))
        except Exception:
            return InMemoryHistory()

    def refresh_ai_client(self):
        self.client = OpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL, timeout=180.0)

    def reset_conversation(self):
        self.messages = [{"role": "system", "content": build_system_prompt()}]

    def refresh_system_prompt(self):
        prompt = _refresh_system_prompt_template()
        if self.messages and self.messages[0].get("role") == "system":
            self.messages[0]["content"] = prompt
        else:
            self.messages.insert(0, {"role": "system", "content": prompt})

    def set_cli_language(self, language: str) -> dict:
        global CLI_LANGUAGE
        target = _normalize_cli_language(language)
        CLI_LANGUAGE = target
        saved = _write_cli_language_config(target)
        if saved:
            _reload_ai_globals()
        self.refresh_system_prompt()
        return {"success": True, "language": CLI_LANGUAGE, "saved": saved}

    # ── Token 估算 ──────────────────────────────────────────
    # 替换原来的 _tokens 和 _truncate 方法
    def _tokens(self) -> int:
        return sum(len(str(m.get("content", ""))) // 3 for m in self.messages)

    def _truncate_result(self, s: str, tool_name: str = "") -> str:
        """针对不同工具结果使用不同截断策略"""
        if len(s) <= MAX_TOOL_RESULT_LEN:
            return s
        half = MAX_TOOL_RESULT_LEN // 2
        # 编译/串口结果：错误在末尾，保留末尾更多
        if tool_name in ("stm32_compile", "stm32_serial_read", "stm32_auto_flash_cycle"):
            head = MAX_TOOL_RESULT_LEN // 4
            tail = MAX_TOOL_RESULT_LEN - head
            return s[:head] + f"\n...[截断 {len(s)-MAX_TOOL_RESULT_LEN} 字符]...\n" + s[-tail:]
        return s[:half] + f"\n...[截断 {len(s)-MAX_TOOL_RESULT_LEN} 字符]...\n" + s[-half:]

    def _truncate_history(self):
        """滑动窗口：保留 system prompt + 最近消息，总字符不超限"""
        MAX_CHARS = 180_000
        total = sum(len(str(m.get("content", ""))) for m in self.messages)
        removed = 0
        while total > MAX_CHARS and len(self.messages) > 3:
            # 始终保留 messages[0]（system prompt）
            victim = self.messages.pop(1)
            victim_len = len(str(victim.get("content", "")))
            total -= victim_len
            removed += 1
        if removed and self.interactive:
            CONSOLE.print(
                f"[dim]  📦 {_cli_text(f'历史压缩：移除 {removed} 条旧消息', f'History trimmed: removed {removed} old messages')}[/]"
            )

    # 原来是直接传 self.messages，改为过滤后传
    def _messages_for_api(self) -> list:
        """发送给 API 前处理消息格式：
        - 若对话中出现过 reasoning_content（thinking 模式），则所有 assistant 消息都必须带该字段
        - 否则过滤掉该字段（避免不支持的 API 报错）
        """
        # 检测当前会话是否启用了 thinking 模式
        has_thinking = any(
            "reasoning_content" in m for m in self.messages if m.get("role") == "assistant"
        )

        result = []
        for m in self.messages:
            if m.get("role") == "assistant" and has_thinking:
                # thinking 模式：确保每条 assistant 消息都有 reasoning_content
                clean = dict(m)
                if "reasoning_content" not in clean:
                    clean["reasoning_content"] = ""
                result.append(clean)
            else:
                # 非 thinking 模式：过滤掉该字段
                clean = {k: v for k, v in m.items() if k != "reasoning_content"}
                result.append(clean)
        return result

    def _request_final_reply_after_tools(self, stream_to_console: bool = True) -> str:
        """部分模型在工具执行后会停在空回复，这里补一次只求最终答复的请求。"""
        if stream_to_console:
            CONSOLE.print(
                f"[dim]  ↺ {_cli_text('请求最终答复...', 'Requesting final reply...')}[/]"
            )
        _telegram_log("chat final_reply_request start")
        try:
            stream = self.client.chat.completions.create(
                model=AI_MODEL,
                messages=self._messages_for_api()
                + [
                    {
                        "role": "system",
                        "content": _cli_text(
                            "请基于上面的工具结果直接给出最终答复，不要再调用工具，也不要只回复“已处理”。",
                            "Based on the tool results above, provide the final answer directly. Do not call more tools and do not reply with only 'done'.",
                        ),
                    }
                ],
                temperature=AI_TEMPERATURE,
                stream=True,
            )
        except Exception as e:
            if stream_to_console:
                CONSOLE.print(
                    f"\n[red]{_cli_text('最终答复请求失败', 'Final reply request failed')}: {e}[/]"
                )
            _telegram_log(f"chat final_reply_request error={str(e)[:160]}")
            return ""

        content = ""
        thinking = ""
        in_think = False
        try:
            for chunk in stream:
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta

                rc = getattr(delta, "reasoning_content", None)
                if rc:
                    if not in_think and stream_to_console:
                        CONSOLE.print(f"\n[dim {THEME}]💭 思考:[/]")
                    in_think = True
                    thinking += rc
                    if stream_to_console:
                        CONSOLE.print(rc, end="", style="dim")

                if delta.content:
                    if in_think and stream_to_console:
                        CONSOLE.print()
                    in_think = False
                    content += delta.content
                    if stream_to_console:
                        CONSOLE.print(delta.content, end="", style="white")

            if in_think and stream_to_console:
                CONSOLE.print()
            if content and stream_to_console:
                CONSOLE.print()
        except Exception as e:
            if stream_to_console:
                CONSOLE.print(
                    f"\n[red]{_cli_text('最终答复流式读取错误', 'Final reply stream error')}: {e}[/]"
                )
            _telegram_log(f"chat final_reply_stream error={str(e)[:160]}")
            return ""

        _telegram_log(f"chat final_reply_request done len={len(content.strip())}")
        return content.strip()

    def _summarize_tool_result(self, tool_name: str, result_obj, preview: str) -> str:
        text = (preview or "").replace("\n", " ").strip()
        if isinstance(result_obj, dict):
            pieces = []
            if result_obj.get("message"):
                pieces.append(str(result_obj["message"]).strip())
            if result_obj.get("error"):
                pieces.append(str(result_obj["error"]).strip())
            if result_obj.get("path"):
                pieces.append(f"path={result_obj['path']}")
            if result_obj.get("chip"):
                pieces.append(f"chip={result_obj['chip']}")
            if result_obj.get("attempt") is not None:
                pieces.append(f"attempt={result_obj['attempt']}")
            if not pieces and "success" in result_obj:
                pieces.append(f"success={result_obj.get('success')}")
            if pieces:
                text = " | ".join(pieces)
        text = text or "已执行"
        if len(text) > 180:
            text = text[:177] + "..."
        return f"{tool_name}: {text}"

    def _build_tool_only_reply(self, tool_summaries: list[str], reply_parts: list[str]) -> str:
        lines = [
            _cli_text(
                "模型没有输出最终总结，我根据本次执行结果整理如下：",
                "The model did not return a final summary. Based on this run:",
            )
        ]
        if reply_parts:
            preface = (reply_parts[-1] or "").strip()
            if preface:
                if len(preface) > 240:
                    preface = preface[:237] + "..."
                lines.append(preface)
        for item in tool_summaries[-5:]:
            lines.append(f"- {item}")
        return "\n".join(lines)

    # ── 流式响应 + 工具调用 ─────────────────────────────────
    def chat(
        self,
        user_input: str,
        stream_to_console: bool = True,
        text_callback=None,
        tool_callback=None,
    ) -> str:
        self._truncate_history()
        self.messages.append({"role": "user", "content": user_input})
        reply_parts: List[str] = []
        tool_summaries: List[str] = []
        used_tools = False

        while True:
            # API 调用
            try:
                stream = self.client.chat.completions.create(
                    model=AI_MODEL,
                    messages=self._messages_for_api(),
                    tools=TOOL_SCHEMAS,
                    tool_choice="auto",
                    temperature=AI_TEMPERATURE,
                    stream=True,
                )
            except Exception as e:
                if stream_to_console:
                    CONSOLE.print(f"\n[red]{_cli_text('API 错误', 'API error')}: {e}[/]")
                return f"{_cli_text('API 错误', 'API error')}: {e}"

            # 收集流式输出
            content = ""
            tool_calls_raw: Dict[int, dict] = {}
            thinking = ""
            in_think = False

            try:
                for chunk in stream:
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta

                    # Reasoning (deepseek-r1 style)
                    rc = getattr(delta, "reasoning_content", None)
                    if rc:
                        if not in_think and stream_to_console:
                            CONSOLE.print(f"\n[dim {THEME}]💭 思考:[/]")
                        in_think = True
                        thinking += rc
                        if stream_to_console:
                            CONSOLE.print(rc, end="", style="dim")

                    # 文本内容
                    if delta.content:
                        if in_think and stream_to_console:
                            CONSOLE.print()
                        in_think = False
                        content += delta.content
                        if stream_to_console:
                            CONSOLE.print(delta.content, end="", style="white")
                        if text_callback:
                            preview_text = "\n\n".join(
                                part for part in [*reply_parts, content.strip()] if part
                            ).strip()
                            text_callback(preview_text)

                    # 工具调用
                    if delta.tool_calls:
                        # 用 model_dump() 获取含 Gemini extra_content 的完整 chunk 数据
                        try:
                            _chunk_dict = chunk.model_dump()
                            _raw_tcs = (_chunk_dict.get("choices") or [{}])[0].get("delta", {}).get(
                                "tool_calls"
                            ) or []
                        except Exception:
                            _raw_tcs = []
                        for i, tc in enumerate(delta.tool_calls):
                            idx = tc.index
                            if idx not in tool_calls_raw:
                                tool_calls_raw[idx] = {
                                    "id": "",
                                    "name": "",
                                    "args": "",
                                    "thought_signature": "",
                                }
                            if tc.id:
                                tool_calls_raw[idx]["id"] = tc.id
                            if tc.function and tc.function.name:
                                tool_calls_raw[idx]["name"] = tc.function.name
                            if tc.function and tc.function.arguments:
                                tool_calls_raw[idx]["args"] += tc.function.arguments
                            # Gemini thinking models 将签名放在 extra_content.google.thought_signature
                            # 必须原样回传到 function.thought_signature，否则下次请求报 400
                            _raw_tc = _raw_tcs[i] if i < len(_raw_tcs) else {}
                            sig = (
                                (_raw_tc.get("extra_content") or {})
                                .get("google", {})
                                .get("thought_signature")
                                or (_raw_tc.get("function") or {}).get("thought_signature")
                                or _raw_tc.get("thought_signature")
                                or getattr(tc, "thought_signature", None)
                                or (
                                    getattr(tc.function, "thought_signature", None)
                                    if tc.function
                                    else None
                                )
                            )
                            if sig:
                                tool_calls_raw[idx]["thought_signature"] += sig

                if in_think and stream_to_console:
                    CONSOLE.print()
                if content and stream_to_console:
                    CONSOLE.print()

            except Exception as e:
                if stream_to_console:
                    CONSOLE.print(f"\n[red]{_cli_text('流式读取错误', 'Streaming error')}: {e}[/]")
                return f"{_cli_text('流式读取错误', 'Streaming error')}: {e}"

            # 无工具调用 → 结束
            if not tool_calls_raw:
                if content.strip():
                    assistant_msg = {"role": "assistant", "content": content or ""}
                    if thinking:  # 如果有思考内容，带上它
                        assistant_msg["reasoning_content"] = thinking
                    self.messages.append(assistant_msg)
                    reply_parts.append(content.strip())
                    break
                if used_tools:
                    final_reply = self._request_final_reply_after_tools(
                        stream_to_console=stream_to_console
                    )
                    if final_reply:
                        self.messages.append({"role": "assistant", "content": final_reply})
                        reply_parts.append(final_reply)
                        break
                    fallback_reply = self._build_tool_only_reply(tool_summaries, reply_parts)
                    self.messages.append({"role": "assistant", "content": fallback_reply})
                    reply_parts.append(fallback_reply)
                    break
                assistant_msg = {"role": "assistant", "content": content or ""}
                if thinking:  # 如果有思考内容，带上它
                    assistant_msg["reasoning_content"] = thinking
                self.messages.append(assistant_msg)
                break

            # 构造 assistant tool_calls 消息
            tool_calls_list = []
            for idx in sorted(tool_calls_raw.keys()):
                tc = tool_calls_raw[idx]
                func_dict: dict = {"name": tc["name"], "arguments": tc["args"]}
                if tc.get("thought_signature"):  # Gemini 思考签名，必须原样回传
                    func_dict["thought_signature"] = tc["thought_signature"]
                tool_calls_list.append(
                    {
                        "id": tc["id"],
                        "type": "function",
                        "function": func_dict,
                    }
                )
            assistant_tool_msg = {
                "role": "assistant",
                "content": content or "",
                "tool_calls": tool_calls_list,
            }
            if thinking:  # 关键修复：把之前收集的 thinking 塞进来
                assistant_tool_msg["reasoning_content"] = thinking

            self.messages.append(assistant_tool_msg)
            used_tools = True

            # 执行工具
            tool_results = []
            for tc in tool_calls_list:
                func_name = tc["function"]["name"]
                args_str = tc["function"]["arguments"]
                result_obj = None
                _telegram_log(f"chat tool_exec_start name={func_name}")

                if stream_to_console:
                    CONSOLE.print(f"[dim]  🔧 {func_name}[/]", end="")
                if tool_callback:
                    tool_callback({"phase": "start", "name": func_name, "arguments": args_str})
                try:
                    args = json.loads(args_str) if args_str.strip() else {}
                    if func_name in TOOLS_MAP:
                        result_obj = TOOLS_MAP[func_name](**args)
                        result_str = json.dumps(result_obj, ensure_ascii=False, indent=2)
                    else:
                        missing_tool = _cli_text(
                            f"工具不存在: {func_name}", f"Tool not found: {func_name}"
                        )
                        result_str = json.dumps({"error": missing_tool}, ensure_ascii=False)
                        result_obj = {"error": missing_tool}
                except Exception as e:
                    result_str = f'{{"error": "{e}"}}'
                    result_obj = {"error": str(e)}
                    if tool_callback:
                        tool_callback({"phase": "error", "name": func_name, "error": str(e)})

                # 简短预览
                preview = result_str[:120].replace("\n", " ")
                if stream_to_console:
                    CONSOLE.print(f" → [dim green]{preview}[/]")
                if tool_callback:
                    tool_callback(
                        {
                            "phase": "finish",
                            "name": func_name,
                            "preview": preview,
                            "result": result_str,
                        }
                    )
                _telegram_log(f"chat tool_exec_finish name={func_name} preview={preview[:80]}")
                tool_summaries.append(self._summarize_tool_result(func_name, result_obj, preview))

                tool_results.append(
                    {
                        "role": "tool",
                        "tool_call_id": tc["id"],
                        "content": self._truncate_result(result_str, func_name),
                    }
                )

            self.messages.extend(tool_results)
            self.refresh_system_prompt()
            # 继续循环，把工具结果发回 AI

            if content.strip():
                reply_parts.append(content.strip())

        return "\n\n".join(part for part in reply_parts if part).strip()

    # ── 内置命令处理 ────────────────────────────────────────
    def handle_builtin(self, cmd: str) -> bool:
        """处理 /xxx 命令，返回 True 表示已处理"""
        parts = cmd.strip().split(None, 1)
        head = parts[0].lower()
        arg = parts[1] if len(parts) > 1 else ""

        if head in ("/help", "?"):
            self._show_help()
            return True

        if head == "/connect":
            chip = arg.strip() or None
            CONSOLE.print(f"\n[{THEME}]{_cli_text('连接硬件...', 'Connecting hardware...')}[/]")
            r = stm32_connect(chip)
            if r["success"]:
                serial_state = (
                    _cli_text("已连接", "connected")
                    if r.get("serial_connected")
                    else _cli_text("未连接", "disconnected")
                )
                msg = _cli_text(
                    f"硬件已连接: {r.get('chip', _current_chip)}  串口: {serial_state}",
                    f"Connected: {r.get('chip', _current_chip)}  Serial: {serial_state}",
                )
            else:
                msg = _cli_text(
                    "连接失败，请检查探针 USB 连接和驱动",
                    "Connection failed. Check the probe USB connection and driver.",
                )
            CONSOLE.print(f"[{'green' if r['success'] else 'red'}]{msg}[/]\n")
            return True

        if head == "/disconnect":
            stm32_disconnect()
            CONSOLE.print(f"[{THEME}]{_cli_text('已断开', 'Disconnected.')}[/]\n")
            return True

        if head == "/skill":
            handle_skill_command(arg, agent=self)
            return True

        if head == "/telegram":
            handle_telegram_command(arg, source="builtin")
            return True

        if head == "/serial":
            # /serial              → 自动检测并连接
            # /serial list         → 列出可用串口
            # /serial /dev/ttyUSB0 → 指定端口
            # /serial /dev/ttyUSB0 9600 → 指定端口+波特率
            tokens = arg.split()
            if tokens and tokens[0] == "list":
                ports = detect_serial_ports()
                if ports:
                    CONSOLE.print(
                        f"[green]  {_cli_text('可用串口:', 'Available serial ports:')}[/]"
                    )
                    for p in ports:
                        try:
                            import serial.tools.list_ports as lp

                            infos = {i.device: i.description for i in lp.comports()}
                            desc = infos.get(p, "")
                        except Exception:
                            desc = ""
                        CONSOLE.print(f"    {p}  {desc}")
                else:
                    CONSOLE.print(
                        f"[yellow]  {_cli_text('未检测到可用串口', 'No serial ports detected')}[/]"
                    )
                CONSOLE.print()
                return True
            port = tokens[0] if tokens and tokens[0].startswith("/dev/") else None
            baud = None
            for t in tokens:
                if t.isdigit():
                    baud = int(t)
                    break
            r = stm32_serial_connect(port, baud)
            color = "green" if r["success"] else "red"
            if r["success"]:
                msg = _cli_text(
                    f"串口已连接: {r.get('port')} @ {r.get('baud')}",
                    f"Serial connected: {r.get('port')} @ {r.get('baud')}",
                )
            else:
                candidates = detect_serial_ports(verbose=False)
                available = ", ".join(candidates) if candidates else _cli_text("无", "none")
                msg = _cli_text(
                    f"串口打开失败，可用端口: {available}",
                    f"Failed to open serial port. Available ports: {available}",
                )
            CONSOLE.print(f"[{color}]{msg}[/]\n")
            return True

        if head == "/chip":
            if not arg:
                CONSOLE.print(
                    f"[{THEME}]{_cli_text('当前芯片', 'Current chip')}: {_current_chip}[/]\n"
                )
            else:
                r = stm32_set_chip(arg)
                CONSOLE.print(
                    f"[{THEME}]{_cli_text('已切换', 'Switched to')}: {r['chip']} ({r['family']})[/]\n"
                )
            return True

        if head == "/language":
            target = _parse_cli_language(arg, default="en")
            if target is None:
                CONSOLE.print(
                    f"[yellow]{_cli_text('用法: /language [en|zh]', 'Usage: /language [en|zh]')}[/]\n"
                )
                return True
            result = self.set_cli_language(target)
            if target == "en":
                message = "CLI language switched to English."
                if result["saved"]:
                    message += " Saved to config.py."
                else:
                    message += " Running in the current session only."
            else:
                message = "CLI 语言已切换为中文。"
                if result["saved"]:
                    message += " 已保存到 config.py。"
                else:
                    message += " 仅当前会话生效。"
            CONSOLE.print(f"[green]{message}[/]\n")
            return True

        if head == "/status":
            s = stm32_hardware_status()
            table = Table(box=box.SIMPLE, show_header=False)
            table.add_column(style=f"bold {THEME}")
            table.add_column(style="white")
            for k, v in s.items():
                value = "not found" if _is_cli_english() and str(v) == "未找到" else str(v)
                table.add_row(k, value)
            CONSOLE.print(table)
            CONSOLE.print()
            return True

        if head == "/probes":
            probes = stm32_list_probes()
            if probes["probes"]:
                for p in probes["probes"]:
                    CONSOLE.print(f"  [{THEME}]{p['description']}[/] ({p['uid']})")
            else:
                CONSOLE.print(
                    f"[yellow]{_cli_text('未检测到任何探针，请检查 USB 连接', 'No probes detected. Check the USB connection.')}[/]"
                )
            CONSOLE.print()
            return True

        if head == "/projects":
            r = stm32_list_projects()
            if r["projects"]:
                table = Table(title=_cli_text("历史项目", "Project History"), box=box.SIMPLE)
                table.add_column(_cli_text("项目名", "Project"), style=f"bold {THEME}")
                table.add_column(_cli_text("芯片", "Chip"), style="cyan")
                table.add_column(_cli_text("描述", "Description"), style="white")
                for p in r["projects"]:
                    table.add_row(p["name"], p["chip"], p["request"][:40])
                CONSOLE.print(table)
            else:
                CONSOLE.print(f"[dim]{_cli_text('暂无历史项目', 'No project history yet')}[/]")
            CONSOLE.print()
            return True

        if head == "/member":
            subcmd = arg.strip().lower()
            if subcmd == "path":
                path = _ensure_member_file()
                CONSOLE.print(
                    f"[{THEME}]{_cli_text('member.md 路径', 'member.md path')}: {path}[/]\n"
                )
                return True
            if subcmd == "reload":
                self.refresh_system_prompt()
                CONSOLE.print(
                    f"[green]{_cli_text('member.md 已重新加载到当前会话', 'member.md reloaded into the current session')}[/]\n"
                )
                return True
            if subcmd and subcmd not in {"path", "reload"}:
                CONSOLE.print(
                    f"[yellow]{_cli_text('用法: /member [path|reload]', 'Usage: /member [path|reload]')}[/]\n"
                )
                return True
            self.refresh_system_prompt()
            title = _cli_text("Gary 经验库", "Gary Memory")
            CONSOLE.print(
                Panel(
                    Markdown(_member_preview_markdown()),
                    title=f"[bold {THEME}]{title}[/]",
                    border_style=THEME,
                )
            )
            CONSOLE.print()
            return True

        if head == "/clear":
            self.reset_conversation()
            CONSOLE.clear()
            self._print_header()
            return True

        if head == "/config":
            configure_ai_cli(agent=self)
            return True

        if head in ("/exit", "/quit"):
            CONSOLE.print(
                f"\n[{THEME}]{_cli_text('正在退出，清理硬件和 Telegram...', 'Exiting, cleaning up hardware and Telegram...')}[/]"
            )
            shutdown = _shutdown_cli_runtime(stop_telegram=True)
            tg = shutdown.get("telegram", {})
            if tg.get("message"):
                color = "green" if tg.get("success") else "yellow"
                CONSOLE.print(f"[{color}]{tg['message']}[/]")
            CONSOLE.print(_cli_text("再见！", "Goodbye!"))
            sys.exit(0)

        return False

    # ── UI ──────────────────────────────────────────────────
    def _print_header(self):
        chip_line = _cli_text(
            f"芯片: [bold]{_current_chip}[/]  |  模型: [bold]{AI_MODEL}[/]",
            f"Chip: [bold]{_current_chip}[/]  |  Model: [bold]{AI_MODEL}[/]",
        )
        hw_line = (
            _cli_text(
                f"硬件: [green]已连接[/]  串口: [{'green' if _serial_connected else 'yellow'}]"
                f"{'已连接' if _serial_connected else '未连接'}[/]",
                f"Hardware: [green]connected[/]  Serial: [{'green' if _serial_connected else 'yellow'}]"
                f"{'connected' if _serial_connected else 'disconnected'}[/]",
            )
            if _hw_connected
            else _cli_text("硬件: [dim]未连接[/]", "Hardware: [dim]disconnected[/]")
        )
        art = (
            "   ██████╗  █████╗ ██████╗ ██╗   ██╗\n"
            "  ██╔════╝ ██╔══██╗██╔══██╗╚██╗ ██╔╝\n"
            "  ██║  ███╗███████║██████╔╝ ╚████╔╝ \n"
            "  ██║   ██║██╔══██║██╔══██╗  ╚██╔╝  \n"
            "  ╚██████╔╝██║  ██║██║  ██║   ██║   \n"
            "   ╚═════╝ ╚═╝  ╚═╝╚═╝  ╚═╝   ╚═╝  "
        )
        panel = Panel(
            f"[bold {THEME}]{art}[/]\n\n"
            f"  {chip_line}\n  {hw_line}\n\n"
            f"  [dim]{_cli_text('输入需求即可生成代码 · Tab 补全命令 · /help 查看命令 · /connect 连接硬件', 'Describe what you want to build · Tab completes commands · /help shows commands · /connect attaches hardware')}[/]",
            title=f"[bold {THEME}]Gary Dev Agent[/]",
            border_style=THEME,
            padding=(0, 1),
        )
        CONSOLE.print(panel)
        CONSOLE.print()

    def _show_help(self):
        table = Table(title=_cli_text("内置命令", "Built-in Commands"), box=box.SIMPLE)
        table.add_column(_cli_text("命令", "Command"), style=f"bold {THEME}")
        table.add_column(_cli_text("说明", "Description"), style="white")
        cmds = [
            (
                "/connect [chip]",
                _cli_text(
                    "连接探针（如 /connect STM32F103C8T6）",
                    "Connect the probe (for example: /connect STM32F103C8T6)",
                ),
            ),
            (
                "/serial [port] [baud]",
                _cli_text(
                    "连接串口（如 /serial /dev/ttyUSB0 115200）",
                    "Connect serial (for example: /serial /dev/ttyUSB0 115200)",
                ),
            ),
            ("/disconnect", _cli_text("断开探针和串口", "Disconnect probe and serial")),
            ("/chip [model]", _cli_text("查看/切换芯片型号", "Show or change the chip model")),
            (
                "/language [en|zh]",
                _cli_text(
                    "切换 CLI 语言，默认一键切到英文",
                    "Switch CLI language. `/language` switches to English immediately",
                ),
            ),
            ("/probes", _cli_text("列出所有可用探针", "List all available probes")),
            ("/status", _cli_text("查看硬件+工具链状态", "Show hardware and toolchain status")),
            (
                "/config",
                _cli_text(
                    "配置 AI 接口（API Key / Model / Base URL）",
                    "Configure AI settings (API Key / Model / Base URL)",
                ),
            ),
            ("/projects", _cli_text("列出历史项目", "List saved projects")),
            (
                "/member [path|reload]",
                _cli_text(
                    "查看经验库；`path` 显示路径；`reload` 重新载入 member.md",
                    "View memory; `path` shows the file location; `reload` refreshes member.md",
                ),
            ),
            (
                "/telegram [subcommand]",
                _cli_text(
                    "Telegram 机器人管理: start/stop/status/allow/remove/reset",
                    "Manage the Telegram bot: start/stop/status/allow/remove/reset",
                ),
            ),
            ("/clear", _cli_text("清空对话历史", "Clear conversation history")),
            ("/exit", _cli_text("退出并停止 Telegram", "Exit and stop Telegram")),
            ("?", _cli_text("显示帮助", "Show help")),
            (
                "Tab",
                _cli_text(
                    "补全命令/子命令/芯片/串口/技能名，历史输入会自动预测",
                    "Complete commands, subcommands, chips, serial ports, and skills; history is suggested automatically",
                ),
            ),
            (
                "/skill [subcommand]",
                _cli_text(
                    "技能管理: list/install/enable/disable/create/export",
                    "Manage skills: list/install/enable/disable/create/export",
                ),
            ),
        ]
        for cmd, desc in cmds:
            table.add_row(cmd, desc)
        CONSOLE.print(table)
        CONSOLE.print()

    def _status_bar(self):
        tokens = self._tokens()
        hw = (
            f"[green]●[/] {_current_chip}"
            if _hw_connected
            else _cli_text("[dim]○ 未连接[/]", "[dim]○ Disconnected[/]")
        )
        CONSOLE.print(
            f"[dim]{hw}  │  {AI_MODEL}  │  {_cli_text('上下文', 'context')}: ~{tokens} tokens[/]"
        )
        CONSOLE.rule(style="dim")

    # ── 主循环 ──────────────────────────────────────────────
    def run(self):
        if self.session is None:
            raise RuntimeError("当前实例未启用交互式 PromptSession")
        telegram_startup = _ensure_cli_telegram_daemon()
        CONSOLE.clear()
        self._print_header()
        if telegram_startup and telegram_startup.get("message"):
            color = "green" if telegram_startup.get("success") else "yellow"
            CONSOLE.print(f"[{color}]  Telegram: {telegram_startup.get('message', '')}[/]")
            CONSOLE.print()
        pt_style = Style.from_dict({"prompt": f"bold {THEME}"})

        while True:
            try:
                self._status_bar()
                user_input = self.session.prompt(
                    HTML(f'<style color="cyan"><b>Gary > </b></style>'),
                    style=pt_style,
                )
                if not user_input.strip():
                    continue

                if user_input.startswith("/") or user_input.strip() == "?":
                    self.handle_builtin(user_input.strip())
                    continue

                self.chat(user_input)

            except KeyboardInterrupt:
                CONSOLE.print(
                    f"\n[dim]{_cli_text('Ctrl+C 中断。/exit 退出。', 'Interrupted by Ctrl+C. Use /exit to quit.')}[/]"
                )
            except EOFError:
                _shutdown_cli_runtime(stop_telegram=True)
                break


# ─────────────────────────────────────────────────────────────
# 入口
# ─────────────────────────────────────────────────────────────
def main():
    global _current_chip

    # 命令行参数
    args = sys.argv[1:]

    # ── Telegram 管理模式：gary telegram ... ─────────────────
    if args and args[0].lower() == "telegram":
        handle_telegram_command(" ".join(args[1:]), source="cli")
        sys.exit(0)
    if "--telegram" in args:
        idx = args.index("--telegram")
        handle_telegram_command(" ".join(args[idx + 1 :]), source="cli")
        sys.exit(0)

    # ── 诊断模式：Gary doctor ────────────────────────────────
    if "--doctor" in args:
        run_doctor()
        sys.exit(0)

    # ── 配置模式：Gary config ────────────────────────────────
    if "--config" in args:
        configure_ai_cli()
        sys.exit(0)

    if "--chip" in args:
        idx = args.index("--chip")
        if idx + 1 < len(args):
            _current_chip = args[idx + 1].upper()

    # 检查依赖
    CONSOLE.print(f"[dim]{_cli_text('检查环境...', 'Checking environment...')}[/]")

    # GCC
    compiler = _get_compiler()
    ci = compiler.check(_current_chip)
    if ci.get("gcc"):
        CONSOLE.print(f"[green]  GCC: {ci['gcc_version']}[/]")
    else:
        CONSOLE.print(
            f"[yellow]  GCC: {_cli_text('未找到 arm-none-eabi-gcc', 'arm-none-eabi-gcc not found')}[/]"
        )

    if ci.get("hal"):
        CONSOLE.print(f"[green]  HAL: {_cli_text('已就绪', 'ready')}[/]")
    else:
        CONSOLE.print(
            f"[yellow]  HAL: {_cli_text('未找到，请运行 setup.sh', 'not found, run setup.sh')}[/]"
        )

    # pyocd
    try:
        import pyocd

        CONSOLE.print(f"[green]  pyocd: {pyocd.__version__}[/]")
    except ImportError:
        CONSOLE.print(
            f"[yellow]  pyocd: {_cli_text('未安装（pip install pyocd）', 'not installed (pip install pyocd)')}[/]"
        )

    # 串口自动扫描
    import platform as _platform, glob as _glob

    serial_candidates = detect_serial_ports(verbose=False)
    if serial_candidates:
        CONSOLE.print(
            f"[green]  {_cli_text('串口', 'Serial')}: {_cli_text(f'检测到 {serial_candidates}（连接时自动选择）', f'detected {serial_candidates} (auto-selected on connect)')}[/]"
        )
    else:
        _plat = _platform.system()
        # 扫不到 → 检查是否权限问题，给出平台相关修复命令
        if _plat == "Linux":
            import grp as _grp

            no_perm = [
                p
                for pat in ["/dev/ttyUSB*", "/dev/ttyACM*", "/dev/ttyS[4-9]*"]
                for p in _glob.glob(pat)
                if os.path.exists(p) and not os.access(p, os.R_OK | os.W_OK)
            ]
            if no_perm:
                try:
                    grp_name = _grp.getgrgid(os.stat(no_perm[0]).st_gid).gr_name
                except Exception:
                    grp_name = "dialout"
                CONSOLE.print(
                    f"[yellow]  {_cli_text('串口', 'Serial')}: {_cli_text(f'发现 {no_perm} 但权限不足', f'found {no_perm} but permissions are insufficient')}[/]"
                )
                CONSOLE.print(
                    f"[yellow]    → sudo usermod -aG {grp_name} $USER && newgrp {grp_name}[/]"
                )
            else:
                CONSOLE.print(
                    f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到串口设备（连接硬件后重试）', 'no serial device detected (connect hardware and retry)')}[/]"
                )
        elif _plat == "Darwin":
            CONSOLE.print(
                f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到串口设备', 'no serial device detected')}[/]"
            )
            CONSOLE.print(
                f"[dim]    {_cli_text('插上 USB 转串口后重启程序，或运行 /serial list 查看', 'Reconnect your USB serial adapter and restart, or run /serial list')}[/]"
            )
        elif _plat == "Windows":
            CONSOLE.print(
                f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到 COM 口', 'no COM ports detected')}[/]"
            )
            CONSOLE.print(
                f"[dim]    {_cli_text('请在设备管理器确认驱动已安装（CH340/CP210x），或运行 /serial list', 'Check the driver in Device Manager (CH340/CP210x), or run /serial list')}[/]"
            )
        else:
            CONSOLE.print(
                f"[dim]  {_cli_text('串口', 'Serial')}: {_cli_text('未检测到串口设备（连接硬件后重试）', 'no serial device detected (connect hardware and retry)')}[/]"
            )

    CONSOLE.print()

    # 单次执行模式：Gary do "任务"  →  python stm32_agent.py --do "任务"
    if "--do" in args:
        idx = args.index("--do")
        task = args[idx + 1] if idx + 1 < len(args) else ""
        if not task:
            CONSOLE.print(
                f"[red]{_cli_text('--do 后需要任务描述，例如: Gary do \"让 PA0 LED 闪烁\"', '--do requires a task description, for example: Gary do \"blink PA0 LED\"')}[/]"
            )
            sys.exit(1)
        if "--connect" in args:
            chip_arg = None
            if "--chip" in args:
                ci_idx = args.index("--chip")
                if ci_idx + 1 < len(args):
                    chip_arg = args[ci_idx + 1]
            stm32_connect(chip_arg)
        agent = STM32Agent()
        CONSOLE.print(f"\n[cyan]  ▶ Gary do: {task}[/]\n")
        agent.chat(task)
        stm32_disconnect()
        sys.exit(0)

    # 自动连接
    if "--connect" in args:
        chip_arg = None
        if "--chip" in args:
            idx = args.index("--chip")
            if idx + 1 < len(args):
                chip_arg = args[idx + 1]
        stm32_connect(chip_arg)

    agent = STM32Agent()
    agent.run()


if __name__ == "__main__":
    main()
